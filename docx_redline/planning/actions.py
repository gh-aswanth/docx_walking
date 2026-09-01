# This Source Code Form is subject to the terms of the Mozilla Public
# License, v. 2.0. If a copy of the MPL was not distributed with this
# file, You can obtain one at https://mozilla.org/MPL/2.0/.
#
# Copyright (c) 2026 Aswanth B S

"""Action items: the vocabulary an LLM emits, and the planner that applies them.

An *action item* is a review decision ("move the governing-law clause up into
the term section"), not an XML edit.  The planner turns each one into tracked
changes and -- crucially -- works out the **consequential** edits the model
never mentioned: when a clause moves, its old siblings close ranks, its new
siblings shift down, and every cross-reference in the document follows.

Pipeline order matters and is fixed:

    resolve  -> content edits -> structural edits -> renumber -> references

Content edits run first so they address the document by its *original* clause
numbers.  Renumbering runs last so it sees the final order of the paragraphs.
"""

from __future__ import annotations

import re
from collections.abc import Iterable, Sequence
from dataclasses import dataclass, field

from lxml import etree

from ..editing.redline import RedlineError, Redliner
from ..oxml import edits
from ..oxml.elements import deepcopy_without, make, make_run
from ..oxml.ns import qn
from ..oxml.textmap import paragraph_text
from ..structure.clauses import (
    Clause,
    ClauseError,
    ClauseTree,
    ReferenceEdit,
    Renumbering,
    collect_reference_edits,
    parse_clause,
)

# ---------------------------------------------------------------------------
# schema
# ---------------------------------------------------------------------------

#: action type -> (required keys, optional keys).  This doubles as the contract
#: handed to the model, so keep the names self-describing.
# fmt: off
ACTION_SCHEMA: dict[str, tuple[tuple[str, ...], tuple[str, ...]]] = {
    # -- text level, no structural consequence -------------------------
    "replace_text": (("find", "replace"), ("clause", "all", "regex")),
    "insert_text": (("text",), ("clause", "anchor", "position")),
    "delete_text": (("find",), ("clause", "all", "regex")),
    "rewrite_clause": (("clause", "text"), ()),
    # -- structural, triggers renumbering ------------------------------
    "insert_clause": (("text",), ("after_clause", "before_clause", "into_section", "title")),
    "delete_clause": (("clause",), ()),
    "move_clause": (("clause",), ("after_clause", "before_clause", "into_section", "position")),
    "reorder_clauses": (("section", "order"), ()),
    "insert_section": (("title",), ("after_section", "text")),
    "delete_section": (("section",), ()),
    "move_section": (("section",), ("after_section", "before_section", "position")),
    # -- tables ---------------------------------------------------------
    "insert_row": (("table", "values"), ("row",)),
    "delete_row": (("table", "row"), ()),
    "update_cell": (("table", "row", "col", "text"), ()),
    # -- presentation & annotation --------------------------------------
    "format_text": (("find",), ("clause", "bold", "italic", "underline", "color", "highlight")),
    "format_clause": (("clause",), ("alignment", "space_after", "space_before", "style")),
    "comment": (("text",), ("clause", "find")),
}
# fmt: on

#: Actions that change the document's structure and therefore its numbering.
STRUCTURAL_ACTIONS = frozenset(
    {
        "insert_clause",
        "delete_clause",
        "move_clause",
        "reorder_clauses",
        "insert_section",
        "delete_section",
        "move_section",
    }
)

#: Annotations run last, after every edit, so they land on the finished text.
ANNOTATION_ACTIONS = frozenset({"comment"})

#: Actions the planner derives itself.  A model must never emit these -- they
#: are consequences, and letting the model guess them is how numbering drifts.
DERIVED_ACTIONS = frozenset({"renumber_clause", "update_cross_reference"})

SEVERITIES = ("low", "medium", "high", "critical")


def allowed_types(segment) -> frozenset[str]:
    """Which action types a segment can support.

    Structural actions resolve through clause numbers. A segment cut on headings
    or on fixed windows has none, so proposing a move inside one is a finding the
    planner could never apply -- better to keep it out of the schema the model
    is given than to reject it afterwards.
    """
    if segment.strategy == "clauses":
        return frozenset(ACTION_SCHEMA)
    return frozenset(ACTION_SCHEMA) - STRUCTURAL_ACTIONS


@dataclass
class ActionResult:
    """What actually happened for one action item."""

    id: str
    type: str
    status: str  # applied | skipped | failed
    detail: str = ""
    rationale: str = ""
    severity: str = ""
    edits: int = 0
    derived: list[dict] = field(default_factory=list)
    #: paragraphs the action touched, used to anchor its rationale in the doc
    anchors: list = field(default_factory=list, repr=False)

    def to_dict(self) -> dict:
        out = {
            "id": self.id,
            "type": self.type,
            "status": self.status,
            "detail": self.detail,
            "edits": self.edits,
        }
        if self.rationale:
            out["rationale"] = self.rationale
        if self.severity:
            out["severity"] = self.severity
        if self.derived:
            out["derived"] = self.derived
        return out


@dataclass
class PlanReport:
    """The whole run, serialisable straight to JSON."""

    results: list[ActionResult] = field(default_factory=list)
    renumbered: list[dict] = field(default_factory=list)
    references: list[dict] = field(default_factory=list)
    dangling_references: list[dict] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)

    @property
    def applied(self) -> int:
        return sum(1 for r in self.results if r.status == "applied")

    @property
    def failed(self) -> int:
        return sum(1 for r in self.results if r.status == "failed")

    def to_dict(self) -> dict:
        return {
            "summary": {
                "actions": len(self.results),
                "applied": self.applied,
                "failed": self.failed,
                "clauses_renumbered": len(self.renumbered),
                "cross_references_updated": len(self.references),
                "dangling_references": len(self.dangling_references),
                "warnings": len(self.warnings),
            },
            "actions": [r.to_dict() for r in self.results],
            "renumbering": self.renumbered,
            "cross_references": self.references,
            "dangling_references": self.dangling_references,
            "warnings": self.warnings,
        }

    def format(self) -> str:
        lines = [
            f"actions: {self.applied} applied, {self.failed} failed, "
            f"{len(self.results) - self.applied - self.failed} skipped",
            f"clauses renumbered: {len(self.renumbered)}",
            f"cross-references updated: {len(self.references)}",
        ]
        if self.dangling_references:
            lines.append(f"dangling references: {len(self.dangling_references)}")
        for warning in self.warnings:
            lines.append(f"  warning: {warning}")
        return "\n".join(lines)


def validate_actions(items: Iterable[dict]) -> list[str]:
    """Schema check with no document access -- safe to run on raw model output."""
    problems: list[str] = []
    seen_ids: set[str] = set()
    for index, item in enumerate(items):
        where = item.get("id") or f"#{index}"
        kind = item.get("type")
        if kind in DERIVED_ACTIONS:
            problems.append(f"{where}: {kind!r} is derived by the planner and must not be supplied")
            continue
        if kind not in ACTION_SCHEMA:
            problems.append(f"{where}: unknown action type {kind!r}")
            continue
        required, optional = ACTION_SCHEMA[kind]
        for key in required:
            if item.get(key) in (None, ""):
                problems.append(f"{where}: {kind!r} requires {key!r}")
        known = set(required) | set(optional) | {"id", "type", "rationale", "severity", "note"}
        for key in item:
            if key not in known:
                problems.append(f"{where}: {kind!r} does not accept {key!r}")
        severity = item.get("severity")
        if severity is not None and severity not in SEVERITIES:
            problems.append(f"{where}: severity {severity!r} not in {SEVERITIES}")
        if item.get("id"):
            if item["id"] in seen_ids:
                problems.append(f"{where}: duplicate action id")
            seen_ids.add(item["id"])
        if kind == "move_clause" and not (
            item.get("after_clause") or item.get("before_clause") or item.get("into_section")
        ):
            problems.append(
                f"{where}: move_clause needs after_clause, before_clause or into_section"
            )
        if kind == "move_section" and not (
            item.get("after_section") or item.get("before_section") or item.get("position")
        ):
            problems.append(
                f"{where}: move_section needs after_section, before_section or position"
            )
        if kind in ("move_clause", "move_section") and item.get("position") not in (
            None,
            "first",
            "last",
        ):
            problems.append(
                f"{where}: position must be 'first' or 'last', not {item['position']!r}"
            )
        if kind == "reorder_clauses":
            order = item.get("order")
            if not isinstance(order, (list, tuple)) or len(order) < 2:
                problems.append(
                    f"{where}: reorder_clauses needs an 'order' list of two or more clause numbers"
                )
            elif len(set(map(str, order))) != len(order):
                problems.append(f"{where}: reorder_clauses 'order' repeats a clause number")
        if kind == "insert_clause" and not (
            item.get("after_clause") or item.get("before_clause") or item.get("into_section")
        ):
            problems.append(
                f"{where}: insert_clause needs after_clause, before_clause or into_section"
            )
    return problems


# ---------------------------------------------------------------------------
# planner
# ---------------------------------------------------------------------------
class ActionPlanner:
    """Applies action items to a :class:`Redliner`, renumbering as it goes."""

    def __init__(
        self,
        rl: Redliner,
        renumber: bool = True,
        strict: bool = False,
        explain: bool = False,
    ) -> None:
        self.rl = rl
        self.renumber_enabled = renumber
        self.strict = strict
        #: attach each action's rationale to the document as a Word comment
        self.explain = explain
        self._anchors: list[etree._Element] = []
        self.body = rl.document.element.body
        self.tree = ClauseTree(self.body)
        #: what every clause was numbered before anything moved
        self.baseline = self.tree.snapshot()
        self.report = PlanReport()
        self._deleted_labels: set[str] = set()

    # -- entry point ----------------------------------------------------
    def run(self, items: Sequence[dict]) -> PlanReport:
        problems = validate_actions(items)
        if problems:
            raise RedlineError("invalid action items:\n  " + "\n  ".join(problems))

        content = [
            i
            for i in items
            if i["type"] not in STRUCTURAL_ACTIONS and i["type"] not in ANNOTATION_ACTIONS
        ]
        structural = [i for i in items if i["type"] in STRUCTURAL_ACTIONS]
        annotations = [i for i in items if i["type"] in ANNOTATION_ACTIONS]

        # Content first: it addresses clauses by the numbers the model saw.
        for item in content:
            self._run_one(item)
        for item in structural:
            self._run_one(item)
        # Annotations last: a comment should point at the finished paragraph,
        # not at text a later action moved or struck.
        for item in annotations:
            self._run_one(item)

        if self.renumber_enabled:
            self.finalize()
        return self.report

    def annotate_rationales(self) -> int:
        """Write each applied action's rationale into the document as a comment.

        Without this the *why* only ever reaches the JSON report: a reviewer
        opening the .docx sees the tracked change but not the reason for it.
        Anchoring on the paragraph the action actually touched puts the
        explanation beside the change in Word's review pane.

        Runs last, after renumbering, so a comment never points at a paragraph
        a later stage moved, struck or renumbered out from under it.
        """
        attached = 0
        for result in self.report.results:
            if result.status != "applied" or not result.rationale:
                continue
            if result.type in ANNOTATION_ACTIONS:
                continue  # a comment explaining a comment helps nobody
            anchor = self._live_anchor(result.anchors)
            if anchor is None:
                self.report.warnings.append(
                    f"{result.id}: no paragraph left to attach the rationale to"
                )
                continue
            label = f"[{result.id}"
            if result.severity:
                label += f" \u00b7 {result.severity}"
            label += "] "
            try:
                self.rl.add_comment(_wrap(self.rl, anchor), label + result.rationale)
            except RedlineError as exc:  # pragma: no cover - defensive
                self.report.warnings.append(f"{result.id}: could not attach rationale ({exc})")
                continue
            attached += 1
        return attached

    @staticmethod
    def _live_anchor(anchors) -> etree._Element | None:
        """First recorded paragraph still in the tree and still *visibly* there.

        ``p.iter(w:r)`` also matches runs inside a ``w:del``, so a struck
        paragraph looks alive and the rationale lands on text the redline is
        removing.  Fall back to that only when nothing visible remains -- a
        deletion still deserves its explanation.
        """
        from ..oxml.textmap import iter_visible_runs

        fallback: etree._Element | None = None
        for p in anchors:
            if p is None or p.getparent() is None:
                continue
            if next(iter_visible_runs(p), None) is not None:
                return p
            if fallback is None and next(p.iter(qn("w:r")), None) is not None:
                fallback = p
        return fallback

    def finalize(self) -> PlanReport:
        """Run the consequential pass: renumber, then repoint cross-references.

        Called automatically by :meth:`run` unless renumbering is switched off.
        Callers that layer other edits into the same document -- a document
        compare, say -- construct the planner *first* (which captures the
        baseline numbering), do their work, and call this at the very end so
        one pass sees the final paragraph order.
        """
        self._renumber()
        if self.explain:
            self.annotate_rationales()
        return self.report

    def _run_one(self, item: dict) -> None:
        kind = item["type"]
        identifier = item.get("id") or kind
        handler = getattr(self, f"_do_{kind}")
        self._anchors = []
        try:
            detail, count = handler(item)
        except (ClauseError, RedlineError, KeyError, IndexError) as exc:
            if self.strict:
                raise
            self.report.results.append(ActionResult(identifier, kind, "failed", str(exc)))
            self.report.warnings.append(f"{identifier}: {exc}")
            return
        status = "applied" if count else "skipped"
        self.report.results.append(
            ActionResult(
                id=identifier,
                type=kind,
                status=status,
                detail=detail,
                rationale=item.get("rationale", ""),
                severity=item.get("severity", ""),
                edits=count,
                anchors=self._anchors,
            )
        )

    def _touch(self, p: etree._Element) -> etree._Element:
        """Note a paragraph an action edited, so its rationale can land there."""
        if p is not None and all(p is not seen for seen in self._anchors):
            self._anchors.append(p)
        return p

    # -- resolution -----------------------------------------------------
    def _clause(self, item: dict, key: str = "clause") -> Clause:
        label = item.get(key)
        if not label:
            raise RedlineError(f"action {item.get('id')} needs {key!r}")
        return self.tree.get(str(label))

    def _require_unique(self, item, pattern, needle: str, scope) -> None:
        """Refuse an unscoped edit whose text appears in more than one paragraph.

        Without a ``clause``, ``_scope`` widens to the entire body and the first
        match wins -- so a phrase quoted from an exhibit on page 240 silently
        edits page 1 instead, and reports ``applied``. Whole-document coverage
        means models now quote unnumbered text routinely, so ambiguity has to be
        an error rather than a coin toss.
        """
        hits = [p for p in scope if pattern.search(paragraph_text(p))]
        if len(hits) <= 1:
            return
        where = "; ".join(paragraph_text(p)[:48].strip() for p in hits[:3])
        raise RedlineError(
            f"{needle!r} appears in {len(hits)} paragraphs -- add a clause, quote more "
            f"surrounding text, or set all=true. Matches: {where}"
        )

    def _scope(self, item: dict) -> list[etree._Element]:
        """Paragraphs an action may touch: one clause, or the whole document."""
        label = item.get("clause")
        if not label:
            return list(self.body.iter(qn("w:p")))
        clause = self.tree.get(str(label))
        return [clause.p] + [child.p for child in clause.walk() if child is not clause]

    # -- content actions -------------------------------------------------
    def _do_replace_text(self, item) -> tuple[str, int]:
        return self._text_edit(item, item["find"], item["replace"])

    def _do_delete_text(self, item) -> tuple[str, int]:
        return self._text_edit(item, item["find"], "")

    def _text_edit(self, item, find: str, replace: str) -> tuple[str, int]:
        regex = bool(item.get("regex"))
        limit = None if item.get("all") else 1
        pattern = re.compile(find if regex else re.escape(find))
        scope = self._scope(item)
        if limit and not item.get("clause"):
            self._require_unique(item, pattern, find, scope)
        count = 0
        for p in scope:
            text = paragraph_text(p)
            spans = [(m.start(), m.end()) for m in pattern.finditer(text)]
            for start, end in reversed(spans):
                edits.replace_range(p, start, end, replace, self.rl.ctx)
                self._touch(p)
                count += 1
                if limit and count >= limit:
                    break
            if limit and count >= limit:
                break
        if not count:
            raise RedlineError(f"text {find!r} not found in scope")
        verb = "deleted" if not replace else "replaced"
        return f"{verb} {find!r}" + (f" -> {replace!r}" if replace else ""), count

    def _do_insert_text(self, item) -> tuple[str, int]:
        anchor = item.get("anchor")
        position = item.get("position", "after")
        text = item["text"]
        scope = self._scope(item)
        if not item.get("clause"):
            if not anchor:
                raise RedlineError(
                    "insert_text without a clause needs an anchor: otherwise it appends to "
                    "the first paragraph of the document"
                )
            self._require_unique(item, re.compile(re.escape(anchor)), anchor, scope)
        for p in scope:
            body = paragraph_text(p)
            if anchor:
                index = body.find(anchor)
                if index < 0:
                    continue
                offset = index if position == "before" else index + len(anchor)
            elif position == "start":
                offset = 0
            else:
                offset = len(body)
            edits.insert_text(p, offset, text, self.rl.ctx)
            self._touch(p)
            return f"inserted {text!r} {position} {anchor or 'clause'}", 1
        raise RedlineError(f"anchor {anchor!r} not found in scope")

    def _do_rewrite_clause(self, item) -> tuple[str, int]:
        clause = self._clause(item)
        # Keep the clause number: the model rewrites the body, the planner owns
        # the numbering.
        prefix = clause.text[: clause.span[1]]
        keep = clause.text[clause.span[1] :]
        lead = keep[: len(keep) - len(keep.lstrip())]
        new_text = prefix + lead + item["text"].strip()
        count = self.rl.set_paragraph_text(_wrap(self.rl, clause.p), new_text)
        self._touch(clause.p)
        return f"rewrote clause {clause.label}", count or 1

    def _do_format_text(self, item) -> tuple[str, int]:
        props = {
            k: v
            for k, v in item.items()
            if k in ("bold", "italic", "underline", "color", "highlight")
        }
        count = 0
        for p in self._scope(item):
            text = paragraph_text(p)
            index = text.find(item["find"])
            if index < 0:
                continue
            from docx.text.run import Run

            from ..oxml.textmap import ParagraphText, split_at

            split_at(p, index + len(item["find"]))
            split_at(p, index)
            view = ParagraphText(p)
            paragraph = _wrap(self.rl, p)
            runs = [Run(r, paragraph) for r in view.runs_in(index, index + len(item["find"]))]
            count += self.rl.format_runs(runs, **props)
            self._touch(p)
            break
        if not count:
            raise RedlineError(f"text {item['find']!r} not found for formatting")
        return f"formatted {item['find']!r}", count

    def _do_format_clause(self, item) -> tuple[str, int]:
        clause = self._clause(item)
        props = {
            k: v
            for k, v in item.items()
            if k in ("alignment", "space_after", "space_before", "style")
        }
        from ..editing.ops import _format_props

        self.rl.format_paragraph(_wrap(self.rl, clause.p), **_format_props(props, exclude=set()))
        self._touch(clause.p)
        return f"reformatted clause {clause.label}", 1

    def _do_comment(self, item) -> tuple[str, int]:
        """Anchor a review comment to a clause, or to a phrase inside one."""
        if item.get("find"):
            for p in self._scope(item):
                text = paragraph_text(p)
                index = text.find(item["find"])
                if index < 0:
                    continue
                from docx.text.run import Run

                from ..oxml.textmap import ParagraphText, split_at

                split_at(p, index + len(item["find"]))
                split_at(p, index)
                paragraph = _wrap(self.rl, p)
                runs = [
                    Run(r, paragraph)
                    for r in ParagraphText(p).runs_in(index, index + len(item["find"]))
                ]
                self.rl.add_comment(paragraph, item["text"], runs=runs)
                return f"commented on {item['find']!r}", 1
            raise RedlineError(f"cannot anchor a comment: {item['find']!r} not found")
        clause = self._clause(item)
        self.rl.add_comment(_wrap(self.rl, clause.p), item["text"])
        return f"commented on clause {clause.label}", 1

    # -- table actions ---------------------------------------------------
    def _do_insert_row(self, item) -> tuple[str, int]:
        table = self.rl.tables()[int(item["table"])]
        row = self.rl.insert_table_row(table, index=item.get("row"), values=item["values"])
        self._touch(next(row._tr.iter(qn("w:p")), None))
        return f"inserted row into table {item['table']}", 1

    def _do_delete_row(self, item) -> tuple[str, int]:
        table = self.rl.tables()[int(item["table"])]
        self.rl.delete_table_row(table, int(item["row"]))
        self._touch(next(table._tbl.findall(qn("w:tr"))[int(item["row"])].iter(qn("w:p")), None))
        return f"deleted row {item['row']} of table {item['table']}", 1

    def _do_update_cell(self, item) -> tuple[str, int]:
        table = self.rl.tables()[int(item["table"])]
        cell = table.cell(int(item["row"]), int(item["col"]))
        count = self.rl.set_cell_text(cell, item["text"])
        self._touch(next(cell._tc.iter(qn("w:p")), None))
        return f"updated cell ({item['row']},{item['col']}) of table {item['table']}", count or 1

    # -- structural actions ----------------------------------------------
    def _do_insert_clause(self, item) -> tuple[str, int]:
        anchor, mode = self._anchor_for_insert(item)
        provisional = self._provisional_label(anchor, mode)
        new_p = self._build_clause_paragraph(anchor, provisional, item.get("title"), item["text"])
        if mode == "before":
            anchor.p.addprevious(new_p)
        else:
            _last_paragraph_of(anchor).addnext(new_p)
        edits.mark_paragraph_inserted(new_p, self.rl.ctx)
        self._touch(new_p)
        self.tree.reload()
        return f"inserted new clause {provisional} ({item.get('title') or 'untitled'})", 1

    def _anchor_for_insert(self, item) -> tuple[Clause, str]:
        if item.get("after_clause"):
            return self.tree.get(str(item["after_clause"])), "after"
        if item.get("before_clause"):
            return self.tree.get(str(item["before_clause"])), "before"
        section = self.tree.get(str(item["into_section"]))
        if section.children:
            return section.children[-1], "after"
        return section, "after"

    def _provisional_label(self, anchor: Clause, mode: str) -> str:
        """A plausible number for a new clause; the renumber pass has the last word."""
        number = list(anchor.number)
        if anchor.level == 1 and mode == "after" and not anchor.children:
            return f"{number[0]}.1"
        number[-1] += 0 if mode == "before" else 1
        return ".".join(str(part) for part in number)

    def _build_clause_paragraph(
        self, template: Clause, label: str, title: str | None, text: str
    ) -> etree._Element:
        """Clone the look of an existing clause: bold number run + body run."""
        new_p = make("w:p")
        ppr = template.p.find(qn("w:pPr"))
        if ppr is not None:
            new_p.append(deepcopy_without(ppr, ("w:rPr", "w:sectPr", "w:pPrChange")))
        runs = list(template.p.iter(qn("w:r")))
        number_rpr = runs[0].find(qn("w:rPr")) if runs else None
        body_rpr = runs[1].find(qn("w:rPr")) if len(runs) > 1 else number_rpr
        body = text.strip()
        if title and not body.lower().startswith(title.lower()):
            body = f"{title}. {body}"
        new_p.append(make_run(_label_prefix(label), number_rpr))
        new_p.append(make_run(body, body_rpr))
        return new_p

    def _do_delete_clause(self, item, kind: str = "clause") -> tuple[str, int]:
        clause = self._clause(item, "section" if kind == "section" else "clause")
        removed = list(clause.walk())

        # A delete takes the whole subtree it finds -- including anything an
        # earlier action moved or inserted into it.  Deleting a clause the
        # reviewer explicitly asked to *keep* is not a decision this layer gets
        # to make silently, and the two actions genuinely contradict.
        newcomers = [n for n in removed if n.p not in self.baseline]
        if newcomers:
            arrived = ", ".join(sorted(n.label for n in newcomers)[:6])
            raise RedlineError(
                f"cannot delete {kind} {clause.label}: {len(newcomers)} clause(s) were moved or "
                f"inserted into it earlier in this run ({arrived}). Deleting would discard them "
                "too -- drop one of the two actions."
            )

        for node in removed:
            self.rl.delete_paragraph(_wrap(self.rl, node.p))
            self._touch(node.p)
            self._deleted_labels.add(self.baseline.get(node.p, node.label))
        self.tree.reload()
        suffix = f" (+{len(removed) - 1} sub-clause(s))" if len(removed) > 1 else ""
        return f"deleted {kind} {clause.label}{suffix}", len(removed)

    def _do_delete_section(self, item) -> tuple[str, int]:
        return self._do_delete_clause(item, kind="section")

    def _do_move_clause(self, item) -> tuple[str, int]:
        clause = self._clause(item)
        subtree = list(clause.walk())
        after, before, where = self._move_target(item, subtree)
        old_label = clause.label

        cursor = after
        for node in subtree:
            clone = self._move_paragraph(
                node.p, after=cursor, before=before if cursor is None else None
            )
            cursor = clone
        self.tree.reload()
        extra = f" with {len(subtree) - 1} sub-clause(s)" if len(subtree) > 1 else ""
        return f"moved clause {old_label} {where}{extra}", len(subtree)

    def _move_target(
        self, item, subtree
    ) -> tuple[etree._Element | None, etree._Element | None, str]:
        """Resolve where a move lands: ``(after, before, human description)``.

        Exactly one of ``after``/``before`` comes back non-None; the caller
        threads its cursor through so a multi-paragraph clause stays together.
        """
        moving = {id(node.p) for node in subtree}

        if item.get("into_section"):
            section = self.tree.get(str(item["into_section"]))
            position = item.get("position", "first")
            others = [c for c in section.children if id(c.p) not in moving]
            if position == "first":
                if others:
                    return None, others[0].p, f"to the top of section {section.label}"
                return section.p, None, f"to the top of section {section.label}"
            anchor = others[-1] if others else section
            return _last_paragraph_of(anchor), None, f"to the end of section {section.label}"

        key = "after_clause" if item.get("after_clause") else "before_clause"
        anchor = self.tree.get(str(item[key]))
        if id(anchor.p) in moving:
            raise RedlineError(f"cannot move clause {item['clause']} relative to itself")
        if key == "after_clause":
            return _last_paragraph_of(anchor), None, f"after {anchor.label}"
        return None, anchor.p, f"before {anchor.label}"

    def _do_move_section(self, item) -> tuple[str, int]:
        section = self._clause(item, "section")
        subtree = list(section.walk())
        if item.get("after_section"):
            anchor = self.tree.get(str(item["after_section"]))
            after, before, where = _last_paragraph_of(anchor), None, f"after section {anchor.label}"
        elif item.get("before_section"):
            anchor = self.tree.get(str(item["before_section"]))
            after, before, where = None, anchor.p, f"before section {anchor.label}"
        elif item.get("position") == "first":
            others = [s for s in self.tree.sections if s.p is not section.p]
            after, before, where = None, others[0].p, "to the top of the document"
        else:
            others = [s for s in self.tree.sections if s.p is not section.p]
            after, before, where = (
                _last_paragraph_of(others[-1]),
                None,
                "to the end of the document",
            )
        if any(node.p is (before if before is not None else after) for node in subtree):
            raise RedlineError(f"cannot move section {section.label} relative to itself")

        cursor = after
        for node in subtree:
            clone = self._move_paragraph(
                node.p, after=cursor, before=before if cursor is None else None
            )
            cursor = clone
        self.tree.reload()
        return f"moved section {section.label} {where}", len(subtree)

    def _do_reorder_clauses(self, item) -> tuple[str, int]:
        """Rearrange a section's clauses into an explicit order.

        Only the clauses that genuinely have to move are moved: the longest
        run that is already in relative order stays put, so a reviewer sees one
        or two move revisions instead of the whole section struck and re-added.
        """
        section = self._clause(item, "section")
        children = section.children
        current = [clause.label for clause in children]
        requested, notes = self._reconcile_order(section, current, item["order"])
        target = _permute_in_place(current, requested)
        if target == current:
            return f"section {section.label} already in the requested order", 0

        # Work in slot positions, not labels.  Two children can legitimately
        # share a label mid-run (a freshly inserted clause borrows its
        # neighbour's number), and a label-keyed map would silently move one
        # clause twice and leave the other where it was.
        order_of = {label: index for index, label in enumerate(target)}
        slots = [order_of[label] for label in current]
        stay = set(_longest_increasing_run(slots))
        by_slot: dict[int, Clause] = {}
        for position, clause in enumerate(children):
            by_slot.setdefault(slots[position], clause)

        cursor = children[0].p.getprevious()
        if cursor is None:
            cursor = section.p
        moved = 0
        for wanted in range(len(target)):
            clause = by_slot.get(wanted)
            if clause is None:
                continue
            if any(slots[position] == wanted for position in stay):
                cursor = _last_paragraph_of(clause)
                continue
            for node in clause.walk():
                cursor = self._move_paragraph(node.p, after=cursor, before=None)
            moved += 1
        self._touch(section.p)
        self.tree.reload()
        detail = (
            f"reordered section {section.label}: {' -> '.join(target)} "
            f"({moved} clause(s) moved, {len(current) - moved} left in place)"
        )
        if notes:
            detail += " [" + "; ".join(notes) + "]"
        return detail, moved

    def _reconcile_order(
        self, section: Clause, current: list[str], order
    ) -> tuple[list[str], list[str]]:
        """Reduce a requested order to the clauses actually in the section now.

        An earlier action in the same run may have moved a clause out of this
        section, or added one the reviewer never saw. Failing the whole reorder
        for that would be unhelpful, so the listed-but-absent clauses are
        dropped with a note -- unless the number does not exist anywhere in the
        document, which means the reviewer invented it.
        """
        present = set(current)
        everywhere = {clause.label for clause in self.tree}
        requested: list[str] = []
        notes: list[str] = []
        for raw in order:
            label = str(raw)
            if label in present:
                requested.append(label)
            elif label in everywhere:
                notes.append(f"{label} has since moved out of section {section.label}, skipped")
            else:
                raise RedlineError(
                    f"section {section.label} has no clause {label}; it currently holds {current}"
                )
        unlisted = [label for label in current if label not in set(requested)]
        if unlisted:
            notes.append("left where they are: " + ", ".join(unlisted))
        for note in notes:
            self.report.warnings.append(f"reorder of section {section.label}: {note}")
        return requested, notes

    def _move_paragraph(
        self,
        source: etree._Element,
        after: etree._Element | None,
        before: etree._Element | None,
    ) -> etree._Element:
        """Relocate one paragraph, noting it when the move had to be downgraded."""
        from ..editing.redline import paragraph_has_revisions

        downgraded = paragraph_has_revisions(source)
        clone = self.rl.relocate_paragraph(source, after=after, before=before)
        self._touch(clone)
        if downgraded:
            self.report.warnings.append(
                "a clause edited earlier in this run was moved: recorded as "
                "delete + insert rather than a Word move revision, because the "
                "destination must not carry the source's strikeouts"
            )
        return clone

    def _do_insert_section(self, item) -> tuple[str, int]:
        if item.get("after_section"):
            anchor = self.tree.get(str(item["after_section"]))
        elif self.tree.sections:
            anchor = self.tree.sections[-1]
        else:
            raise RedlineError(
                "insert_section needs after_section: this document has no numbered sections"
            )
        provisional = str(anchor.number[0] + 1)
        heading = self._build_section_paragraph(anchor, provisional, item["title"])
        _last_paragraph_of(anchor).addnext(heading)
        edits.mark_paragraph_inserted(heading, self.rl.ctx)
        self._touch(heading)
        created = 1
        if item.get("text"):
            first = self._build_clause_paragraph(
                anchor.children[0] if anchor.children else anchor,
                f"{provisional}.1",
                None,
                item["text"],
            )
            heading.addnext(first)
            edits.mark_paragraph_inserted(first, self.rl.ctx)
            created += 1
        self.tree.reload()
        return f"inserted new section {provisional} ({item['title']})", created

    def _build_section_paragraph(self, template: Clause, label: str, title: str) -> etree._Element:
        new_p = make("w:p")
        ppr = template.p.find(qn("w:pPr"))
        if ppr is not None:
            new_p.append(deepcopy_without(ppr, ("w:rPr", "w:sectPr", "w:pPrChange")))
        runs = list(template.p.iter(qn("w:r")))
        new_p.append(make_run(f"{label}. {title}", runs[0].find(qn("w:rPr")) if runs else None))
        return new_p

    # -- the consequential pass ------------------------------------------
    def _renumber(self) -> None:
        """Re-derive every clause number from position, then fix the references.

        This is the step the model does not have to reason about.  It runs once,
        after all structural work, so a move that shifts two separate groups of
        siblings is handled in a single consistent pass.
        """
        self.tree.reload()
        for label in sorted(self.tree.duplicates):
            self.report.warnings.append(
                f"clause number {label} is carried by more than one paragraph; "
                "renumbering resolves it, but check the result"
            )
        result: Renumbering = self.tree.renumber(self.baseline)

        for clause, old_label, new_label in result.relabelled:
            self._rewrite_label(clause, old_label, new_label)
            self.report.renumbered.append(
                {"from": old_label, "to": new_label, "title": clause.title[:60]}
            )

        vanished = self._vanished_labels(result)
        self._assert_no_clause_vanished(vanished)

        # Cross-references are not confined to the body: headers, footers and
        # foot/endnotes cite sections too, and `Redliner` already edits them.
        removed = self._deleted_labels | vanished
        reference_edits: list[ReferenceEdit] = []
        dangling: list = []
        for root in self.rl._all_roots():
            found, missing = collect_reference_edits(root, result.mapping, removed=removed)
            reference_edits.extend(found)
            dangling.extend(missing)
        # Apply per paragraph, right to left, so earlier offsets stay valid.
        by_paragraph: dict[int, list[ReferenceEdit]] = {}
        for edit in reference_edits:
            by_paragraph.setdefault(id(edit.p), []).append(edit)
        for group in by_paragraph.values():
            for edit in sorted(group, key=lambda e: e.start, reverse=True):
                self._rewrite_reference(edit)
                self.report.references.append(
                    {
                        "from": edit.old,
                        "to": edit.new,
                        "context": edit.context,
                        "text": paragraph_text(edit.p)[:80],
                    }
                )
        for p, number in dangling:
            self.report.dangling_references.append(
                {"reference": number, "text": paragraph_text(p)[:100]}
            )
            self.report.warnings.append(
                f"reference to deleted clause {number} left untouched -- needs a human decision"
            )

    def _assert_no_clause_vanished(self, vanished: set[str]) -> None:
        """A clause may only disappear because something asked it to.

        Independent edits can combine to erase a clause nobody targeted -- most
        easily by moving one into a section that a later action deletes, since
        ``delete_clause`` removes the whole subtree it finds.  Everything needed
        to notice is already computed, so notice loudly.
        """
        unexplained = vanished - self._deleted_labels
        if not unexplained:
            return
        listed = ", ".join(sorted(unexplained)[:8])
        more = f" (+{len(unexplained) - 8} more)" if len(unexplained) > 8 else ""
        raise RedlineError(
            f"{len(unexplained)} clause(s) disappeared that no action deleted: {listed}{more}. "
            "This usually means one action moved a clause into a subtree another action removed."
        )

    def _vanished_labels(self, result: Renumbering) -> set[str]:
        """Baseline clauses no surviving clause answers to any more.

        ``delete_clause`` records its own targets, but a document compare can
        drop a clause too, and a reference to it is just as dangling.  A *move*
        must not count: the destination copy still carries the old label, so it
        shows up in ``claimed``.
        """
        baseline_labels = {label for _, label in self.baseline.items()}
        return baseline_labels - result.claimed

    def _rewrite_label(self, clause: Clause, old_label: str, new_label: str) -> None:
        """Rewrite a clause's own number, tracked unless the clause is itself new."""
        current = parse_clause(clause.p)
        if current is None:
            return
        start, end = current.span
        if clause.inserted or clause.moved:
            # The paragraph is already an insertion/move: a tracked edit inside
            # it would strike text the original document never had.
            edits.replace_in_place(clause.p, start, end, new_label)
        else:
            edits.replace_range(clause.p, start, end, new_label, self.rl.ctx)

    def _rewrite_reference(self, edit: ReferenceEdit) -> None:
        ppr = edit.p.find(qn("w:pPr"))
        rpr = ppr.find(qn("w:rPr")) if ppr is not None else None
        inside_revision = rpr is not None and (
            rpr.find(qn("w:ins")) is not None or rpr.find(qn("w:moveTo")) is not None
        )
        if inside_revision:
            edits.replace_in_place(edit.p, edit.start, edit.end, edit.new)
        else:
            edits.replace_range(edit.p, edit.start, edit.end, edit.new, self.rl.ctx)


def _wrap(rl: Redliner, p: etree._Element):
    from docx.text.paragraph import Paragraph

    return Paragraph(p, rl.document)


def _permute_in_place(current: Sequence[str], requested: Sequence[str]) -> list[str]:
    """Apply ``requested`` to the slots those clauses already occupy.

    Clauses the reviewer did not mention keep their exact position, so an
    unlisted clause never drifts to the end of the section as a side effect.
    """
    listed = set(requested)
    slots = [index for index, label in enumerate(current) if label in listed]
    result = list(current)
    for slot, label in zip(slots, requested, strict=False):
        result[slot] = label
    return result


def _label_prefix(label: str) -> str:
    """Render a clause number so ``parse_clause`` can read it back.

    ``SECTION_RE`` requires the trailing dot on a top-level number (``4.``),
    while ``CLAUSE_RE`` accepts ``4.1`` with or without one.  Emitting ``"4  "``
    produced a paragraph the clause tree could not see at all, so a new section
    was never renumbered.
    """
    return f"{label}. " if "." not in label else f"{label}  "


def _longest_increasing_run(values: Sequence[int]) -> list[int]:
    """Indices of a longest increasing subsequence -- the items that need not move."""
    if not values:
        return []
    tails: list[int] = []  # index in `values` of the smallest tail per length
    previous: list[int] = [-1] * len(values)
    for index, value in enumerate(values):
        low, high = 0, len(tails)
        while low < high:
            middle = (low + high) // 2
            if values[tails[middle]] < value:
                low = middle + 1
            else:
                high = middle
        previous[index] = tails[low - 1] if low else -1
        if low == len(tails):
            tails.append(index)
        else:
            tails[low] = index
    result: list[int] = []
    cursor = tails[-1]
    while cursor != -1:
        result.append(cursor)
        cursor = previous[cursor]
    return sorted(result)


def _last_paragraph_of(clause: Clause) -> etree._Element:
    """The final paragraph belonging to a clause, including its sub-clauses."""
    nodes = list(clause.walk())
    return nodes[-1].p


def apply_actions(
    rl: Redliner,
    items: Sequence[dict],
    renumber: bool = True,
    strict: bool = False,
    explain: bool = False,
) -> PlanReport:
    """Convenience wrapper around :class:`ActionPlanner`.

    ``explain=True`` also writes every action's ``rationale`` into the document
    as a Word comment; the pipeline turns it on by default.
    """
    return ActionPlanner(rl, renumber=renumber, strict=strict, explain=explain).run(items)
