# This Source Code Form is subject to the terms of the Mozilla Public
# License, v. 2.0. If a copy of the MPL was not distributed with this
# file, You can obtain one at https://mozilla.org/MPL/2.0/.
#
# Copyright (c) 2026 Aswanth B S

"""The end-to-end pipeline: a .docx in, a redlined .docx plus a report out.

    EXTRACT -> PROPOSE -> VALIDATE -> PLAN -> APPLY -> VERIFY -> REPORT

Each stage has one job and hands the next a plain data structure, so any stage
can be run, inspected or replaced on its own:

* **extract**  read the document into a clause tree and an outline
* **propose**  ask a reviewer (Claude, or the offline rules) for action items
* **validate** schema-check the items before the document is touched
* **plan/apply** turn the items into tracked changes, then derive the
  consequential renumbering and cross-reference edits
* **verify**   accept/reject the result and confirm the numbering is sound
* **report**   emit machine-readable JSON alongside the .docx

The verify stage is what makes this safe to automate: it proves the redline
resolves to the intended document and that the original is recoverable.
"""

from __future__ import annotations

import io
import json
from collections.abc import Callable, Sequence
from dataclasses import dataclass, field
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

import docx

from ..editing.compare import CompareStats, compare_into
from ..editing.redline import RedlineError, Redliner
from ..oxml.ns import qn
from ..oxml.textmap import paragraph_text
from ..structure.clauses import ClauseTree
from .actions import ActionPlanner, PlanReport, validate_actions
from .agent import Proposal, Reviewer, get_reviewer, load_proposal

DEFAULT_BRIEF = (
    "Review this agreement on behalf of the Customer. Push back on payment terms, "
    "liability, auto-renewal and security-incident timelines. Restructure where the "
    "ordering is unhelpful and add any covenant a customer would expect."
)


@dataclass
class StageLog:
    """One pipeline stage: what it did and whether it worked."""

    name: str
    ok: bool
    detail: str = ""
    data: dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> dict:
        out = {"stage": self.name, "ok": self.ok, "detail": self.detail}
        if self.data:
            out.update(self.data)
        return out


@dataclass
class PipelineResult:
    source: Path
    output: Path | None = None
    compare: CompareStats | None = None
    actions_path: Path | None = None
    report_path: Path | None = None
    proposal: Proposal | None = None
    plan: PlanReport | None = None
    stages: list[StageLog] = field(default_factory=list)
    checks: list[tuple[str, bool, str]] = field(default_factory=list)

    @property
    def ok(self) -> bool:
        return all(stage.ok for stage in self.stages) and all(
            passed for _, passed, _ in self.checks
        )

    def to_dict(self) -> dict:
        return {
            "document": str(self.source),
            "output": str(self.output) if self.output else None,
            "action_items": str(self.actions_path) if self.actions_path else None,
            "generated_at": datetime.now(UTC).strftime("%Y-%m-%dT%H:%M:%SZ"),
            "ok": self.ok,
            "stages": [s.to_dict() for s in self.stages],
            "proposal": {
                "generated_by": self.proposal.source if self.proposal else None,
                "model": self.proposal.model if self.proposal else None,
                "summary": self.proposal.summary if self.proposal else "",
                "usage": self.proposal.usage if self.proposal else {},
                "count": len(self.proposal.action_items) if self.proposal else 0,
            },
            "compare": self.compare.__dict__ if self.compare else None,
            "plan": self.plan.to_dict() if self.plan else {},
            "verification": [
                {"check": name, "passed": passed, "detail": detail}
                for name, passed, detail in self.checks
            ],
        }

    def format(self) -> str:
        lines = []
        for stage in self.stages:
            mark = "ok " if stage.ok else "FAIL"
            lines.append(f"  [{mark}] {stage.name:<10} {stage.detail}")
        if self.checks:
            lines.append("")
            for name, passed, detail in self.checks:
                mark = "PASS" if passed else "FAIL"
                lines.append(
                    f"  [{mark}] {name}" + (f" -- {detail}" if detail and not passed else "")
                )
        return "\n".join(lines)


class RedlinePipeline:
    """Runs the stages in order and collects a report."""

    def __init__(
        self,
        source: str | Path,
        author: str = "AI Contract Reviewer",
        date: str | None = None,
        reviewer: Reviewer | str = "rules",
        brief: str = DEFAULT_BRIEF,
        renumber: bool = True,
        strict: bool = False,
        explain: bool = True,
        revised: str | Path | None = None,
        comments: Sequence[dict] | None = None,
        similarity_floor: float = 0.45,
        on_stage: Callable[[StageLog], None] | None = None,
    ) -> None:
        self.source = Path(source)
        self.author = author
        self.date = date
        self.brief = brief
        self.renumber = renumber
        self.strict = strict
        #: write each action's rationale into the .docx as a Word comment
        self.explain = explain
        self.revised = Path(revised) if revised else None
        self.comments = list(comments or [])
        self.similarity_floor = similarity_floor
        self.reviewer = get_reviewer(reviewer) if isinstance(reviewer, str) else reviewer
        self.on_stage = on_stage
        self.result = PipelineResult(source=self.source)

    # -- stage plumbing --------------------------------------------------
    def _stage(self, name: str, ok: bool, detail: str = "", **data) -> StageLog:
        log = StageLog(name, ok, detail, data)
        self.result.stages.append(log)
        if self.on_stage:
            self.on_stage(log)
        return log

    # -- 1. extract ------------------------------------------------------
    def extract(self) -> ClauseTree:
        document = docx.Document(str(self.source))
        tree = ClauseTree(document.element.body)
        paragraphs = sum(1 for _ in document.element.body.iter(qn("w:p")))
        self._stage(
            "extract",
            True,
            f"{len(tree.sections)} sections, {len(tree.all())} clauses, "
            f"{paragraphs} paragraphs, {len(document.tables)} tables",
            sections=len(tree.sections),
            clauses=len(tree.all()),
            paragraphs=paragraphs,
        )
        self._tree = tree
        return tree

    # -- 2. propose ------------------------------------------------------
    def propose(self, actions_file: str | Path | None = None) -> Proposal:
        if actions_file and Path(actions_file).exists():
            proposal = load_proposal(actions_file)
            detail = (
                f"loaded {len(proposal.action_items)} action items from {Path(actions_file).name}"
            )
        elif self.reviewer is None:
            proposal = Proposal(action_items=[], summary="", source="none")
            detail = "no reviewer: applying compare and comments only"
        else:
            proposal = self.reviewer.propose(self._tree, self.brief)
            detail = f"{proposal.source} produced {len(proposal.action_items)} action items"
            if proposal.usage:
                detail += (
                    f" ({proposal.usage.get('input_tokens', 0)} in / "
                    f"{proposal.usage.get('output_tokens', 0)} out tokens)"
                )
        self._stage(
            "propose", True, detail, count=len(proposal.action_items), source=proposal.source
        )
        self.result.proposal = proposal
        return proposal

    # -- 3. validate -----------------------------------------------------
    def validate(self, proposal: Proposal) -> list[str]:
        problems = validate_actions(proposal.action_items)
        self._stage(
            "validate",
            not problems,
            "schema clean" if not problems else f"{len(problems)} problem(s)",
            problems=problems,
        )
        if problems and self.strict:
            raise RedlineError("action items failed validation:\n  " + "\n  ".join(problems))
        return problems

    # -- 4/5. plan + apply -----------------------------------------------
    def apply(self, proposal: Proposal) -> tuple[Redliner, PlanReport]:
        """Compare, run the action plan, annotate -- then renumber once.

        The planner is built *before* the compare so its baseline is the
        original numbering, and renumbering is deferred to the very end so a
        single pass sees the paragraph order both stages produced.
        """
        rl = Redliner(self.source, author=self.author, date=self.date)
        planner = ActionPlanner(rl, renumber=False, strict=self.strict, explain=self.explain)

        if self.revised is not None:
            stats = compare_into(rl, self.revised, similarity_floor=self.similarity_floor)
            self.result.compare = stats
            self._stage("compare", True, stats.format(), **stats.__dict__)

        items = list(proposal.action_items) + [
            {"type": "comment", "id": f"CM-{index + 1:04d}", **comment}
            for index, comment in enumerate(self.comments)
        ]
        plan = planner.run(items)
        self._stage(
            "plan",
            plan.failed == 0,
            f"{plan.applied}/{len(plan.results)} actions applied"
            + (f", {plan.failed} failed" if plan.failed else ""),
            applied=plan.applied,
            failed=plan.failed,
        )

        if self.renumber:
            planner.finalize()
        elif self.explain:
            planner.annotate_rationales()
        self._stage(
            "renumber",
            True,
            f"{len(plan.renumbered)} clause number(s) rewritten, "
            f"{len(plan.references)} cross-reference(s) updated"
            + (f", {len(plan.dangling_references)} dangling" if plan.dangling_references else ""),
            renumbered=len(plan.renumbered),
            references=len(plan.references),
        )

        revisions = len(rl.summary())
        annotations = _count_comments(rl)
        detail = f"{revisions} tracked changes written"
        if annotations:
            detail += f", {annotations} comment(s) attached"
        # A comments-only run legitimately produces no revisions at all, so the
        # bar is "something happened", not "something was struck".
        self._stage(
            "apply",
            bool(revisions or plan.applied),
            detail,
            revisions=revisions,
            comments=annotations,
        )
        self.result.plan = plan
        return rl, plan

    # -- 6. verify -------------------------------------------------------
    def verify(self, rl: Redliner) -> list[tuple[str, bool, str]]:
        """Resolve the redline both ways and check the result holds together."""
        buf = io.BytesIO()
        rl.document.save(buf)

        original = _body_text(docx.Document(str(self.source)))

        buf.seek(0)
        accepted = Redliner(buf, track_changes=False)
        accepted.accept_all()
        buf.seek(0)
        rejected = Redliner(buf, track_changes=False)
        rejected.reject_all()

        checks: list[tuple[str, bool, str]] = []
        checks.append(("redline reopens as a valid docx", True, ""))
        checks.append(
            ("reject restores the original document", _body_text(rejected.document) == original, "")
        )
        # A run that only attaches comments makes no edits by design, so the
        # assertion flips: the text must come back *unchanged*.
        if len(rl.summary()):
            checks.append(
                ("accept changes the document", _body_text(accepted.document) != original, "")
            )
        else:
            checks.append(
                (
                    "annotation-only run leaves the text unchanged",
                    _body_text(accepted.document) == original,
                    "",
                )
            )
        checks.append(("no tracked changes survive accept", len(accepted.summary()) == 0, ""))

        rejected_tree = ClauseTree(rejected.document.element.body)
        original_tree = ClauseTree(docx.Document(str(self.source)).element.body)
        checks.append(
            (
                "reject restores the original numbering",
                [c.label for c in rejected_tree] == [c.label for c in original_tree],
                "",
            )
        )

        # Numbering assertions only make sense when the renumbering pass ran;
        # --no-renumber deliberately leaves the document inconsistent.
        if self.renumber:
            accepted_tree = ClauseTree(accepted.document.element.body)
            gaps = _numbering_gaps(accepted_tree)
            # One bad nesting near the front makes every later clause report a
            # gap, so lead with the first divergence and a count rather than
            # four arbitrary entries.
            detail = (
                f"{gaps[0]} (+{len(gaps) - 1} more)" if len(gaps) > 1 else (gaps[0] if gaps else "")
            )
            checks.append(("accepted numbering is contiguous", not gaps, detail))
            # Measure against the original, not against nothing. A real contract
            # cites other agreements ("Section 5 of the Master Agreement"),
            # statutes, and exhibit-local numbering -- all of which look dangling
            # and none of which this redline broke. An absolute check fails on
            # run one and gets ignored from then on.
            original_doc = docx.Document(str(self.source))
            baseline = _broken_references(
                ClauseTree(original_doc.element.body), original_doc.element.body
            )
            broken = _broken_references(accepted_tree, accepted.document.element.body)
            introduced = broken - baseline
            checks.append(
                (
                    "no cross-reference was broken",
                    not introduced,
                    "newly dangling: " + ", ".join(sorted(introduced)[:6]),
                )
            )

        self.result.checks = checks
        passed = sum(1 for _, ok, _ in checks if ok)
        self._stage(
            "verify",
            passed == len(checks),
            f"{passed}/{len(checks)} checks passed",
            passed=passed,
            total=len(checks),
        )
        self._accepted, self._rejected = accepted, rejected
        return checks

    # -- 7. report -------------------------------------------------------
    def write(
        self,
        rl: Redliner,
        output: str | Path,
        report_path: str | Path | None = None,
        actions_path: str | Path | None = None,
        proposal: Proposal | None = None,
    ) -> PipelineResult:
        self.result.output = Path(rl.save(output))
        if actions_path and proposal:
            payload = proposal.to_dict(document=self.source.name, author=self.author)
            Path(actions_path).write_text(json.dumps(payload, indent=2) + "\n", encoding="utf-8")
            self.result.actions_path = Path(actions_path)
        if report_path:
            Path(report_path).write_text(
                json.dumps(self.result.to_dict(), indent=2) + "\n", encoding="utf-8"
            )
            self.result.report_path = Path(report_path)
        self._stage("report", True, f"wrote {self.result.output.name}")
        return self.result

    # -- everything ------------------------------------------------------
    def run(
        self,
        output: str | Path,
        actions_file: str | Path | None = None,
        report_path: str | Path | None = None,
        write_actions: bool = True,
    ) -> PipelineResult:
        self.extract()
        proposal = self.propose(actions_file)
        self.validate(proposal)
        rl, _ = self.apply(proposal)
        self.verify(rl)
        return self.write(
            rl,
            output,
            report_path=report_path,
            actions_path=actions_file if write_actions else None,
            proposal=proposal,
        )


# ---------------------------------------------------------------------------
def _count_comments(rl: Redliner) -> int:
    """How many comments the document part now holds."""
    for rel in rl.document.part.rels.values():
        if rel.reltype.endswith("/comments"):
            try:
                return len(rel.target_part.element.findall(qn("w:comment")))
            except AttributeError:  # pragma: no cover
                return 0
    return 0


def _body_text(document) -> list[str]:
    return [paragraph_text(p) for p in document.element.body.iter(qn("w:p"))]


def _numbering_gaps(tree: ClauseTree) -> list[str]:
    """Report any place where sibling numbers are not 1..n in order."""
    problems: list[str] = []

    def check(siblings, parent_label: str) -> None:
        for index, clause in enumerate(siblings, start=1):
            expected = f"{parent_label}.{index}" if parent_label else str(index)
            if clause.label != expected:
                problems.append(f"expected {expected}, found {clause.label}")
            check(clause.children, clause.label)

    check(tree.sections, "")
    return problems


def _broken_references(tree: ClauseTree, root) -> set[str]:
    """Cross-references that point at a clause number that no longer exists."""
    from ..structure.clauses import iter_references

    known = {clause.label for clause in tree}
    broken: set[str] = set()
    for p in root.iter(qn("w:p")):
        for _, _, number, _ in iter_references(paragraph_text(p)):
            if number not in known:
                broken.add(number)
    return broken


# ---------------------------------------------------------------------------
# the everything-at-once entry point
# ---------------------------------------------------------------------------
def full_redline(
    source: str | Path,
    output: str | Path,
    *,
    revised: str | Path | None = None,
    actions: Sequence[dict] | str | Path | None = None,
    reviewer: Reviewer | str | None = None,
    brief: str = DEFAULT_BRIEF,
    comments: Sequence[dict] | None = None,
    author: str = "AI Contract Reviewer",
    date: str | None = None,
    renumber: bool = True,
    strict: bool = False,
    explain: bool = True,
    similarity_floor: float = 0.45,
    report_path: str | Path | None = None,
    actions_path: str | Path | None = None,
    on_stage: Callable[[StageLog], None] | None = None,
) -> PipelineResult:
    """Run every layer of the redliner over one document, in the right order.

    ``EXTRACT -> PROPOSE -> VALIDATE -> COMPARE -> PLAN -> COMMENT ->
    RENUMBER -> VERIFY -> REPORT``

    Any subset of the three edit sources may be supplied and they compose:

    * ``revised``   -- a second ``.docx`` to diff against, Word-Compare style.
    * ``actions``   -- action items, as a list or a path to a JSON file.
    * ``reviewer``  -- ``"rules"``, ``"claude"``, ``"openai"``, or an instance,
      asked for action items when ``actions`` is not given.
    * ``comments``  -- ``[{"clause": "3.2", "text": "..."}]`` or
      ``[{"find": "phrase", "text": "..."}]``, applied after every edit.

    ``explain`` (on by default) additionally writes each action's ``rationale``
    into the document as a Word comment, so the *why* is visible in the review
    pane rather than only in the JSON report.

    Order is the point. The compare runs first so scripted actions address the
    document the diff produced; comments run after every edit so they anchor to
    finished text; and clause renumbering runs once at the very end, over the
    combined result, against a baseline captured before any of it.

    Returns the same :class:`PipelineResult` the staged API returns --
    ``result.ok`` is false if any stage or verification check failed.
    """
    if revised is None and actions is None and reviewer is None and not comments:
        raise RedlineError(
            "full_redline needs something to do: pass revised=, actions=, reviewer= or comments="
        )

    inline_actions: Sequence[dict] | None = None
    actions_file: str | Path | None = None
    if isinstance(actions, (str, Path)):
        actions_file = actions
    elif actions is not None:
        inline_actions = list(actions)

    pipeline = RedlinePipeline(
        source,
        author=author,
        date=date,
        # Pass None straight through: `RedlinePipeline` defaults to "rules"
        # when constructed directly, but here None means "no reviewer" -- a
        # compare-only or comments-only run must not silently gain a plan.
        reviewer=reviewer,
        brief=brief,
        renumber=renumber,
        strict=strict,
        explain=explain,
        revised=revised,
        comments=comments,
        similarity_floor=similarity_floor,
        on_stage=on_stage,
    )
    pipeline.extract()

    if inline_actions is not None:
        proposal = Proposal(action_items=list(inline_actions), summary="", source="inline")
        pipeline._stage(
            "propose",
            True,
            f"{len(inline_actions)} action items supplied inline",
            count=len(inline_actions),
            source="inline",
        )
        pipeline.result.proposal = proposal
    else:
        proposal = pipeline.propose(actions_file)

    problems = pipeline.validate(proposal)
    if problems:
        # A schema problem is not a per-action miss to be reported and skipped:
        # the input is malformed, so nothing should run against the document.
        raise RedlineError("action items failed validation:\n  " + "\n  ".join(problems))

    rl, _ = pipeline.apply(proposal)
    pipeline.verify(rl)
    return pipeline.write(
        rl, output, report_path=report_path, actions_path=actions_path, proposal=proposal
    )
