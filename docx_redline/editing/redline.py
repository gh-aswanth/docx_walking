# This Source Code Form is subject to the terms of the Mozilla Public
# License, v. 2.0. If a copy of the MPL was not distributed with this
# file, You can obtain one at https://mozilla.org/MPL/2.0/.
#
# Copyright (c) 2026 Aswanth B S

"""``Redliner`` -- the public, python-docx-flavoured API for tracked changes.

    rl = Redliner("contract.docx", author="Outside Counsel")
    rl.replace_text("thirty (30) days", "forty-five (45) days")
    rl.insert_paragraph_after(rl.find_paragraph("3.4"), "3.5  New clause ...")
    rl.save("contract_redlined.docx")

Everything the class does ends up as real OOXML revision markup, so Word,
Google Docs and LibreOffice all show it in their native review UI and the
changes can be accepted or rejected one by one.
"""

from __future__ import annotations

import copy
import re
from collections.abc import Iterable, Sequence
from dataclasses import dataclass
from pathlib import Path

import docx
from docx.document import Document as _Document
from docx.table import Table, _Cell, _Row
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from lxml import etree

from ..errors import RedlineError
from ..oxml import edits
from ..oxml.diffing import diff_ops
from ..oxml.elements import PPR_ORDER, deepcopy_without, get_or_add, make, make_run
from ..oxml.ns import qn
from ..oxml.revisions import Author, RevisionContext, RevisionIds
from ..oxml.textmap import ParagraphText, iter_visible_runs, paragraph_text
from . import review

__all__ = ["RedlineError", "Redliner"]


@dataclass(frozen=True)
class Match:
    """A located run of text inside a paragraph."""

    paragraph: Paragraph
    start: int
    end: int

    @property
    def text(self) -> str:
        return paragraph_text(self.paragraph._p)[self.start : self.end]


class Redliner:
    """Applies tracked changes to a ``.docx``.

    Parameters
    ----------
    source:
        Path to a ``.docx``, a file-like object, or an existing
        :class:`docx.document.Document`.
    author / initials / date:
        Attribution stamped on every revision.  ``date`` accepts a
        ``datetime`` or a pre-formatted ISO-8601 UTC string; it defaults to
        "now", and pinning it makes output byte-reproducible.
    track_changes:
        Also switch on Word's *Track Changes* toggle in ``settings.xml`` so
        that edits a human makes afterwards keep being recorded.
    scope:
        Which stories document-wide operations touch.  Defaults to body plus
        headers/footers plus foot/endnotes.
    """

    def __init__(
        self,
        source,
        author: str = "Redline",
        initials: str | None = None,
        date=None,
        track_changes: bool = True,
        scope: Sequence[str] = ("body", "headers", "footers", "notes"),
    ) -> None:
        self.document: _Document = (
            source if isinstance(source, _Document) else docx.Document(source)
        )
        self.scope = tuple(scope)
        self.ctx = RevisionContext(
            author=Author(name=author, initials=initials, date=date),
            ids=RevisionIds.from_elements(self._all_roots()),
        )
        if track_changes:
            self.enable_track_changes()

    # ------------------------------------------------------------------
    # parts & traversal
    # ------------------------------------------------------------------
    def _all_roots(self) -> list[etree._Element]:
        """Every XML root we may write revisions into."""
        roots = [self.document.element.body]
        for section in self.document.sections:
            for attr in (
                "header",
                "footer",
                "even_page_header",
                "even_page_footer",
                "first_page_header",
                "first_page_footer",
            ):
                try:
                    part = getattr(section, attr)
                except (KeyError, ValueError):  # pragma: no cover - odd sections
                    continue
                if part is not None and part.is_linked_to_previous is False:
                    roots.append(part._element)
        for rel_type in ("footnotes", "endnotes"):
            root = self._notes_root(rel_type)
            if root is not None:
                roots.append(root)
        return roots

    def _notes_root(self, kind: str) -> etree._Element | None:
        part = self.document.part
        for rel in part.rels.values():
            if rel.reltype.endswith("/" + kind):
                try:
                    return rel.target_part.element
                except AttributeError:  # pragma: no cover
                    return None
        return None

    def _scoped_roots(self) -> list[etree._Element]:
        roots: list[etree._Element] = []
        if "body" in self.scope:
            roots.append(self.document.element.body)
        for root in self._all_roots()[1:]:
            tag = root.tag.split("}")[-1]
            if (
                (tag in ("hdr",) and "headers" in self.scope)
                or (tag in ("ftr",) and "footers" in self.scope)
                or (tag in ("footnotes", "endnotes") and "notes" in self.scope)
            ):
                roots.append(root)
        return roots

    def paragraphs(self, include_tables: bool = True) -> list[Paragraph]:
        """Every paragraph in scope, in document order (table cells included)."""
        out: list[Paragraph] = []
        for root in self._scoped_roots():
            for p in root.iter(qn("w:p")):
                if not include_tables and p.getparent().tag == qn("w:tc"):
                    continue
                if self._is_inside_deleted_row(p):
                    continue
                out.append(Paragraph(p, self.document))
        return out

    @staticmethod
    def _is_inside_deleted_row(p: etree._Element) -> bool:
        tr = p.getparent()
        while tr is not None and tr.tag != qn("w:tr"):
            tr = tr.getparent()
        if tr is None:
            return False
        trpr = tr.find(qn("w:trPr"))
        return trpr is not None and trpr.find(qn("w:del")) is not None

    @staticmethod
    def text_of(paragraph: Paragraph) -> str:
        """Current text of a paragraph, tracked insertions included.

        Use this rather than python-docx's ``Paragraph.text``: that property
        only looks at ``w:r`` children of the paragraph, so it silently misses
        anything wrapped in a ``w:ins`` -- i.e. every insertion this library
        makes -- while still showing text that has been struck out.
        """
        return paragraph_text(paragraph._p)

    def tables(self) -> list[Table]:
        out: list[Table] = []
        for root in self._scoped_roots():
            for tbl in root.iter(qn("w:tbl")):
                out.append(Table(tbl, self.document))
        return out

    # ------------------------------------------------------------------
    # locating
    # ------------------------------------------------------------------
    def find_paragraphs(
        self,
        contains: str | None = None,
        regex: str | None = None,
        style: str | None = None,
        startswith: str | None = None,
        ignore_case: bool = False,
        include_tables: bool = True,
    ) -> list[Paragraph]:
        """Filter paragraphs by text / regex / style.  All filters are ANDed."""
        flags = re.IGNORECASE if ignore_case else 0
        pattern = re.compile(regex, flags) if regex else None
        results = []
        for para in self.paragraphs(include_tables=include_tables):
            text = paragraph_text(para._p)
            probe = text.lower() if ignore_case else text
            if contains is not None:
                needle = contains.lower() if ignore_case else contains
                if needle not in probe:
                    continue
            if startswith is not None:
                needle = startswith.lower() if ignore_case else startswith
                if not probe.lstrip().startswith(needle):
                    continue
            if pattern is not None and not pattern.search(text):
                continue
            if style is not None and (para.style is None or para.style.name != style):
                continue
            results.append(para)
        return results

    def find_paragraph(self, *args, **kwargs) -> Paragraph:
        """Like :meth:`find_paragraphs` but requires exactly one hit."""
        found = self.find_paragraphs(*args, **kwargs)
        if not found:
            raise RedlineError(f"no paragraph matched {args or kwargs}")
        return found[0]

    def find_text(
        self,
        needle: str,
        regex: bool = False,
        ignore_case: bool = False,
        limit: int | None = None,
    ) -> list[Match]:
        """Locate ``needle`` across the document, returning character spans."""
        flags = re.IGNORECASE if ignore_case else 0
        pattern = re.compile(needle if regex else re.escape(needle), flags)
        matches: list[Match] = []
        for para in self.paragraphs():
            text = paragraph_text(para._p)
            for m in pattern.finditer(text):
                if m.start() == m.end():
                    continue
                matches.append(Match(para, m.start(), m.end()))
                if limit is not None and len(matches) >= limit:
                    return matches
        return matches

    # ------------------------------------------------------------------
    # 1. INSERT text
    # ------------------------------------------------------------------
    def insert_text(self, paragraph: Paragraph, offset: int, text: str) -> None:
        """Insert ``text`` at character ``offset`` of ``paragraph``, tracked."""
        edits.insert_text(paragraph._p, self._clamp(paragraph, offset), text, self.ctx)

    def insert_text_before(self, needle: str, text: str, **kw) -> int:
        return self._insert_relative(needle, text, before=True, **kw)

    def insert_text_after(self, needle: str, text: str, **kw) -> int:
        return self._insert_relative(needle, text, before=False, **kw)

    def _insert_relative(
        self, needle: str, text: str, before: bool, count: int | None = 1, **kw
    ) -> int:
        matches = self.find_text(needle, limit=count, **kw)
        for match in reversed(matches):
            offset = match.start if before else match.end
            edits.insert_text(match.paragraph._p, offset, text, self.ctx)
        return len(matches)

    def append_text(self, paragraph: Paragraph, text: str) -> None:
        """Append ``text`` to the end of a paragraph, tracked."""
        edits.insert_text(paragraph._p, len(paragraph_text(paragraph._p)), text, self.ctx)

    # ------------------------------------------------------------------
    # 2. DELETE text
    # ------------------------------------------------------------------
    def delete_text(self, paragraph: Paragraph, start: int, end: int) -> None:
        edits.delete_range(paragraph._p, start, end, self.ctx)

    def delete_matching(
        self, needle: str, regex: bool = False, ignore_case: bool = False, count: int | None = None
    ) -> int:
        """Strike every occurrence of ``needle`` (or the first ``count``)."""
        matches = self.find_text(needle, regex=regex, ignore_case=ignore_case, limit=count)
        for match in reversed(matches):
            edits.delete_range(match.paragraph._p, match.start, match.end, self.ctx)
        return len(matches)

    # ------------------------------------------------------------------
    # 3. REPLACE text
    # ------------------------------------------------------------------
    def replace_text(
        self,
        old: str,
        new: str,
        regex: bool = False,
        ignore_case: bool = False,
        count: int | None = 1,
        insertion_first: bool = False,
    ) -> int:
        """Redline ``old`` -> ``new``.  Returns the number of replacements.

        ``count=None`` replaces every occurrence.  Matches are applied back to
        front so earlier offsets stay valid.
        """
        matches = self.find_text(old, regex=regex, ignore_case=ignore_case, limit=count)
        by_para: dict[int, list[Match]] = {}
        for match in matches:
            by_para.setdefault(id(match.paragraph._p), []).append(match)
        for group in by_para.values():
            for match in sorted(group, key=lambda m: m.start, reverse=True):
                replacement = new
                if regex:
                    text = paragraph_text(match.paragraph._p)
                    flags = re.IGNORECASE if ignore_case else 0
                    m = re.compile(old, flags).match(text, match.start, match.end)
                    if m is not None:
                        replacement = m.expand(new)
                edits.replace_range(
                    match.paragraph._p,
                    match.start,
                    match.end,
                    replacement,
                    self.ctx,
                    insertion_first,
                )
        return len(matches)

    # ------------------------------------------------------------------
    # 4. UPDATE a whole paragraph via word-level diff
    # ------------------------------------------------------------------
    def set_paragraph_text(
        self, paragraph: Paragraph, new_text: str, ignore_case: bool = False
    ) -> int:
        """Rewrite a paragraph to ``new_text``, marking only the words that moved.

        This is the workhorse for "here is the revised clause" edits: a
        word-level diff keeps the markup minimal and reviewable instead of
        striking the whole paragraph.
        """
        p = paragraph._p
        original = paragraph_text(p)
        ops = diff_ops(original, new_text, ignore_case=ignore_case)
        for op in sorted(ops, key=lambda o: o.start, reverse=True):
            if op.kind == "insert":
                edits.insert_text(p, op.start, op.text, self.ctx)
            elif op.kind == "delete":
                edits.delete_range(p, op.start, op.end, self.ctx)
            else:
                edits.replace_range(p, op.start, op.end, op.text, self.ctx)
        return len(ops)

    # ------------------------------------------------------------------
    # 5. INSERT / DELETE whole paragraphs
    # ------------------------------------------------------------------
    def insert_paragraph_after(
        self,
        reference: Paragraph,
        text: str = "",
        style: str | None = None,
        copy_format: bool = True,
    ) -> Paragraph:
        """Add a brand-new paragraph after ``reference`` as a tracked insertion."""
        return self._insert_paragraph(reference, text, style, copy_format, after=True)

    def insert_paragraph_before(
        self,
        reference: Paragraph,
        text: str = "",
        style: str | None = None,
        copy_format: bool = True,
    ) -> Paragraph:
        return self._insert_paragraph(reference, text, style, copy_format, after=False)

    def _insert_paragraph(self, reference, text, style, copy_format, after: bool) -> Paragraph:
        ref = reference._p
        new_p = make("w:p")
        if copy_format:
            ref_ppr = ref.find(qn("w:pPr"))
            if ref_ppr is not None:
                # Numbering (w:numPr) is deliberately carried over: a new item
                # inserted into a numbered list should be numbered, and Word
                # renumbers the rest of the list on accept.
                new_p.append(deepcopy_without(ref_ppr, ("w:rPr", "w:sectPr", "w:pPrChange")))
        if text:
            new_p.append(make_run(text, self._reference_rpr(ref)))
        if after:
            ref.addnext(new_p)
        else:
            ref.addprevious(new_p)
        paragraph = Paragraph(new_p, reference._parent)
        if style is not None:
            self.apply_style(paragraph, style)
        edits.mark_paragraph_inserted(new_p, self.ctx)
        return paragraph

    def append_paragraph(self, text: str, style: str | None = None) -> Paragraph:
        """Append a tracked new paragraph at the end of the body."""
        body = self.document.element.body
        paragraphs = body.findall(qn("w:p"))
        if paragraphs:
            return self.insert_paragraph_after(
                Paragraph(paragraphs[-1], self.document), text, style
            )
        paragraph = self.document.add_paragraph(text, style=style)
        edits.mark_paragraph_inserted(paragraph._p, self.ctx)
        return paragraph

    def delete_paragraph(self, paragraph: Paragraph) -> None:
        """Strike a whole paragraph, including its paragraph mark.

        Deleting the mark is what makes the paragraph actually disappear on
        accept; striking only the runs would leave an empty paragraph behind.
        """
        edits.mark_paragraph_deleted(paragraph._p, self.ctx)

    def delete_paragraphs(self, paragraphs: Iterable[Paragraph]) -> int:
        count = 0
        for paragraph in list(paragraphs):
            self.delete_paragraph(paragraph)
            count += 1
        return count

    # ------------------------------------------------------------------
    # 6. MOVE (Word's move-from / move-to pair)
    # ------------------------------------------------------------------
    def move_paragraph(self, paragraph: Paragraph, after: Paragraph) -> Paragraph:
        """Relocate a paragraph, recorded as a true Word *move* revision."""
        clone = self.relocate_paragraph(paragraph._p, after=after._p)
        return Paragraph(clone, after._parent)

    def relocate_paragraph(
        self,
        source: etree._Element,
        after: etree._Element | None = None,
        before: etree._Element | None = None,
    ) -> etree._Element:
        """Move ``source`` next to ``after``/``before``; return the new paragraph.

        A clean paragraph becomes a proper ``moveFrom``/``moveTo`` pair, which
        Word labels "Moved from"/"Moved to" in the review pane.

        A paragraph that *already* carries revisions cannot: OOXML has no
        meaningful nesting for "this text was inserted, and then the paragraph
        was moved", and copying the source's ``w:del`` runs to the destination
        would show reviewers a strikeout on text that never existed there.  For
        those, the move is recorded the way Word's own Compare records an
        edited-and-moved block -- struck at the source, inserted at the
        destination, with the destination carrying the post-edit text.
        """
        if paragraph_has_revisions(source):
            return self._cut_and_paste(source, after, before)
        return self._move_revision(source, after, before)

    def _move_revision(self, source, after, before) -> etree._Element:
        name = f"move{self.ctx.ids.next()}"
        clone = copy.deepcopy(source)
        _place(clone, after, before)

        edits._wrap_group_all(list(iter_visible_runs(clone)), "w:moveTo", self.ctx)
        edits.mark_paragraph_mark(clone, "w:moveTo", self.ctx)
        self._wrap_range(clone, "moveTo", name)

        runs = list(iter_visible_runs(source))
        if runs:
            edits._wrap_group_all(runs, "w:moveFrom", self.ctx)
        edits.mark_paragraph_mark(source, "w:moveFrom", self.ctx)
        self._wrap_range(source, "moveFrom", name)
        return clone

    def _cut_and_paste(self, source, after, before) -> etree._Element:
        clone = _accepted_copy(source)
        _place(clone, after, before)
        edits.mark_paragraph_inserted(clone, self.ctx)
        self.delete_paragraph(Paragraph(source, self.document))
        return clone

    def _wrap_range(self, p: etree._Element, kind: str, name: str) -> None:
        start = self.ctx.stamp(make(f"w:{kind}RangeStart"))
        start.set(qn("w:name"), name)
        end = make(f"w:{kind}RangeEnd")
        end.set(qn("w:id"), start.get(qn("w:id")))
        ppr = p.find(qn("w:pPr"))
        p.insert(1 if ppr is not None else 0, start)
        p.append(end)

    # ------------------------------------------------------------------
    # 7. TABLES
    # ------------------------------------------------------------------
    def insert_table_row(
        self, table: Table, index: int | None = None, values: Sequence[str] | None = None
    ) -> _Row:
        """Insert a tracked new row, cloning an existing row for cell geometry."""
        rows = table._tbl.findall(qn("w:tr"))
        if not rows:
            raise RedlineError("table has no rows to clone geometry from")
        if index is None or index >= len(rows):
            template, append_after = rows[-1], True
            index = len(rows)
        else:
            template, append_after = rows[index], False
        new_tr = self._blank_row_like(template)
        if append_after:
            rows[-1].addnext(new_tr)
        else:
            rows[index].addprevious(new_tr)
        if values:
            self._fill_row(new_tr, values)
        edits.mark_row_inserted(new_tr, self.ctx)
        return _Row(new_tr, table)

    def _blank_row_like(self, template: etree._Element) -> etree._Element:
        """Clone a row's structure (cell widths, borders, shading) but no text."""
        clone = deepcopy_without(template, ("w:trPr",))
        trpr = template.find(qn("w:trPr"))
        if trpr is not None:
            clone.insert(0, deepcopy_without(trpr, ("w:ins", "w:del", "w:trPrChange")))
        for tc in clone.findall(qn("w:tc")):
            keep = tc.find(qn("w:p"))
            for p in tc.findall(qn("w:p"))[1:]:
                tc.remove(p)
            for tbl in tc.findall(qn("w:tbl")):
                tc.remove(tbl)
            if keep is None:
                tc.append(make("w:p"))
                continue
            for child in list(keep):
                if child.tag != qn("w:pPr"):
                    keep.remove(child)
            ppr = keep.find(qn("w:pPr"))
            if ppr is not None:
                for stale in ppr.findall(qn("w:rPr")):
                    ppr.remove(stale)
        return clone

    def _fill_row(self, tr: etree._Element, values: Sequence[str]) -> None:
        cells = tr.findall(qn("w:tc"))
        for cell, value in zip(cells, values, strict=False):
            p = cell.find(qn("w:p"))
            if p is None:
                p = make("w:p")
                cell.append(p)
            if value:
                p.append(make_run(str(value)))

    def delete_table_row(self, table: Table, index: int) -> None:
        rows = table._tbl.findall(qn("w:tr"))
        if not 0 <= index < len(rows):
            raise RedlineError(f"row index {index} out of range (0..{len(rows) - 1})")
        edits.mark_row_deleted(rows[index], self.ctx)

    def delete_row(self, row: _Row) -> None:
        edits.mark_row_deleted(row._tr, self.ctx)

    def set_cell_text(self, cell: _Cell, new_text: str) -> int:
        """Redline a cell's text, diffing against what is currently there."""
        paragraphs = cell._tc.findall(qn("w:p"))
        if not paragraphs:
            p = make("w:p")
            cell._tc.append(p)
            paragraphs = [p]
        first = Paragraph(paragraphs[0], cell)
        changes = self.set_paragraph_text(first, new_text)
        for extra in paragraphs[1:]:
            self.delete_paragraph(Paragraph(extra, cell))
            changes += 1
        return changes

    # ------------------------------------------------------------------
    # 8. FORMATTING revisions (rPrChange / pPrChange)
    # ------------------------------------------------------------------
    def format_runs(self, runs: Iterable[Run], **props) -> int:
        """Apply character formatting as a tracked *formatting* revision.

        Accepts the python-docx run/font attribute names, e.g.
        ``bold=True, italic=False, underline=True, name="Arial", size=Pt(12),
        color="FF0000", highlight=WD_COLOR_INDEX.YELLOW, strike=True``.
        """
        count = 0
        for run in list(runs):
            rpr = run._r.find(qn("w:rPr"))
            baseline = copy.deepcopy(rpr) if rpr is not None else None
            self._apply_run_props(run, props)
            edits.record_rpr_change(run._r, baseline, self.ctx)
            count += 1
        return count

    @staticmethod
    def _apply_run_props(run: Run, props: dict) -> None:
        font_only = {"name", "size", "color", "highlight_color", "subscript", "superscript"}
        for key, value in props.items():
            if key == "color":
                from docx.shared import RGBColor

                run.font.color.rgb = (
                    value if isinstance(value, RGBColor) else RGBColor.from_string(str(value))
                )
            elif key in ("highlight", "highlight_color"):
                # The schema advertises this field to models, so it has to
                # accept the name as a plain string, not just the enum member.
                run.font.highlight_color = _highlight(value)
            elif key in font_only:
                setattr(run.font, key, value)
            elif hasattr(run, key):
                setattr(run, key, value)
            else:
                setattr(run.font, key, value)

    def format_paragraph_text(self, paragraph: Paragraph, **props) -> int:
        """Formatting revision over every (visible) run of a paragraph."""
        runs = [Run(r, paragraph) for r in iter_visible_runs(paragraph._p)]
        return self.format_runs(runs, **props)

    def format_matching(
        self, needle: str, regex: bool = False, count: int | None = None, **props
    ) -> int:
        """Format just the matched text, splitting runs at the match boundaries."""
        matches = self.find_text(needle, regex=regex, limit=count)
        total = 0
        for match in reversed(matches):
            p = match.paragraph._p
            view = ParagraphText(p)
            from ..oxml.textmap import split_at

            split_at(p, match.end)
            split_at(p, match.start)
            view.refresh()
            runs = [Run(r, match.paragraph) for r in view.runs_in(match.start, match.end)]
            total += self.format_runs(runs, **props)
        return total

    def format_paragraph(self, paragraph: Paragraph, **props) -> None:
        """Paragraph-level formatting revision (``w:pPrChange``).

        Accepts ``style=...`` plus any :class:`docx.text.parfmt.ParagraphFormat`
        attribute, e.g. ``alignment``, ``space_after``, ``left_indent``.
        """
        ppr = paragraph._p.find(qn("w:pPr"))
        baseline = copy.deepcopy(ppr) if ppr is not None else None
        for key, value in props.items():
            if key == "style":
                self.apply_style(paragraph, value)
            elif hasattr(paragraph.paragraph_format, key):
                setattr(paragraph.paragraph_format, key, value)
            else:
                raise RedlineError(f"unknown paragraph property {key!r}")
        edits.record_ppr_change(paragraph._p, baseline, self.ctx)

    # ------------------------------------------------------------------
    # 9. COMMENTS (annotate without changing text)
    # ------------------------------------------------------------------
    def add_comment(
        self,
        paragraph: Paragraph,
        text: str,
        runs: Sequence[Run] | None = None,
        author: str | None = None,
        initials: str | None = None,
    ):
        """Attach a review comment.  Requires python-docx >= 1.2.

        ``author`` overrides the Redliner's own attribution for this one comment,
        so a multi-agent review can sign each note with the agent that raised it.
        """
        if not hasattr(self.document, "add_comment"):  # pragma: no cover
            raise RedlineError("comments require python-docx >= 1.2.0")
        anchors = (
            list(runs) if runs else [Run(r, paragraph) for r in iter_visible_runs(paragraph._p)]
        )
        if not anchors:
            # A struck paragraph has no visible runs, but commenting on the
            # strikeout is exactly what a reviewer wants to see there.
            anchors = [Run(r, paragraph) for r in paragraph._p.iter(qn("w:r"))]
        if not anchors:
            raise RedlineError("cannot anchor a comment to an empty paragraph")
        name = author or self.ctx.author.name
        if initials is None:
            initials = (
                self.ctx.author.initials
                if author is None
                else "".join(word[0] for word in name.split() if word)[:3].upper()
            )
        return self.document.add_comment(runs=anchors, text=text, author=name, initials=initials)

    # ------------------------------------------------------------------
    # settings / output
    # ------------------------------------------------------------------
    def enable_track_changes(self, enabled: bool = True) -> None:
        """Toggle ``<w:trackChanges/>`` so later human edits are recorded too."""
        settings = self.document.settings.element
        existing = settings.find(qn("w:trackChanges"))
        if enabled and existing is None:
            settings.insert(0, make("w:trackChanges"))
        elif not enabled and existing is not None:
            settings.remove(existing)

    def summary(self) -> review.RevisionSummary:
        """Every tracked change currently in the document."""
        total = review.RevisionSummary()
        for root in self._all_roots():
            name = root.tag.split("}")[-1]
            total.revisions.extend(review.summarize(root, part=name).revisions)
        return total

    def accept_all(self) -> Redliner:
        for root in self._all_roots():
            review.accept_all(root)
        return self

    def reject_all(self) -> Redliner:
        for root in self._all_roots():
            review.reject_all(root)
        return self

    def text(self) -> str:
        """Current (all-markup-accepted) body text -- handy for assertions."""
        return "\n".join(paragraph_text(p) for p in self.document.element.body.iter(qn("w:p")))

    def save(self, path: str | Path) -> Path:
        path = Path(path)
        path.parent.mkdir(parents=True, exist_ok=True)
        self.document.save(str(path))
        return path

    # ------------------------------------------------------------------
    def apply_style(self, paragraph: Paragraph, name: str) -> None:
        """Set a paragraph style by name, styleId, or either in any case.

        python-docx maps built-in style names through its own table, so a
        document that writes ``<w:name w:val="Heading 4"/>`` where the built-in
        spelling is ``heading 4`` raises ``KeyError`` even though the style is
        right there in ``styles.xml``.  When its lookup fails we resolve the
        style ourselves and write ``w:pStyle`` directly.
        """
        try:
            paragraph.style = name
            return
        except KeyError:
            pass
        style_id = self._resolve_style_id(name)
        if style_id is None:
            raise RedlineError(
                f"unknown paragraph style {name!r}; this document defines "
                + ", ".join(repr(n) for n in self.paragraph_style_names())
            )
        ppr = get_or_add(paragraph._p, "w:pPr", ())
        if paragraph._p.index(ppr) != 0:
            paragraph._p.remove(ppr)
            paragraph._p.insert(0, ppr)
        pstyle = get_or_add(ppr, "w:pStyle", PPR_ORDER)
        pstyle.set(qn("w:val"), style_id)

    def _style_elements(self):
        root = self.document.styles.element
        for style in root.findall(qn("w:style")):
            if style.get(qn("w:type")) in (None, "paragraph"):
                yield style

    @staticmethod
    def _style_key(value: str) -> str:
        return value.replace(" ", "").replace("-", "").lower()

    def _resolve_style_id(self, name: str) -> str | None:
        wanted = self._style_key(name)
        for style in self._style_elements():
            style_id = style.get(qn("w:styleId")) or ""
            label = style.find(qn("w:name"))
            label = label.get(qn("w:val")) if label is not None else ""
            if wanted in (self._style_key(style_id), self._style_key(label)):
                return style_id
        return None

    def paragraph_style_names(self) -> list[str]:
        """Every paragraph style this document actually defines."""
        names = []
        for style in self._style_elements():
            label = style.find(qn("w:name"))
            names.append(
                label.get(qn("w:val")) if label is not None else style.get(qn("w:styleId"))
            )
        return sorted(n for n in names if n)

    def _clamp(self, paragraph: Paragraph, offset: int) -> int:
        length = len(paragraph_text(paragraph._p))
        if offset < 0:
            offset += length
        return max(0, min(offset, length))

    @staticmethod
    def _reference_rpr(ref: etree._Element) -> etree._Element | None:
        for run in iter_visible_runs(ref):
            rpr = run.find(qn("w:rPr"))
            if rpr is not None:
                return deepcopy_without(copy.deepcopy(rpr), ("w:rPrChange",))
        return None


#: Wrappers whose presence makes a paragraph unsafe to copy into a move revision.
_CONTENT_REVISIONS = ("w:ins", "w:del", "w:moveFrom", "w:moveTo")


def paragraph_has_revisions(p: etree._Element) -> bool:
    """True when the paragraph already carries insert/delete/move markup."""
    wanted = {qn(tag) for tag in _CONTENT_REVISIONS}
    return any(el.tag in wanted for el in p.iter())


def _place(
    node: etree._Element, after: etree._Element | None, before: etree._Element | None
) -> None:
    if after is not None:
        after.addnext(node)
    elif before is not None:
        before.addprevious(node)
    else:  # pragma: no cover - callers always supply one
        raise RedlineError("relocate_paragraph needs an `after` or a `before` anchor")


def _accepted_copy(p: etree._Element) -> etree._Element:
    """A deep copy of ``p`` with its existing revisions resolved as accepted.

    The moved block should arrive at its destination reading the way it will
    read once the redline is accepted -- not carrying the source's strikeouts.
    """
    holder = make("w:body")
    holder.append(copy.deepcopy(p))
    review.accept_all(holder)
    if not len(holder):  # the paragraph resolved away entirely
        raise RedlineError("cannot move a paragraph whose content is already deleted")
    clone = holder[0]
    holder.remove(clone)
    return clone


def _highlight(value):
    """Coerce ``"YELLOW"`` / ``"bright_green"`` to a ``WD_COLOR_INDEX`` member."""
    from docx.enum.text import WD_COLOR_INDEX

    if value is None or not isinstance(value, str):
        return value
    key = value.strip().upper().replace(" ", "_").replace("-", "_")
    try:
        return WD_COLOR_INDEX[key]
    except KeyError:
        raise RedlineError(
            f"unknown highlight {value!r}; expected one of "
            + ", ".join(sorted(m.name for m in WD_COLOR_INDEX if m.name))
        ) from None
