# This Source Code Form is subject to the terms of the Mozilla Public
# License, v. 2.0. If a copy of the MPL was not distributed with this
# file, You can obtain one at https://mozilla.org/MPL/2.0/.
#
# Copyright (c) 2026 Aswanth B S

"""Declarative edit plans.

A redline is often produced by something that is not Python -- a review
service, a rules engine, an LLM.  This module accepts that plan as plain
JSON-able dicts and applies it, so the caller never has to touch OOXML::

    apply_operations(
        rl,
        [
            {"op": "replace_text", "old": "thirty (30) days", "new": "forty-five (45) days"},
            {"op": "insert_paragraph_after", "match": "3.4", "text": "3.5  Audit rights ..."},
            {"op": "delete_paragraph", "match": "Late Payment"},
        ],
    )
"""

from __future__ import annotations

from collections.abc import Iterable
from dataclasses import dataclass
from typing import Any

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from .redline import RedlineError, Redliner

#: op name -> required keys
SCHEMA: dict[str, tuple[str, ...]] = {
    "insert_text": ("match", "text"),
    "delete_text": ("match",),
    "replace_text": ("old", "new"),
    "update_paragraph": ("match", "text"),
    "insert_paragraph": ("match", "text"),
    "delete_paragraph": ("match",),
    "move_paragraph": ("match", "after"),
    "insert_row": ("table",),
    "delete_row": ("table", "row"),
    "update_cell": ("table", "row", "col", "text"),
    "format_text": ("match",),
    "format_paragraph": ("match",),
    "comment": ("match", "text"),
}

_ALIGNMENTS = dict(WD_ALIGN_PARAGRAPH.__members__.items())


@dataclass
class OpResult:
    op: str
    applied: int
    detail: str = ""


def validate(operations: Iterable[dict]) -> list[str]:
    """Return a list of human-readable problems; empty means the plan is valid."""
    problems: list[str] = []
    for index, spec in enumerate(operations):
        name = spec.get("op")
        if name not in SCHEMA:
            problems.append(f"[{index}] unknown op {name!r}; expected one of {sorted(SCHEMA)}")
            continue
        for key in SCHEMA[name]:
            if key not in spec:
                problems.append(f"[{index}] op {name!r} is missing required key {key!r}")
    return problems


def apply_operations(
    rl: Redliner, operations: Iterable[dict], strict: bool = True
) -> list[OpResult]:
    """Apply an edit plan in order.  With ``strict``, a no-match op is an error."""
    operations = list(operations)
    problems = validate(operations)
    if problems:
        raise RedlineError("invalid operation plan:\n  " + "\n  ".join(problems))
    results: list[OpResult] = []
    for index, spec in enumerate(operations):
        result = _dispatch(rl, dict(spec))
        if strict and result.applied == 0:
            raise RedlineError(f"[{index}] op {spec['op']!r} matched nothing: {spec}")
        results.append(result)
    return results


def _dispatch(rl: Redliner, spec: dict) -> OpResult:
    name = spec.pop("op")
    handler = _HANDLERS[name]
    return handler(rl, spec)


# ---------------------------------------------------------------------------


def _paragraph(rl: Redliner, spec: dict, key: str = "match"):
    needle = spec[key]
    found = rl.find_paragraphs(contains=needle) or rl.find_paragraphs(regex=needle)
    if not found:
        raise RedlineError(f"no paragraph contains {needle!r}")
    occurrence = spec.get("occurrence", 0)
    return found[occurrence]


def _op_insert_text(rl, spec):
    where = spec.get("where", "after")
    count = spec.get("count", 1)
    if where == "before":
        n = rl.insert_text_before(spec["match"], spec["text"], count=count)
    elif where == "end":
        rl.append_text(_paragraph(rl, spec), spec["text"])
        n = 1
    else:
        n = rl.insert_text_after(spec["match"], spec["text"], count=count)
    return OpResult("insert_text", n, f"{where} {spec['match']!r}")


def _op_delete_text(rl, spec):
    n = rl.delete_matching(spec["match"], regex=spec.get("regex", False), count=spec.get("count"))
    return OpResult("delete_text", n, spec["match"])


def _op_replace_text(rl, spec):
    n = rl.replace_text(
        spec["old"],
        spec["new"],
        regex=spec.get("regex", False),
        ignore_case=spec.get("ignore_case", False),
        count=spec.get("count", 1),
    )
    return OpResult("replace_text", n, f"{spec['old']!r} -> {spec['new']!r}")


def _op_update_paragraph(rl, spec):
    n = rl.set_paragraph_text(_paragraph(rl, spec), spec["text"])
    return OpResult("update_paragraph", max(n, 1), spec["match"])


def _op_insert_paragraph(rl, spec):
    reference = _paragraph(rl, spec)
    where = spec.get("where", "after")
    if where == "before":
        rl.insert_paragraph_before(reference, spec["text"], style=spec.get("style"))
    else:
        rl.insert_paragraph_after(reference, spec["text"], style=spec.get("style"))
    return OpResult("insert_paragraph", 1, f"{where} {spec['match']!r}")


def _op_delete_paragraph(rl, spec):
    if spec.get("all"):
        targets = rl.find_paragraphs(contains=spec["match"])
        return OpResult("delete_paragraph", rl.delete_paragraphs(targets), spec["match"])
    rl.delete_paragraph(_paragraph(rl, spec))
    return OpResult("delete_paragraph", 1, spec["match"])


def _op_move_paragraph(rl, spec):
    rl.move_paragraph(_paragraph(rl, spec), _paragraph(rl, spec, "after"))
    return OpResult("move_paragraph", 1, f"{spec['match']!r} after {spec['after']!r}")


def _op_insert_row(rl, spec):
    table = rl.tables()[spec["table"]]
    rl.insert_table_row(table, index=spec.get("row"), values=spec.get("values", []))
    return OpResult("insert_row", 1, f"table {spec['table']}")


def _op_delete_row(rl, spec):
    rl.delete_table_row(rl.tables()[spec["table"]], spec["row"])
    return OpResult("delete_row", 1, f"table {spec['table']} row {spec['row']}")


def _op_update_cell(rl, spec):
    table = rl.tables()[spec["table"]]
    cell = table.cell(spec["row"], spec["col"])
    n = rl.set_cell_text(cell, spec["text"])
    return OpResult(
        "update_cell", max(n, 1), f"table {spec['table']} ({spec['row']},{spec['col']})"
    )


def _op_format_text(rl, spec):
    props = _format_props(spec, exclude={"match", "regex", "count"})
    n = rl.format_matching(
        spec["match"], regex=spec.get("regex", False), count=spec.get("count"), **props
    )
    return OpResult("format_text", n, spec["match"])


def _op_format_paragraph(rl, spec):
    props = _format_props(spec, exclude={"match", "occurrence"}, paragraph=True)
    rl.format_paragraph(_paragraph(rl, spec), **props)
    return OpResult("format_paragraph", 1, spec["match"])


def _op_comment(rl, spec):
    rl.add_comment(_paragraph(rl, spec), spec["text"])
    return OpResult("comment", 1, spec["match"])


def _format_props(spec: dict, exclude: set[str], paragraph: bool = False) -> dict[str, Any]:
    """Coerce JSON-friendly values into python-docx types."""
    props: dict[str, Any] = {}
    for key, value in spec.items():
        if key in exclude:
            continue
        if key == "alignment" and isinstance(value, str):
            value = _ALIGNMENTS[value.upper()]
        elif (
            key in ("size", "space_before", "space_after") and isinstance(value, (int, float))
        ) or (
            key in ("left_indent", "right_indent", "first_line_indent")
            and isinstance(value, (int, float))
        ):
            value = Pt(value)
        props[key] = value
    return props


_HANDLERS = {
    "insert_text": _op_insert_text,
    "delete_text": _op_delete_text,
    "replace_text": _op_replace_text,
    "update_paragraph": _op_update_paragraph,
    "insert_paragraph": _op_insert_paragraph,
    "delete_paragraph": _op_delete_paragraph,
    "move_paragraph": _op_move_paragraph,
    "insert_row": _op_insert_row,
    "delete_row": _op_delete_row,
    "update_cell": _op_update_cell,
    "format_text": _op_format_text,
    "format_paragraph": _op_format_paragraph,
    "comment": _op_comment,
}
