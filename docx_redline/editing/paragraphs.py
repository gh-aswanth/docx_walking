# This Source Code Form is subject to the terms of the Mozilla Public
# License, v. 2.0. If a copy of the MPL was not distributed with this
# file, You can obtain one at https://mozilla.org/MPL/2.0/.
#
# Copyright (c) 2026 Aswanth B S

"""Stable paragraph addressing, match normalisation, and plan verification.

The rest of this package addresses the document the way a lawyer does -- by
clause number and quoted phrase.  That is the right model for a reviewer and the
wrong one for a machine: clause numbers move, sections get renumbered, and a
document with no numbering at all has nothing to address.  This module adds the
other half of the contract:

* :func:`fold` -- index-preserving normalisation.  A model that quotes
  ``"twelve (12) months"`` with a straight apostrophe still matches text Word
  stored with a smart quote.  Every mapping is exactly one character to one
  character, which is what makes an offset computed on the folded string valid
  on the raw string it came from.
* :class:`ParagraphIndex` -- every ``w:p`` in scope gets a stable integer id in
  document order.  An edit needs only that id plus an exact quoted span; it
  never needs a clause number, so routing quality and edit correctness stop
  being the same problem.
* :func:`fingerprint` / :func:`verify_plan` -- binds a plan to the document
  version it was computed against, so a stale plan is refused rather than
  silently landing on text that moved.

Edits are located against a snapshot of every paragraph and only then applied,
right to left, so no applied edit invalidates a pending offset.  Locating and
applying must stay separate passes: locate lazily and the second edit to touch a
paragraph sees XML the first one already rewrote, and the miss gets misdiagnosed
as a parse problem rather than the collision it is.

Review notes are the exception.  They are anchors, not rewrites, so they are
validated up front but re-located against the finished text and applied last --
the same order :mod:`docx_redline.planning.actions` uses, and for the same reason: a
comment should point at the paragraph as it ends up, not as it started.
"""

from __future__ import annotations

import hashlib
import unicodedata
from collections.abc import Iterable, Sequence
from dataclasses import dataclass, field, replace
from enum import Enum

from docx.text.paragraph import Paragraph
from docx.text.run import Run
from lxml import etree

from ..errors import RedlineError, StalePlanError
from ..oxml import edits as _edits
from ..oxml.ns import qn
from ..oxml.textmap import paragraph_text, split_range
from ..structure.clauses import Clause, ClauseTree

__all__ = [
    "ApplyReport",
    "EditResult",
    "ParagraphIndex",
    "ParagraphRef",
    "RedlineEdit",
    "Rejection",
    "ReviewNote",
    "fingerprint",
    "fold",
    "load_edits",
    "validate_edits",
    "verify_plan",
]


# ---------------------------------------------------------------------------
# normalisation
# ---------------------------------------------------------------------------

#: One character in, one character out.  Nothing may be added here that changes
#: the length of the string -- the whole point is that offsets survive the fold.
# fmt: off
_CHAR_FOLD: dict[str, str] = {
    # quotes
    "‘": "'", "’": "'", "‚": "'", "‛": "'",
    "′": "'", "´": "'", "ʼ": "'",
    "“": '"', "”": '"', "„": '"', "‟": '"', "″": '"',
    # dashes and hyphens
    "‐": "-", "‑": "-", "‒": "-", "–": "-",
    "—": "-", "―": "-", "−": "-", "⁃": "-",
    # spaces
    " ": " ", " ": " ", " ": " ", " ": " ", " ": " ",
    " ": " ", " ": " ", " ": " ", " ": " ", " ": " ",
    " ": " ", " ": " ", " ": " ", " ": " ", " ": " ",
    "　": " ", "​": " ", "﻿": " ", "\t": " ",
}
# fmt: on


def fold(text: str) -> str:
    """Normalise for matching.  Guaranteed to preserve string length.

    NFKC is applied one character at a time and kept only when it maps to a
    single character, so ``"Ａ"`` folds to ``"A"`` while the ``"ﬁ"`` ligature is
    left alone rather than expanding and shifting every offset after it.
    """
    out: list[str] = []
    for ch in text:
        mapped = _CHAR_FOLD.get(ch)
        if mapped is None:
            mapped = unicodedata.normalize("NFKC", ch)
            if len(mapped) != 1:
                mapped = ch
        out.append(mapped)
    return "".join(out)


#: A target shorter than this that occurs more than once in its paragraph is
#: rejected as ambiguous rather than resolved to the first hit.  Short fragments
#: ("the Agreement", "30 days") are what a model quotes when it has not read
#: carefully; a long repeated quote is far more likely to be deliberate.
AMBIGUITY_MIN_CHARS = 25


# ---------------------------------------------------------------------------
# results
# ---------------------------------------------------------------------------
class Rejection(str, Enum):  # noqa: UP042 -- StrEnum changes str(); .value is public API
    PARAGRAPH_NOT_FOUND = "paragraph_not_found"
    TARGET_NOT_FOUND = "target_not_found"
    TARGET_AMBIGUOUS = "target_ambiguous"
    TARGET_ALREADY_STRUCK = "target_already_struck"
    SPAN_CONFLICT = "span_conflict"
    EMPTY_TARGET = "empty_target"


_EXPLAIN = {
    Rejection.TARGET_NOT_FOUND: (
        "quoted span does not match the paragraph; re-quote it, do not fuzzy-match"
    ),
    Rejection.TARGET_AMBIGUOUS: (
        f"span occurs more than once and is under {AMBIGUITY_MIN_CHARS} characters; "
        "quote more context, set occurrence=n, or set occurrence=0 for all of them"
    ),
    Rejection.TARGET_ALREADY_STRUCK: (
        "span exists only inside an existing tracked deletion; it is already on its way out"
    ),
    Rejection.SPAN_CONFLICT: (
        "an earlier edit already claimed this span; send both findings to reconciliation"
    ),
    Rejection.EMPTY_TARGET: "target is empty or whitespace only",
}


@dataclass
class RedlineEdit:
    """A proposed tracked change.

    ``target`` must be an exact quoted span from the paragraph.  Never a
    character offset, never a line number -- offsets do not survive a reparse,
    quoted text does, and a quote that no longer matches is a signal the document
    moved on rather than a bug to route around.

    ``occurrence`` is 1-based; ``0`` means every occurrence in the paragraph.
    An empty ``replacement`` is a pure deletion.
    """

    para_id: int
    target: str
    replacement: str = ""
    rationale: str = ""
    agent: str = ""
    severity: str = "medium"
    occurrence: int = 1
    insertion_first: bool = False

    def __post_init__(self) -> None:
        if self.occurrence < 0:
            raise ValueError("occurrence is 1-based, or 0 for every occurrence")


@dataclass
class ReviewNote:
    """A flag rather than an edit.  Rendered as a Word comment on the span."""

    para_id: int
    target: str
    body: str
    agent: str = ""
    severity: str = "medium"
    occurrence: int = 1

    def __post_init__(self) -> None:
        if self.occurrence < 0:
            raise ValueError("occurrence is 1-based, or 0 for every occurrence")


@dataclass
class EditResult:
    item: RedlineEdit | ReviewNote
    applied: bool
    reason: Rejection | None = None
    detail: str = ""
    spans: int = 0

    def __str__(self) -> str:
        head = "APPLIED " if self.applied else "REJECTED"
        who = f"[{self.item.agent}] " if self.item.agent else ""
        why = self.reason.value if self.reason is not None else "rejected"
        tail = "" if self.applied else f"  <- {why}: {self.detail}"
        return f"{head} {who}p{self.item.para_id}: {self.item.target[:52]!r}{tail}"


@dataclass
class ApplyReport:
    results: list[EditResult] = field(default_factory=list)

    @property
    def applied(self) -> list[EditResult]:
        return [r for r in self.results if r.applied]

    @property
    def rejected(self) -> list[EditResult]:
        return [r for r in self.results if not r.applied]

    def summary(self) -> str:
        lines = [f"{len(self.applied)} applied, {len(self.rejected)} rejected"]
        lines += [f"  {r}" for r in self.results]
        return "\n".join(lines)

    def to_dict(self) -> dict:
        return {
            "applied": len(self.applied),
            "rejected": len(self.rejected),
            "results": [
                {
                    "para_id": r.item.para_id,
                    "target": r.item.target,
                    "agent": r.item.agent,
                    "applied": r.applied,
                    "reason": r.reason.value if r.reason else None,
                    "detail": r.detail,
                    "spans": r.spans,
                }
                for r in self.results
            ],
        }


# ---------------------------------------------------------------------------
# the index
# ---------------------------------------------------------------------------
@dataclass(frozen=True)
class ParagraphRef:
    """One addressable paragraph."""

    para_id: int
    text: str
    style: str | None = None
    numbered: bool = False
    in_table: bool = False
    table_index: int | None = None
    clause_label: str | None = None
    level: int = 0

    @property
    def is_empty(self) -> bool:
        return not self.text.strip()


class ParagraphIndex:
    """A stable integer address space over a :class:`~docx_redline.editing.redline.Redliner`.

    Ids are positions in ``rl.paragraphs()`` and stay valid for as long as no
    paragraph is inserted or deleted.  Text-level edits (replace / delete /
    comment) never move them; :meth:`refresh` re-indexes after structural work.
    """

    def __init__(self, rl, include_tables: bool = True) -> None:
        self.rl = rl
        self.include_tables = include_tables
        self.refs: list[ParagraphRef] = []
        self._paragraphs: list[Paragraph] = []
        self._snapshot: list[str] = []
        self._folded: list[str] = []
        self._claimed: dict[int, list[tuple[int, int]]] = {}
        # Kept alive on purpose: these hold the lxml proxies whose id() we key on.
        self._tables: list = []
        self._tree: ClauseTree | None = None
        self.refresh()

    # -- construction ---------------------------------------------------
    def refresh(self) -> ParagraphIndex:
        """Re-read the document and drop every span claim."""
        self._paragraphs = self.rl.paragraphs(include_tables=self.include_tables)
        self._tables = self.rl.tables()
        table_at = {id(t._tbl): i for i, t in enumerate(self._tables)}

        self._tree = ClauseTree(self.rl.document.element.body)
        clause_at: dict[int, Clause] = {id(c.p): c for c in self._tree}

        self.refs = []
        self._snapshot = []
        self._folded = []
        self._claimed = {}

        current: Clause | None = None
        for para_id, para in enumerate(self._paragraphs):
            p = para._p
            text = paragraph_text(p)
            self._snapshot.append(text)
            self._folded.append(fold(text))
            clause = clause_at.get(id(p))
            if clause is not None:
                current = clause
            tbl = self._owning_table(p)
            self.refs.append(
                ParagraphRef(
                    para_id=para_id,
                    text=text,
                    style=self._style_of(p),
                    numbered=p.find(f"{qn('w:pPr')}/{qn('w:numPr')}") is not None,
                    in_table=tbl is not None,
                    table_index=table_at.get(id(tbl)) if tbl is not None else None,
                    clause_label=current.label if current is not None else None,
                    level=clause.level if clause is not None else 0,
                )
            )
        return self

    @staticmethod
    def _owning_table(p: etree._Element) -> etree._Element | None:
        for ancestor in p.iterancestors():
            if ancestor.tag == qn("w:tbl"):
                return ancestor
        return None

    @staticmethod
    def _style_of(p: etree._Element) -> str | None:
        style = p.find(f"{qn('w:pPr')}/{qn('w:pStyle')}")
        return style.get(qn("w:val")) if style is not None else None

    # -- access ---------------------------------------------------------
    def __len__(self) -> int:
        return len(self.refs)

    def __iter__(self):
        return iter(self.refs)

    def __getitem__(self, para_id: int) -> ParagraphRef:
        return self.refs[para_id]

    @property
    def clauses(self) -> list[Clause]:
        """Every numbered clause the document carries, in body order."""
        return list(self._tree) if self._tree is not None else []

    def paragraph(self, para_id: int) -> Paragraph:
        """The live ``python-docx`` paragraph behind an id."""
        if not 0 <= para_id < len(self._paragraphs):
            raise RedlineError(f"para_id {para_id} out of range (0-{len(self.refs) - 1})")
        return self._paragraphs[para_id]

    def find(self, needle: str, ignore_case: bool = False) -> list[int]:
        """Ids of every paragraph containing ``needle``, folded on both sides."""
        probe = fold(needle)
        probe = probe.lower() if ignore_case else probe
        out = []
        for ref in self.refs:
            hay = self._folded[ref.para_id]
            hay = hay.lower() if ignore_case else hay
            if probe in hay:
                out.append(ref.para_id)
        return out

    # -- fingerprint ----------------------------------------------------
    def fingerprint(self) -> str:
        """Binds a redline plan to the document version it was computed against."""
        blob = "\n".join(self._snapshot).encode("utf-8")
        return hashlib.sha256(blob).hexdigest()[:16]

    # -- rendering ------------------------------------------------------
    def render(
        self,
        para_ids: Iterable[int] | None = None,
        with_clause_labels: bool = True,
        skip_empty: bool = True,
    ) -> str:
        """Numbered-paragraph rendering -- the cacheable prefix a model reads.

        Every line starts with the ``[id]`` an edit addresses, so the model can
        only ever point at a paragraph that exists.
        """
        wanted = set(para_ids) if para_ids is not None else None
        out: list[str] = []
        last_label: str | None = None
        for ref in self.refs:
            if wanted is not None and ref.para_id not in wanted:
                continue
            if skip_empty and ref.is_empty:
                continue
            if with_clause_labels and ref.clause_label != last_label:
                last_label = ref.clause_label
                if last_label:
                    out.append(f"\n<<clause {last_label}>>")
            marker = f" [table {ref.table_index}]" if ref.in_table else ""
            out.append(f"[{ref.para_id}]{marker} {ref.text}")
        return "\n".join(out).strip()

    def manifest(self) -> str:
        """Routing input: one line per clause, id range included."""
        rows: list[str] = []
        spans: dict[str, list[int]] = {}
        for ref in self.refs:
            if ref.clause_label:
                spans.setdefault(ref.clause_label, []).append(ref.para_id)
        for clause in self._tree or []:
            ids = spans.get(clause.label)
            if not ids:
                continue
            title = clause.title or clause.body[:60]
            rows.append(f"{clause.label}\tparas {min(ids)}-{max(ids)}\t{title}")
        return "\n".join(rows)

    # -- locating -------------------------------------------------------
    def locate(
        self, para_id: int, target: str, occurrence: int = 1
    ) -> list[tuple[int, int]] | Rejection:
        """Resolve a quoted span to character offsets, or say why it cannot be.

        Offsets are in the same flat coordinate space :mod:`docx_redline.oxml.edits`
        uses, so text inside a hyperlink, a ``w:ins`` or a content control
        resolves normally.  Text already inside a ``w:del`` does not -- it is not
        part of the paragraph's current reading.
        """
        if not target.strip():
            return Rejection.EMPTY_TARGET
        if not 0 <= para_id < len(self.refs):
            return Rejection.PARAGRAPH_NOT_FOUND

        haystack = self._folded[para_id]
        needle = fold(target)
        hits: list[int] = []
        cursor = haystack.find(needle)
        while cursor != -1:
            hits.append(cursor)
            cursor = haystack.find(needle, cursor + 1)

        if not hits:
            if needle in fold(self._all_text(self._paragraphs[para_id]._p)):
                return Rejection.TARGET_ALREADY_STRUCK
            return Rejection.TARGET_NOT_FOUND
        if occurrence == 0:
            return [(h, h + len(needle)) for h in hits]
        if occurrence > len(hits):
            return Rejection.TARGET_NOT_FOUND
        if len(hits) > 1 and occurrence == 1 and len(needle) < AMBIGUITY_MIN_CHARS:
            return Rejection.TARGET_AMBIGUOUS
        start = hits[occurrence - 1]
        return [(start, start + len(needle))]

    @staticmethod
    def _all_text(p: etree._Element) -> str:
        """Every character in the paragraph, struck-out content included."""
        parts = []
        for node in p.iter():
            local = node.tag.split("}")[-1] if isinstance(node.tag, str) else ""
            if local in ("t", "delText") and node.text:
                parts.append(node.text)
            elif local in ("tab", "br", "cr"):
                parts.append(" ")
        return "".join(parts)

    def _explain(self, why: Rejection) -> str:
        if why is Rejection.PARAGRAPH_NOT_FOUND:
            return f"para_id out of range (0-{len(self.refs) - 1})"
        return _EXPLAIN.get(why, "")

    # -- claims ---------------------------------------------------------
    def _free(self, para_id: int, span: tuple[int, int]) -> bool:
        return not any(span[0] < hi and lo < span[1] for lo, hi in self._claimed.get(para_id, []))

    def _claim(self, para_id: int, span: tuple[int, int]) -> None:
        self._claimed.setdefault(para_id, []).append(span)

    # -- applying -------------------------------------------------------
    def apply(self, items: Sequence[RedlineEdit | ReviewNote]) -> ApplyReport:
        """Verify every item, then apply the survivors.

        Nothing is written until every item has been located, so a plan that is
        half wrong cannot leave the document half edited by the time it is found
        out.  Content edits go in right to left; review notes are re-located
        against the finished text and applied last.
        """
        report = ApplyReport()
        staged: list[tuple[int, int, int, RedlineEdit]] = []
        notes: list[tuple[ReviewNote, EditResult]] = []

        for item in items:
            found = self.locate(item.para_id, item.target, item.occurrence)

            if isinstance(found, Rejection):
                # Text an earlier pass struck is still worth commenting on --
                # often the comment is *about* the strike. Anchor a note at the
                # paragraph rather than dropping the reviewer's finding.
                if found is Rejection.TARGET_ALREADY_STRUCK and isinstance(item, ReviewNote):
                    spans: list[tuple[int, int]] = []
                else:
                    report.results.append(EditResult(item, False, found, self._explain(found)))
                    continue
            else:
                spans = found

            if isinstance(item, ReviewNote):
                # An anchor cannot invalidate an offset, so notes never claim.
                result = EditResult(item, True, spans=len(spans))
                report.results.append(result)
                notes.append((item, result))
                continue

            if any(not self._free(item.para_id, span) for span in spans):
                report.results.append(
                    EditResult(
                        item, False, Rejection.SPAN_CONFLICT, self._explain(Rejection.SPAN_CONFLICT)
                    )
                )
                continue
            for span in spans:
                self._claim(item.para_id, span)
                staged.append((item.para_id, span[0], span[1], item))
            report.results.append(EditResult(item, True, spans=len(spans)))

        for para_id, start, end, edit in sorted(staged, reverse=True, key=lambda s: s[:2]):
            _edits.replace_range(
                self._paragraphs[para_id]._p,
                start,
                end,
                edit.replacement,
                self.rl.ctx,
                edit.insertion_first,
            )

        if notes:
            self._reindex_text()
            for note, result in notes:
                result.detail = self._attach(note)

        if staged or notes:
            self._reindex_text()
        return report

    def _reindex_text(self) -> None:
        """Re-read paragraph text in place; ids and claims are untouched."""
        for para_id, para in enumerate(self._paragraphs):
            text = paragraph_text(para._p)
            self._snapshot[para_id] = text
            self._folded[para_id] = fold(text)
            self.refs[para_id] = replace(self.refs[para_id], text=text)

    def _attach(self, note: ReviewNote) -> str:
        """Anchor a note to its span in the *finished* text.

        Re-located rather than replayed: an edit applied a moment ago may have
        struck the very words the note quotes, and a comment on a paragraph is
        far more use than a comment on a stale offset.
        """
        paragraph = self._paragraphs[note.para_id]
        located = self.locate(note.para_id, note.target, note.occurrence)
        runs: list[Run] | None = None
        detail = ""
        if isinstance(located, Rejection):
            detail = f"anchored to the whole paragraph ({located.value})"
        else:
            elements: list = []
            for start, end in sorted(located, reverse=True):
                view = split_range(paragraph._p, start, end)
                elements.extend(view.runs_in(start, end))
            runs = [Run(element, paragraph) for element in elements] or None
            if runs is None:
                detail = "anchored to the whole paragraph (span held no runs)"
        self.rl.add_comment(
            paragraph,
            note.body,
            runs=runs,
            author=note.agent or None,
        )
        return detail


# ---------------------------------------------------------------------------
# plan verification
# ---------------------------------------------------------------------------
def fingerprint(source) -> str:
    """Fingerprint a ``ParagraphIndex``, a ``Redliner``, or a path to a ``.docx``."""
    if isinstance(source, ParagraphIndex):
        return source.fingerprint()
    if hasattr(source, "paragraphs") and hasattr(source, "ctx"):
        return ParagraphIndex(source).fingerprint()
    from .redline import Redliner

    return ParagraphIndex(Redliner(source, track_changes=False)).fingerprint()


def verify_plan(source, plan_fingerprint: str) -> str:
    """Guard against applying a v3 redline to a v4 document."""
    current = fingerprint(source)
    if current != plan_fingerprint:
        raise StalePlanError(
            f"document moved on: plan was computed against {plan_fingerprint}, "
            f"document is now {current}. Re-run the review."
        )
    return current


# ---------------------------------------------------------------------------
# model output
# ---------------------------------------------------------------------------
_EDIT_KEYS = {
    "kind",
    "para_id",
    "target",
    "replacement",
    "rationale",
    "agent",
    "severity",
    "occurrence",
    "insertion_first",
}
_NOTE_KEYS = {"kind", "para_id", "target", "body", "agent", "severity", "occurrence"}
SEVERITIES = ("low", "medium", "high", "critical")


def validate_edits(payload: Iterable[dict]) -> list[str]:
    """Schema check with no document access -- safe to run on raw model output."""
    problems: list[str] = []
    for index, item in enumerate(payload):
        where = f"#{index}"
        kind = item.get("kind", "edit")
        if kind not in ("edit", "note"):
            problems.append(f"{where}: unknown kind {kind!r}, expected 'edit' or 'note'")
            continue
        allowed = _EDIT_KEYS if kind == "edit" else _NOTE_KEYS
        required = ("para_id", "target") + (() if kind == "edit" else ("body",))
        for key in required:
            if item.get(key) in (None, ""):
                problems.append(f"{where}: {kind!r} requires {key!r}")
        for key in item:
            if key not in allowed:
                problems.append(f"{where}: {kind!r} does not accept {key!r}")
        if kind == "edit" and "replacement" not in item:
            problems.append(f"{where}: 'edit' requires 'replacement' (use \"\" to delete)")
        severity = item.get("severity")
        if severity is not None and severity not in SEVERITIES:
            problems.append(f"{where}: severity {severity!r} not in {SEVERITIES}")
        try:
            if int(item.get("occurrence", 1)) < 0:
                problems.append(f"{where}: occurrence must be >= 0")
        except (TypeError, ValueError):
            problems.append(f"{where}: occurrence must be an integer")
    return problems


def load_edits(payload: Iterable[dict], strict: bool = True) -> list[RedlineEdit | ReviewNote]:
    """Coerce structured model JSON into edits.

    ``strict`` raises on the first malformed batch rather than silently dropping
    items -- a dropped finding is a review that quietly did less than it claimed.
    """
    payload = list(payload)
    problems = validate_edits(payload)
    if problems and strict:
        raise RedlineError("malformed edit payload:\n  " + "\n  ".join(problems))
    bad = {int(p.split(":", 1)[0].lstrip("#")) for p in problems}

    out: list[RedlineEdit | ReviewNote] = []
    for index, item in enumerate(payload):
        if index in bad:
            continue
        if item.get("kind", "edit") == "edit":
            out.append(
                RedlineEdit(
                    para_id=int(item["para_id"]),
                    target=item["target"],
                    replacement=item.get("replacement", ""),
                    rationale=item.get("rationale", ""),
                    agent=item.get("agent", ""),
                    severity=item.get("severity", "medium"),
                    occurrence=int(item.get("occurrence", 1)),
                    insertion_first=bool(item.get("insertion_first", False)),
                )
            )
        else:
            out.append(
                ReviewNote(
                    para_id=int(item["para_id"]),
                    target=item["target"],
                    body=item["body"],
                    agent=item.get("agent", ""),
                    severity=item.get("severity", "medium"),
                    occurrence=int(item.get("occurrence", 1)),
                )
            )
    return out
