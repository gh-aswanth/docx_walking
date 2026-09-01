"""02 · Every Redliner constructor option.

Redliner(source, author="Redline", initials=None, date=None,
         track_changes=True, scope=("body", "headers", "footers", "notes"))
"""

import docx
from _shared import SOURCE, banner, section

from docx_redline import Redliner

banner("02 · Redliner options")

section("source: a path, a file object, or an existing Document")
print("path     ->", len(Redliner(SOURCE).paragraphs()), "paragraphs")
with open(SOURCE, "rb") as handle:
    print("file-like->", len(Redliner(handle).paragraphs()), "paragraphs")
print("Document ->", len(Redliner(docx.Document(SOURCE)).paragraphs()), "paragraphs")

section("author / initials — attribution stamped on every revision")
rl = Redliner(SOURCE, author="Jordan Blake")
rl.replace_text("thirty (30) days", "forty-five (45) days")
print("derived initials:", rl.ctx.author.initials)
rl = Redliner(SOURCE, author="Jordan Blake", initials="JBQ")
print("explicit initials:", rl.ctx.author.initials)

section("date — pin it and the output is byte-reproducible")
one = Redliner(SOURCE, author="A", date="2026-01-01T00:00:00Z")
two = Redliner(SOURCE, author="A", date="2026-01-01T00:00:00Z")
for r in (one, two):
    r.replace_text("thirty (30) days", "forty-five (45) days")
print("same stamp     :", one.ctx.author.date == two.ctx.author.date, one.ctx.author.date)
print("floating stamp :", Redliner(SOURCE).ctx.author.date[:4] + "-...")

section("track_changes — also set Word's own toggle in settings.xml")
for flag in (True, False):
    rl = Redliner(SOURCE, track_changes=flag)
    settings = rl.document.settings.element
    on = settings.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}trackChanges")
    print(f"track_changes={flag!s:<5} -> <w:trackChanges/> present: {on is not None}")

section("scope — which stories document-wide operations touch")
# The sample contract has no headers, footers or notes, so every scope sees the
# same 92 body paragraphs. On a document that has them, narrowing the scope is
# what stops a document-wide replace_text rewriting the running header too.
for scope in [("body",), ("body", "headers", "footers"), ("body", "headers", "footers", "notes")]:
    rl = Redliner(SOURCE, scope=scope)
    print(f"{scope!s:<44} {len(rl.paragraphs()):>3} paragraphs in scope")

section("include_tables — table cells are paragraphs too")
rl = Redliner(SOURCE)
print("with tables   :", len(rl.paragraphs(include_tables=True)))
print("without tables:", len(rl.paragraphs(include_tables=False)))
