import io
import sys
from pathlib import Path

import docx
import pytest

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx_redline import Redliner
from docx_redline.oxml.ns import qn
from docx_redline.oxml.textmap import paragraph_text

DATE = "2026-01-01T00:00:00Z"

#: Test fixtures live beside the examples that document them, so a change to
#: the sample contract shows up in both places at once.
DATA = Path(__file__).resolve().parent.parent / "examples" / "data"
SAMPLE = DATA / "Sample_Software_License_Agreement.docx"
PLAN = DATA / "action_items.json"


def texts(document) -> list[str]:
    """Body paragraph text, ignoring anything already struck out."""
    return [paragraph_text(p) for p in document.element.body.iter(qn("w:p"))]


def roundtrip(rl: Redliner) -> io.BytesIO:
    buf = io.BytesIO()
    rl.document.save(buf)
    buf.seek(0)
    return buf


def document_xml(rl: Redliner) -> str:
    """The saved ``word/document.xml`` as text -- for asserting on raw markup."""
    import zipfile

    with zipfile.ZipFile(roundtrip(rl)) as zf:
        return zf.read("word/document.xml").decode()


def resolved(rl: Redliner, how: str) -> list[str]:
    """Save, reopen, then accept or reject everything -- returns body text."""
    other = Redliner(roundtrip(rl), track_changes=False)
    getattr(other, f"{how}_all")()
    return texts(other.document)


CONTRACT = [
    ("1. Definitions", True),
    ('1.1  "Services" means the hosted platform.', False),
    ('1.2  "Order Form" means the ordering document.', False),
    ("2. Fees and Payment", True),
    ("2.1  Fees. Customer shall pay the fees in the Order Form.", False),
    ("2.2  Invoicing. Payment is due within thirty (30) days.", False),
    ("2.3  Taxes. Fees exclude applicable taxes.", False),
    ("3. Term and Termination", True),
    ("3.1  Term. This Agreement runs for the Subscription Term.", False),
    ("3.2  Renewal. The term renews automatically.", False),
    ("4. General", True),
    ("4.1  Governing Law. Delaware law governs, as noted in Section 2.2.", False),
    ("4.2  Notices. Notices must be in writing.", False),
    ("4.3  Assignment. Neither Party may assign without consent.", False),
    ("Exhibit A. Renewal notice is due per Section 3.2 and Sections 2.2 and 2.3.", False),
]


@pytest.fixture
def agreement_bytes() -> io.BytesIO:
    d = docx.Document()
    for text, heading in CONTRACT:
        para = d.add_paragraph()
        number, _, body = text.partition("  ") if "  " in text else (text, "", "")
        if body:
            para.add_run(number + "  ").bold = True
            para.add_run(body)
        else:
            para.add_run(text).bold = heading
    buf = io.BytesIO()
    d.save(buf)
    buf.seek(0)
    return buf


@pytest.fixture
def agreement(tmp_path, agreement_bytes):
    path = tmp_path / "agreement.docx"
    path.write_bytes(agreement_bytes.getvalue())
    return path


@pytest.fixture
def rl_agreement(agreement_bytes) -> Redliner:
    return Redliner(agreement_bytes, author="Reviewer", date=DATE)


@pytest.fixture
def contract_bytes() -> io.BytesIO:
    d = docx.Document()
    d.add_paragraph("1. Fees. Customer shall pay the fees in the Order Form.")
    p = d.add_paragraph("2. Invoicing. Payment is due within ")
    p.add_run("thirty (30) days").bold = True
    p.add_run(" of the invoice date.")
    d.add_paragraph("3. Late Payment. Interest accrues at 1.5% per month.")
    d.add_paragraph("4. Termination. Either Party may terminate for material breach.")
    table = d.add_table(rows=3, cols=3)
    for col, value in enumerate(("Tier", "Users", "Fee")):
        table.cell(0, col).text = value
    for col, value in enumerate(("Standard", "50", "$25,000")):
        table.cell(1, col).text = value
    for col, value in enumerate(("Premium", "150", "$60,000")):
        table.cell(2, col).text = value
    buf = io.BytesIO()
    d.save(buf)
    buf.seek(0)
    return buf


@pytest.fixture
def rl(contract_bytes) -> Redliner:
    return Redliner(contract_bytes, author="Tester", date=DATE)


@pytest.fixture
def original_text(contract_bytes) -> list[str]:
    contract_bytes.seek(0)
    out = texts(docx.Document(contract_bytes))
    contract_bytes.seek(0)
    return out
