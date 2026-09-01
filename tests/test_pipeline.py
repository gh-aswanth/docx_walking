"""Clause structure, the renumbering cascade, and the end-to-end pipeline.

The invariant under test throughout: after a structural action, the *accepted*
document must number 1..n with no gaps and every cross-reference must resolve,
while the *rejected* document must be byte-identical in text and numbering to
what we started with.
"""

import io
import json

import docx
import pytest
from conftest import DATE, roundtrip, texts

from docx_redline import Redliner
from docx_redline.errors import ClauseError, RedlineError
from docx_redline.planning.actions import apply_actions, validate_actions
from docx_redline.planning.agent import RuleBasedReviewer, load_proposal, normalize_items
from docx_redline.planning.pipeline import RedlinePipeline, _numbering_gaps
from docx_redline.structure.clauses import ClauseTree, iter_references, parse_clause


# ---------------------------------------------------------------------------
# fixtures
# ---------------------------------------------------------------------------
def labels(document) -> list[str]:
    return [clause.label for clause in ClauseTree(document.element.body)]


def titled(document) -> list[tuple[str, str]]:
    return [(c.label, c.title) for c in ClauseTree(document.element.body)]


def resolve(rl: Redliner, how: str) -> Redliner:
    other = Redliner(roundtrip(rl), track_changes=False)
    getattr(other, f"{how}_all")()
    return other


# ---------------------------------------------------------------------------
# clause parsing
# ---------------------------------------------------------------------------
def test_tree_shape(rl_agreement):
    tree = ClauseTree(rl_agreement.document.element.body)
    assert [s.label for s in tree.sections] == ["1", "2", "3", "4"]
    assert [c.label for c in tree.get("2").children] == ["2.1", "2.2", "2.3"]
    assert tree.get("4.1").title == "Governing Law"
    assert tree.get("2.2").body.startswith("Invoicing.")


@pytest.mark.parametrize(
    "text,expected",
    [
        ("12.1  Governing Law. ...", "12.1"),
        ("12. General Provisions", "12"),
        ("1.2.3  Deeply nested clause.", "1.2.3"),
        ("99.0%-99.89% uptime: 5% service credit", None),
        ("Monthly Uptime Commitment: 99.9%", None),
        ("Subscription Term: 12 months, commencing on the Effective Date", None),
        ("", None),
    ],
)
def test_clause_number_detection(text, expected):
    d = docx.Document()
    p = d.add_paragraph(text)
    clause = parse_clause(p._p)
    assert (clause.label if clause else None) == expected


@pytest.mark.parametrize(
    "text,expected",
    [
        ("as set out in Section 4.2", ["4.2"]),
        ("subject to Section 10.", ["10"]),
        ("breach of Sections 10.2 and 10.3", ["10.2", "10.3"]),
        ("see Sections 2.2, 2.3 and 3.1", ["2.2", "2.3", "3.1"]),
        ("SECTION 9, THE SERVICES ARE PROVIDED", ["9"]),
        ("within thirty (30) days", []),
    ],
)
def test_reference_detection(text, expected):
    assert [num for _, _, num, _ in iter_references(text)] == expected


# ---------------------------------------------------------------------------
# the renumbering cascade
# ---------------------------------------------------------------------------
def test_move_clause_renumbers_both_groups(rl_agreement):
    """The scenario that motivates the whole module.

    4.1 leaves section 4 and lands after 3.1, so it becomes 3.2; the old 3.2
    shifts to 3.3, and 4.2/4.3 close ranks behind it.
    """
    report = apply_actions(
        rl_agreement, [{"id": "M", "type": "move_clause", "clause": "4.1", "after_clause": "3.1"}]
    )
    mapping = {c["from"]: c["to"] for c in report.renumbered}
    assert mapping == {"4.1": "3.2", "3.2": "3.3", "4.2": "4.1", "4.3": "4.2"}

    accepted = resolve(rl_agreement, "accept")
    assert titled(accepted.document) == [
        ("1", "Definitions"),
        ("1.1", '"Services" means the hosted platform'),
        ("1.2", '"Order Form" means the ordering document'),
        ("2", "Fees and Payment"),
        ("2.1", "Fees"),
        ("2.2", "Invoicing"),
        ("2.3", "Taxes"),
        ("3", "Term and Termination"),
        ("3.1", "Term"),
        ("3.2", "Governing Law"),
        ("3.3", "Renewal"),
        ("4", "General"),
        ("4.1", "Notices"),
        ("4.2", "Assignment"),
    ]
    assert not _numbering_gaps(ClauseTree(accepted.document.element.body))


def test_move_clause_rejects_cleanly(rl_agreement, agreement_bytes):
    before = texts(docx.Document(agreement_bytes))
    apply_actions(
        rl_agreement, [{"id": "M", "type": "move_clause", "clause": "4.1", "after_clause": "3.1"}]
    )
    rejected = resolve(rl_agreement, "reject")
    assert texts(rejected.document) == before
    assert labels(rejected.document) == [
        "1",
        "1.1",
        "1.2",
        "2",
        "2.1",
        "2.2",
        "2.3",
        "3",
        "3.1",
        "3.2",
        "4",
        "4.1",
        "4.2",
        "4.3",
    ]


def test_move_updates_cross_references(rl_agreement):
    """A reference to a clause that moved must follow it, in exhibits too."""
    report = apply_actions(
        rl_agreement, [{"id": "M", "type": "move_clause", "clause": "4.1", "after_clause": "3.1"}]
    )
    assert {(r["from"], r["to"]) for r in report.references} == {("3.2", "3.3")}
    accepted = "\n".join(texts(resolve(rl_agreement, "accept").document))
    assert "per Section 3.3" in accepted  # Renewal moved from 3.2 to 3.3
    assert "Sections 2.2 and 2.3" in accepted  # untouched clauses keep their numbers


def test_insert_clause_shifts_following_siblings(rl_agreement):
    report = apply_actions(
        rl_agreement,
        [
            {
                "id": "I",
                "type": "insert_clause",
                "after_clause": "2.1",
                "title": "Invoicing Currency",
                "text": "All amounts are stated in U.S. dollars.",
            }
        ],
    )
    assert {c["from"]: c["to"] for c in report.renumbered} == {"2.2": "2.3", "2.3": "2.4"}
    accepted = resolve(rl_agreement, "accept")
    tree = ClauseTree(accepted.document.element.body)
    assert tree.get("2.2").title == "Invoicing Currency"
    assert tree.get("2.3").title == "Invoicing"
    assert not _numbering_gaps(tree)


def test_inserted_clause_number_is_not_double_tracked(rl_agreement):
    """The new clause's own number is inside the insertion -- no strikeout on it."""
    apply_actions(
        rl_agreement,
        [{"id": "I", "type": "insert_clause", "after_clause": "2.1", "text": "New obligation."}],
    )
    rejected = resolve(rl_agreement, "reject")
    assert labels(rejected.document) == [
        "1",
        "1.1",
        "1.2",
        "2",
        "2.1",
        "2.2",
        "2.3",
        "3",
        "3.1",
        "3.2",
        "4",
        "4.1",
        "4.2",
        "4.3",
    ]
    assert "New obligation." not in "\n".join(texts(rejected.document))


def test_delete_clause_closes_the_gap(rl_agreement):
    report = apply_actions(rl_agreement, [{"id": "D", "type": "delete_clause", "clause": "2.2"}])
    assert {c["from"]: c["to"] for c in report.renumbered} == {"2.3": "2.2"}
    accepted = resolve(rl_agreement, "accept")
    tree = ClauseTree(accepted.document.element.body)
    assert tree.get("2.2").title == "Taxes"
    assert not _numbering_gaps(tree)


def test_delete_clause_flags_dangling_references(rl_agreement):
    report = apply_actions(rl_agreement, [{"id": "D", "type": "delete_clause", "clause": "2.2"}])
    # cited from clause 4.1 and again from the exhibit -- both are flagged
    assert [d["reference"] for d in report.dangling_references] == ["2.2", "2.2"]
    assert any("needs a human decision" in w for w in report.warnings)


def test_delete_section_takes_its_children(rl_agreement):
    apply_actions(rl_agreement, [{"id": "D", "type": "delete_section", "section": "2"}])
    accepted = resolve(rl_agreement, "accept")
    assert [s.label for s in ClauseTree(accepted.document.element.body).sections] == ["1", "2", "3"]
    assert "Fees and Payment" not in "\n".join(texts(accepted.document))


def test_move_section_carries_sub_clauses(rl_agreement):
    apply_actions(
        rl_agreement, [{"id": "M", "type": "move_section", "section": "4", "after_section": "1"}]
    )
    accepted = resolve(rl_agreement, "accept")
    tree = ClauseTree(accepted.document.element.body)
    assert [s.title for s in tree.sections][:3] == ["Definitions", "General", "Fees and Payment"]
    assert [c.label for c in tree.get("2").children] == ["2.1", "2.2", "2.3"]
    assert not _numbering_gaps(tree)


def test_insert_section_appends_and_numbers(rl_agreement):
    apply_actions(
        rl_agreement,
        [
            {
                "id": "S",
                "type": "insert_section",
                "after_section": "4",
                "title": "Insurance",
                "text": "Provider shall maintain cyber cover.",
            }
        ],
    )
    accepted = resolve(rl_agreement, "accept")
    tree = ClauseTree(accepted.document.element.body)
    assert tree.get("5").title == "Insurance"
    assert tree.get("5.1").body.startswith("Provider shall maintain")


def test_renumbering_can_be_switched_off(rl_agreement):
    report = apply_actions(
        rl_agreement,
        [{"id": "M", "type": "move_clause", "clause": "4.1", "after_clause": "3.1"}],
        renumber=False,
    )
    assert report.renumbered == [] and report.references == []


def test_two_structural_actions_compose(rl_agreement):
    """A move and an insert in the same run must not double-count each other."""
    apply_actions(
        rl_agreement,
        [
            {"id": "M", "type": "move_clause", "clause": "4.1", "after_clause": "3.1"},
            {"id": "I", "type": "insert_clause", "after_clause": "2.1", "text": "Currency is USD."},
        ],
    )
    accepted = resolve(rl_agreement, "accept")
    tree = ClauseTree(accepted.document.element.body)
    assert not _numbering_gaps(tree)
    assert tree.get("3.2").title == "Governing Law"
    assert tree.get("2.2").body == "Currency is USD."


# ---------------------------------------------------------------------------
# content actions
# ---------------------------------------------------------------------------
def test_content_actions(rl_agreement, agreement_bytes):
    before = texts(docx.Document(agreement_bytes))
    apply_actions(
        rl_agreement,
        [
            {
                "id": "A",
                "type": "replace_text",
                "clause": "2.2",
                "find": "thirty (30) days",
                "replace": "forty-five (45) days",
            },
            {
                "id": "B",
                "type": "insert_text",
                "clause": "2.2",
                "anchor": "days",
                "position": "after",
                "text": " of receipt",
            },
            {"id": "C", "type": "delete_text", "clause": "3.2", "find": " automatically"},
            {
                "id": "D",
                "type": "rewrite_clause",
                "clause": "4.2",
                "text": "Notices. Notices must be in writing and sent by courier.",
            },
        ],
    )
    accepted = "\n".join(texts(resolve(rl_agreement, "accept").document))
    assert "forty-five (45) days of receipt" in accepted
    assert "renews automatically" not in accepted
    assert "sent by courier" in accepted
    assert texts(resolve(rl_agreement, "reject").document) == before


def test_rewrite_clause_keeps_the_number(rl_agreement):
    apply_actions(
        rl_agreement,
        [
            {
                "id": "R",
                "type": "rewrite_clause",
                "clause": "2.2",
                "text": "Invoicing. Payment is due on receipt.",
            }
        ],
    )
    accepted = resolve(rl_agreement, "accept")
    clause = ClauseTree(accepted.document.element.body).get("2.2")
    assert clause.text.startswith("2.2")
    assert clause.body == "Invoicing. Payment is due on receipt."


def test_action_targeting_a_clause_stays_in_scope(rl_agreement):
    """`clause` scopes the search, so a phrase in two clauses is unambiguous."""
    apply_actions(
        rl_agreement,
        [
            {
                "id": "A",
                "type": "replace_text",
                "clause": "2.1",
                "find": "Order Form",
                "replace": "Order",
            }
        ],
    )
    accepted = texts(resolve(rl_agreement, "accept").document)
    assert "means the ordering document" in "\n".join(accepted)  # 1.2 untouched
    assert "fees in the Order." in "\n".join(accepted)


def test_formatting_and_comment_actions(rl_agreement):
    report = apply_actions(
        rl_agreement,
        [
            {
                "id": "F",
                "type": "format_text",
                "clause": "4.1",
                "find": "Delaware",
                "bold": True,
                "color": "C00000",
            },
            {"id": "P", "type": "format_clause", "clause": "2.1", "space_after": 18},
            {"id": "C", "type": "comment", "clause": "2.2", "text": "Check with finance."},
        ],
    )
    assert report.failed == 0
    counts = rl_agreement.summary().counts
    assert counts["format:rPrChange"] and counts["format:pPrChange"]


def test_table_actions(rl_agreement, agreement_bytes):
    d = docx.Document(agreement_bytes)
    table = d.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Tier"
    table.cell(0, 1).text = "Fee"
    table.cell(1, 0).text = "Standard"
    table.cell(1, 1).text = "$10,000"
    buf = io.BytesIO()
    d.save(buf)
    buf.seek(0)
    rl = Redliner(buf, author="Reviewer", date=DATE)
    report = apply_actions(
        rl,
        [
            {"id": "T1", "type": "insert_row", "table": 0, "values": ["Premium", "$25,000"]},
            {"id": "T2", "type": "update_cell", "table": 0, "row": 1, "col": 1, "text": "$12,000"},
        ],
    )
    assert report.failed == 0
    accepted = "\n".join(texts(resolve(rl, "accept").document))
    assert "Premium" in accepted and "$12,000" in accepted


# ---------------------------------------------------------------------------
# validation
# ---------------------------------------------------------------------------
def test_validation_catches_bad_action_items():
    problems = validate_actions(
        [
            {"id": "1", "type": "nonsense"},
            {"id": "2", "type": "replace_text", "find": "x"},
            {"id": "3", "type": "move_clause", "clause": "4.1"},
            {
                "id": "7",
                "type": "move_clause",
                "clause": "4.1",
                "into_section": "3",
                "position": "middle",
            },
            {"id": "8", "type": "reorder_clauses", "section": "2", "order": ["2.1"]},
            {"id": "9", "type": "reorder_clauses", "section": "2", "order": ["2.1", "2.1", "2.2"]},
            {"id": "10", "type": "move_section", "section": "2"},
            {"id": "4", "type": "renumber_clause", "clause": "4.1"},
            {"id": "5", "type": "replace_text", "find": "a", "replace": "b", "bogus": 1},
            {"id": "5", "type": "delete_text", "find": "a"},
            {"id": "6", "type": "delete_text", "find": "a", "severity": "urgent"},
        ]
    )
    joined = "\n".join(problems)
    assert "unknown action type" in joined
    assert "requires 'replace'" in joined
    assert "needs after_clause, before_clause or into_section" in joined
    assert "position must be 'first' or 'last'" in joined
    assert "'order' list of two or more" in joined
    assert "'order' repeats a clause number" in joined
    assert "move_section needs after_section, before_section or position" in joined
    assert "derived by the planner" in joined
    assert "does not accept 'bogus'" in joined
    assert "duplicate action id" in joined
    assert "severity 'urgent'" in joined


def test_invalid_plan_is_rejected_before_touching_the_document(rl_agreement):
    with pytest.raises(RedlineError, match="invalid action items"):
        apply_actions(rl_agreement, [{"id": "X", "type": "replace_text", "find": "a"}])
    assert len(rl_agreement.summary()) == 0


def test_unresolvable_action_is_reported_not_raised(rl_agreement):
    report = apply_actions(
        rl_agreement,
        [{"id": "X", "type": "replace_text", "clause": "2.2", "find": "nope", "replace": "y"}],
    )
    assert report.failed == 1
    assert report.results[0].status == "failed"
    assert report.warnings


def test_strict_mode_raises(rl_agreement):
    with pytest.raises(ClauseError):
        apply_actions(
            rl_agreement,
            [{"id": "X", "type": "delete_clause", "clause": "99.9"}],
            strict=True,
        )


def test_normalize_strips_nulls_and_assigns_ids():
    items = normalize_items([{"type": "delete_text", "find": "x", "clause": None, "regex": ""}])
    assert items == [{"type": "delete_text", "find": "x", "id": "AI-0001"}]


# ---------------------------------------------------------------------------
# reviewers + pipeline
# ---------------------------------------------------------------------------
def test_rule_based_reviewer_emits_valid_items(rl_agreement):
    proposal = RuleBasedReviewer().propose(ClauseTree(rl_agreement.document.element.body), "brief")
    assert proposal.action_items
    assert validate_actions(proposal.action_items) == []
    assert proposal.source == "rule-based"


def test_load_proposal_round_trip(tmp_path):
    path = tmp_path / "actions.json"
    path.write_text(
        json.dumps(
            {
                "summary": "s",
                "action_items": [{"id": "A", "type": "delete_text", "find": "x", "clause": None}],
            }
        )
    )
    proposal = load_proposal(path)
    assert proposal.action_items == [{"id": "A", "type": "delete_text", "find": "x"}]


def test_load_proposal_rejects_bad_file(tmp_path):
    path = tmp_path / "bad.json"
    path.write_text(json.dumps([{"id": "A", "type": "replace_text", "find": "x"}]))
    with pytest.raises(ValueError, match="invalid action items"):
        load_proposal(path)


def test_pipeline_end_to_end(tmp_path, agreement):
    actions = tmp_path / "action_items.json"
    output = tmp_path / "out.docx"
    report = tmp_path / "report.json"

    pipeline = RedlinePipeline(agreement, author="Bot", date=DATE)
    result = pipeline.run(output, actions_file=actions, report_path=report)

    assert result.ok, result.format()
    assert [s.name for s in result.stages] == [
        "extract",
        "propose",
        "validate",
        "plan",
        "renumber",
        "apply",
        "verify",
        "report",
    ]
    assert output.exists() and report.exists() and actions.exists()
    assert all(passed for _, passed, _ in result.checks)

    payload = json.loads(report.read_text())
    assert payload["ok"] is True
    assert payload.get("summary", True)
    assert payload["plan"]["summary"]["actions"] >= 1


def test_pipeline_replays_a_saved_action_file(tmp_path, agreement):
    actions = tmp_path / "action_items.json"
    actions.write_text(
        json.dumps(
            {
                "action_items": [
                    {"id": "A", "type": "move_clause", "clause": "4.1", "after_clause": "3.1"},
                ]
            }
        )
    )
    pipeline = RedlinePipeline(agreement, date=DATE)
    result = pipeline.run(tmp_path / "out.docx", actions_file=actions, write_actions=False)
    assert result.ok
    assert result.proposal.source == "file"
    assert {c["from"]: c["to"] for c in result.plan.renumbered}["4.1"] == "3.2"


def test_numbering_gap_detector_actually_detects_gaps(rl_agreement):
    """The verification check must fail on bad numbering, not rubber-stamp it."""
    tree = ClauseTree(rl_agreement.document.element.body)
    assert _numbering_gaps(tree) == []
    tree.get("2.2").label = "2.7"
    assert "expected 2.2, found 2.7" in _numbering_gaps(tree)


def test_pipeline_reports_not_ok_when_an_action_fails(tmp_path, agreement):
    actions = tmp_path / "actions.json"
    actions.write_text(
        json.dumps(
            {
                "action_items": [
                    {
                        "id": "A",
                        "type": "replace_text",
                        "clause": "2.2",
                        "find": "absent",
                        "replace": "x",
                    },
                ]
            }
        )
    )
    result = RedlinePipeline(agreement, date=DATE).run(
        tmp_path / "out.docx", actions_file=actions, write_actions=False
    )
    assert not result.ok
    assert result.plan.failed == 1


def test_structural_actions_do_not_reuse_stale_paragraph_identities(rl_agreement):
    """Regression: the label snapshot must survive several tree reloads.

    Each structural action reloads the clause tree, freeing the previous batch
    of lxml proxies. If the snapshot only stored their ``id()`` values, CPython
    would recycle those addresses and later clauses would be diffed against the
    wrong baseline label -- producing nonsense like ``11.1 -> 11``.
    """
    report = apply_actions(
        rl_agreement,
        [
            {"id": "M", "type": "move_clause", "clause": "4.1", "after_clause": "3.1"},
            {"id": "D", "type": "delete_clause", "clause": "1.2"},
            {"id": "I", "type": "insert_clause", "after_clause": "2.1", "text": "Currency is USD."},
        ],
    )
    assert {c["from"]: c["to"] for c in report.renumbered} == {
        "4.1": "3.2",  # moved
        "3.2": "3.3",  # destination sibling shifts down
        "2.2": "2.3",  # insertion pushes the rest of section 2 down
        "2.3": "2.4",
        "4.2": "4.1",  # source siblings close ranks
        "4.3": "4.2",
    }
    accepted = resolve(rl_agreement, "accept")
    assert not _numbering_gaps(ClauseTree(accepted.document.element.body))


def test_repeated_runs_are_deterministic(agreement_bytes):
    """Same input, same action items -> byte-identical renumbering, every time."""
    plans = []
    for _ in range(3):
        agreement_bytes.seek(0)
        rl = Redliner(agreement_bytes, author="Reviewer", date=DATE)
        report = apply_actions(
            rl,
            [
                {"id": "M", "type": "move_clause", "clause": "4.1", "after_clause": "3.1"},
                {"id": "I", "type": "insert_clause", "after_clause": "2.1", "text": "USD only."},
            ],
        )
        plans.append([(c["from"], c["to"]) for c in report.renumbered])
    assert plans[0] == plans[1] == plans[2]


# ---------------------------------------------------------------------------
# move to top / reorder by number
# ---------------------------------------------------------------------------
def test_move_clause_to_top_of_section(rl_agreement):
    """`into_section` + `position: first` -- the "move it to the top" ask."""
    report = apply_actions(
        rl_agreement,
        [
            {
                "id": "T",
                "type": "move_clause",
                "clause": "2.3",
                "into_section": "2",
                "position": "first",
            }
        ],
    )
    assert report.results[0].detail == "moved clause 2.3 to the top of section 2"
    assert {c["from"]: c["to"] for c in report.renumbered} == {
        "2.3": "2.1",
        "2.1": "2.2",
        "2.2": "2.3",
    }
    accepted = resolve(rl_agreement, "accept")
    tree = ClauseTree(accepted.document.element.body)
    assert [(c.label, c.title) for c in tree.get("2").children] == [
        ("2.1", "Taxes"),
        ("2.2", "Fees"),
        ("2.3", "Invoicing"),
    ]
    assert not _numbering_gaps(tree)


def test_move_clause_to_top_updates_references_to_it(rl_agreement):
    """4.1 cites 2.2; moving 2.2 to the top must repoint that citation."""
    report = apply_actions(
        rl_agreement,
        [
            {
                "id": "T",
                "type": "move_clause",
                "clause": "2.2",
                "into_section": "2",
                "position": "first",
            }
        ],
    )
    # 2.2 is cited twice (from clause 4.1 and from the exhibit); nothing cites 2.1
    assert {(r["from"], r["to"]) for r in report.references} == {("2.2", "2.1")}
    assert len(report.references) == 2
    accepted = "\n".join(texts(resolve(rl_agreement, "accept").document))
    assert "as noted in Section 2.1" in accepted


def test_move_clause_to_bottom_of_section(rl_agreement):
    report = apply_actions(
        rl_agreement,
        [
            {
                "id": "T",
                "type": "move_clause",
                "clause": "2.1",
                "into_section": "2",
                "position": "last",
            }
        ],
    )
    assert {c["from"]: c["to"] for c in report.renumbered} == {
        "2.2": "2.1",
        "2.3": "2.2",
        "2.1": "2.3",
    }
    tree = ClauseTree(resolve(rl_agreement, "accept").document.element.body)
    assert tree.get("2.3").title == "Fees"


def test_move_clause_into_another_section_top(rl_agreement):
    apply_actions(
        rl_agreement,
        [
            {
                "id": "T",
                "type": "move_clause",
                "clause": "4.3",
                "into_section": "3",
                "position": "first",
            }
        ],
    )
    tree = ClauseTree(resolve(rl_agreement, "accept").document.element.body)
    assert [c.title for c in tree.get("3").children] == ["Assignment", "Term", "Renewal"]
    assert [c.title for c in tree.get("4").children] == ["Governing Law", "Notices"]
    assert not _numbering_gaps(tree)


def test_move_section_to_top_of_document(rl_agreement):
    report = apply_actions(
        rl_agreement, [{"id": "T", "type": "move_section", "section": "3", "position": "first"}]
    )
    assert report.results[0].detail == "moved section 3 to the top of the document"
    tree = ClauseTree(resolve(rl_agreement, "accept").document.element.body)
    assert [s.title for s in tree.sections] == [
        "Term and Termination",
        "Definitions",
        "Fees and Payment",
        "General",
    ]
    assert [c.label for c in tree.get("1").children] == ["1.1", "1.2"]
    assert not _numbering_gaps(tree)


def test_move_section_before_another(rl_agreement):
    apply_actions(
        rl_agreement, [{"id": "T", "type": "move_section", "section": "4", "before_section": "2"}]
    )
    tree = ClauseTree(resolve(rl_agreement, "accept").document.element.body)
    assert [s.title for s in tree.sections] == [
        "Definitions",
        "General",
        "Fees and Payment",
        "Term and Termination",
    ]


def test_reorder_clauses_by_number(rl_agreement):
    report = apply_actions(
        rl_agreement,
        [{"id": "R", "type": "reorder_clauses", "section": "2", "order": ["2.3", "2.1", "2.2"]}],
    )
    assert {c["from"]: c["to"] for c in report.renumbered} == {
        "2.3": "2.1",
        "2.1": "2.2",
        "2.2": "2.3",
    }
    tree = ClauseTree(resolve(rl_agreement, "accept").document.element.body)
    assert [c.title for c in tree.get("2").children] == ["Taxes", "Fees", "Invoicing"]
    assert not _numbering_gaps(tree)


def test_reorder_moves_only_what_has_to_move(rl_agreement):
    """[A,B,C] -> [C,A,B] is one move, not three: A and B keep their order."""
    report = apply_actions(
        rl_agreement,
        [{"id": "R", "type": "reorder_clauses", "section": "2", "order": ["2.3", "2.1", "2.2"]}],
    )
    assert report.results[0].edits == 1
    assert "1 clause(s) moved, 2 left in place" in report.results[0].detail
    assert rl_agreement.summary().counts["move-from"] == 2  # one run + its paragraph mark


def test_reorder_reversal_moves_the_minimum(rl_agreement):
    report = apply_actions(
        rl_agreement,
        [{"id": "R", "type": "reorder_clauses", "section": "2", "order": ["2.3", "2.2", "2.1"]}],
    )
    assert report.results[0].edits == 2  # a full reversal keeps exactly one clause put
    tree = ClauseTree(resolve(rl_agreement, "accept").document.element.body)
    assert [c.title for c in tree.get("2").children] == ["Taxes", "Invoicing", "Fees"]


def test_reorder_rejects_cleanly(rl_agreement, agreement_bytes):
    before = texts(docx.Document(agreement_bytes))
    apply_actions(
        rl_agreement,
        [{"id": "R", "type": "reorder_clauses", "section": "2", "order": ["2.2", "2.3", "2.1"]}],
    )
    rejected = resolve(rl_agreement, "reject")
    assert texts(rejected.document) == before
    assert [c.label for c in ClauseTree(rejected.document.element.body).get("2").children] == [
        "2.1",
        "2.2",
        "2.3",
    ]


def test_reorder_that_is_already_correct_is_a_no_op(rl_agreement):
    report = apply_actions(
        rl_agreement,
        [{"id": "R", "type": "reorder_clauses", "section": "2", "order": ["2.1", "2.2", "2.3"]}],
    )
    assert report.results[0].status == "skipped"
    assert len(rl_agreement.summary()) == 0


def test_reorder_skips_a_clause_that_moved_out_of_the_section(rl_agreement):
    """An earlier action may have taken a clause away; that is not a failure."""
    report = apply_actions(
        rl_agreement,
        [
            {
                "id": "M",
                "type": "move_clause",
                "clause": "2.1",
                "into_section": "3",
                "position": "first",
            },
            {"id": "R", "type": "reorder_clauses", "section": "2", "order": ["2.3", "2.1", "2.2"]},
        ],
    )
    assert report.failed == 0
    assert any("2.1 has since moved out of section 2" in w for w in report.warnings)
    tree = ClauseTree(resolve(rl_agreement, "accept").document.element.body)
    assert [c.title for c in tree.get("2").children] == ["Taxes", "Invoicing"]
    assert not _numbering_gaps(tree)


def test_reorder_rejects_a_clause_number_that_exists_nowhere(rl_agreement):
    report = apply_actions(
        rl_agreement,
        [{"id": "R", "type": "reorder_clauses", "section": "2", "order": ["2.1", "2.2", "2.9"]}],
    )
    assert report.failed == 1
    assert "has no clause 2.9" in report.results[0].detail


def test_reorder_leaves_unlisted_clauses_where_they_are(rl_agreement):
    """Only the listed clauses are permuted; an unmentioned one must not drift."""
    report = apply_actions(
        rl_agreement,
        [{"id": "R", "type": "reorder_clauses", "section": "2", "order": ["2.3", "2.1"]}],
    )
    assert any("left where they are: 2.2" in w for w in report.warnings)
    tree = ClauseTree(resolve(rl_agreement, "accept").document.element.body)
    # 2.2 (Invoicing) held slot 2 and keeps it; Taxes and Fees swap around it
    assert [c.title for c in tree.get("2").children] == ["Taxes", "Invoicing", "Fees"]


@pytest.mark.parametrize(
    "order",
    [
        ["2.1", "2.2", "2.3"],
        ["2.1", "2.3", "2.2"],
        ["2.2", "2.1", "2.3"],
        ["2.2", "2.3", "2.1"],
        ["2.3", "2.1", "2.2"],
        ["2.3", "2.2", "2.1"],
    ],
)
def test_every_permutation_reorders_correctly(agreement_bytes, order):
    agreement_bytes.seek(0)
    rl = Redliner(agreement_bytes, author="Reviewer", date=DATE)
    titles = {"2.1": "Fees", "2.2": "Invoicing", "2.3": "Taxes"}
    apply_actions(rl, [{"id": "R", "type": "reorder_clauses", "section": "2", "order": order}])
    tree = ClauseTree(resolve(rl, "accept").document.element.body)
    assert [c.title for c in tree.get("2").children] == [titles[label] for label in order]
    assert not _numbering_gaps(tree)


# ---------------------------------------------------------------------------
# moving a clause that was already edited
# ---------------------------------------------------------------------------
def test_moving_an_already_edited_clause_round_trips(rl_agreement, agreement_bytes):
    """Regression: a move revision cannot carry the source's own strikeouts.

    Deep-copying a paragraph that already holds ``w:ins``/``w:del`` leaves the
    deleted runs *outside* the ``w:moveTo`` wrapper, so rejecting the redline
    resurrected them into the following paragraph. Such moves are now recorded
    as delete + insert instead.
    """
    before = texts(docx.Document(agreement_bytes))
    report = apply_actions(
        rl_agreement,
        [
            {
                "id": "C",
                "type": "replace_text",
                "clause": "2.2",
                "find": "thirty (30) days",
                "replace": "forty-five (45) days",
            },
            {
                "id": "M",
                "type": "move_clause",
                "clause": "2.2",
                "into_section": "2",
                "position": "first",
            },
        ],
    )
    assert any("delete + insert" in w for w in report.warnings)

    accepted = resolve(rl_agreement, "accept")
    tree = ClauseTree(accepted.document.element.body)
    assert tree.get("2.1").title == "Invoicing"
    assert "forty-five (45) days" in tree.get("2.1").text
    assert "thirty (30) days" not in "\n".join(texts(accepted.document))
    assert not _numbering_gaps(tree)

    assert texts(resolve(rl_agreement, "reject").document) == before


def test_clean_move_still_uses_move_revisions(rl_agreement):
    apply_actions(
        rl_agreement,
        [
            {
                "id": "M",
                "type": "move_clause",
                "clause": "2.2",
                "into_section": "2",
                "position": "first",
            }
        ],
    )
    counts = rl_agreement.summary().counts
    assert counts["move-from"] and counts["move-to"]
    assert not counts["paragraph-mark-delete"]
