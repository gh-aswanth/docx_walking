"""Paragraph addressing: stable ids, folded matching, and plan verification.

The invariant under test throughout is that an edit is *verified* before it is
applied. Every rejection here is a case where guessing would have produced a
document that looked redlined and was wrong.
"""

import docx
import pytest
from conftest import DATE, document_xml, texts

from docx_redline import (
    ParagraphIndex,
    RedlineEdit,
    Redliner,
    Rejection,
    ReviewNote,
    StalePlanError,
    fold,
    load_edits,
    validate_edits,
    verify_plan,
)
from docx_redline.errors import RedlineError


def build(*paragraphs) -> Redliner:
    document = docx.Document()
    for text in paragraphs:
        document.add_paragraph(text)
    return Redliner(document, author="Reviewer", date=DATE)


# ---------------------------------------------------------------------------
# fold
# ---------------------------------------------------------------------------
@pytest.mark.parametrize(
    "raw",
    [
        "“twelve (12)–month” term",
        "a non-breaking space",
        "an apostrophe’s curl",
        "Ａ fullwidth",
        "a\ttab",
    ],
)
def test_fold_preserves_length(raw):
    """Offsets computed on the folded string must be valid on the raw one."""
    assert len(fold(raw)) == len(raw)


def test_fold_leaves_ligatures_alone():
    """NFKC would expand this and shift every offset after it."""
    assert fold("ofﬁce") == "ofﬁce"


def test_fold_matches_straight_quotes_against_smart_ones():
    rl = build("Payment is due within “thirty (30) days” of invoice.")
    index = ParagraphIndex(rl)
    report = index.apply([RedlineEdit(0, '"thirty (30) days"', '"forty-five (45) days"')])
    assert report.applied
    assert "forty-five (45) days" in index[0].text


# ---------------------------------------------------------------------------
# the address space
# ---------------------------------------------------------------------------
def test_ids_are_document_order_and_render_carries_them():
    rl = build("First.", "Second.", "Third.")
    index = ParagraphIndex(rl)
    assert [ref.para_id for ref in index] == [0, 1, 2]
    assert index.render().splitlines() == ["[0] First.", "[1] Second.", "[2] Third."]


def test_render_returns_the_paragraph_text_not_just_clause_labels():
    """The prototype built the body into a dict it never read; this is that bug."""
    rl = build("1. Definitions", "1.1  “User” means a person.")
    body = ParagraphIndex(rl).render()
    assert "“User” means a person." in body


def test_index_sees_text_inside_an_existing_insertion():
    rl = build("Payment is due within thirty (30) days.")
    rl.replace_text("thirty (30) days", "forty-five (45) days")
    index = ParagraphIndex(rl)
    assert "forty-five (45) days" in index[0].text
    assert index.apply([RedlineEdit(0, "forty-five (45) days", "sixty (60) days")]).applied


# ---------------------------------------------------------------------------
# rejections
# ---------------------------------------------------------------------------
def reason(report):
    return report.rejected[0].reason


def test_unknown_para_id_is_rejected_with_the_valid_range():
    rl = build("Only paragraph.")
    report = ParagraphIndex(rl).apply([RedlineEdit(7, "Only", "One")])
    assert reason(report) is Rejection.PARAGRAPH_NOT_FOUND
    assert "0-0" in report.rejected[0].detail


def test_a_quote_that_does_not_match_is_never_fuzzy_matched():
    rl = build("Payment is due within thirty (30) days.")
    report = ParagraphIndex(rl).apply([RedlineEdit(0, "thirty days", "sixty days")])
    assert reason(report) is Rejection.TARGET_NOT_FOUND
    assert texts(rl.document) == ["Payment is due within thirty (30) days."]


def test_a_short_repeated_span_is_ambiguous_rather_than_the_first_hit():
    rl = build("The Provider shall notify the Provider's agent.")
    report = ParagraphIndex(rl).apply([RedlineEdit(0, "Provider", "Supplier")])
    assert reason(report) is Rejection.TARGET_AMBIGUOUS


def test_occurrence_resolves_a_repeated_span():
    rl = build("The Provider shall notify the Provider's agent.")
    index = ParagraphIndex(rl)
    assert index.apply([RedlineEdit(0, "Provider", "Supplier", occurrence=2)]).applied
    assert index[0].text == "The Provider shall notify the Supplier's agent."


def test_occurrence_zero_replaces_every_hit():
    rl = build("The Provider shall notify the Provider's agent.")
    index = ParagraphIndex(rl)
    report = index.apply([RedlineEdit(0, "Provider", "Supplier", occurrence=0)])
    assert report.applied[0].spans == 2
    assert index[0].text == "The Supplier shall notify the Supplier's agent."


def test_two_edits_to_one_span_collide_instead_of_stacking():
    rl = build("Payment is due within thirty (30) days.")
    index = ParagraphIndex(rl)
    report = index.apply(
        [
            RedlineEdit(0, "thirty (30) days", "forty-five (45) days", agent="a"),
            RedlineEdit(0, "thirty (30) days", "sixty (60) days", agent="b"),
        ]
    )
    assert len(report.applied) == 1
    assert reason(report) is Rejection.SPAN_CONFLICT
    assert index[0].text == "Payment is due within forty-five (45) days."


def test_disjoint_edits_in_one_paragraph_both_land():
    rl = build("Payment is due within thirty (30) days of the invoice date.")
    index = ParagraphIndex(rl)
    report = index.apply(
        [
            RedlineEdit(0, "thirty (30) days", "forty-five (45) days"),
            RedlineEdit(0, "of the invoice date", "of receipt of a valid invoice"),
        ]
    )
    assert len(report.applied) == 2
    assert index[0].text == (
        "Payment is due within forty-five (45) days of receipt of a valid invoice."
    )


def test_text_that_is_already_struck_is_reported_as_such():
    rl = build("Payment is due within thirty (30) days.")
    rl.delete_matching("thirty (30) days")
    report = ParagraphIndex(rl).apply([RedlineEdit(0, "thirty (30) days", "sixty (60) days")])
    assert reason(report) is Rejection.TARGET_ALREADY_STRUCK


def test_nothing_is_written_when_an_edit_is_rejected():
    rl = build("Alpha.", "Beta.")
    index = ParagraphIndex(rl)
    index.apply([RedlineEdit(0, "Gamma", "Delta")])
    assert rl.summary().revisions == []


# ---------------------------------------------------------------------------
# review notes
# ---------------------------------------------------------------------------
def test_a_note_anchors_to_its_span_and_signs_with_its_agent():
    rl = build("Provider shall invoice Customer annually in advance.")
    index = ParagraphIndex(rl)
    report = index.apply(
        [ReviewNote(0, "annually in advance", "Confirm AP can honour this.", agent="Payment Terms")]
    )
    assert report.applied
    comment = next(iter(rl.document.comments))
    assert (comment.author, comment.initials) == ("Payment Terms", "PT")
    assert comment.text == "Confirm AP can honour this."
    # the range brackets the quoted span, not the whole paragraph
    body = document_xml(rl)
    assert body.index("commentRangeStart") < body.index("annually in advance")
    assert body.index("annually in advance") < body.index("commentRangeEnd")


def test_a_note_over_text_an_edit_struck_falls_back_to_the_paragraph():
    rl = build("Payment is due within thirty (30) days.")
    index = ParagraphIndex(rl)
    report = index.apply(
        [
            RedlineEdit(0, "thirty (30) days", "forty-five (45) days"),
            ReviewNote(0, "thirty (30) days", "Superseded by the 45-day edit."),
        ]
    )
    assert len(report.applied) == 2
    assert "target_already_struck" in report.applied[1].detail


def test_a_note_over_text_struck_in_an_earlier_pass_still_lands():
    """The comment is usually *about* the strike; dropping it loses the finding."""
    rl = build("Payment is due within thirty (30) days.")
    index = ParagraphIndex(rl)
    index.apply([RedlineEdit(0, "thirty (30) days", "forty-five (45) days")])
    report = index.apply([ReviewNote(0, "thirty (30) days", "Superseded; kept for audit.")])
    assert report.applied and not report.rejected
    assert "target_already_struck" in report.applied[0].detail
    assert next(iter(rl.document.comments)).text == "Superseded; kept for audit."


def test_an_edit_over_text_struck_in_an_earlier_pass_is_still_refused():
    rl = build("Payment is due within thirty (30) days.")
    index = ParagraphIndex(rl)
    index.apply([RedlineEdit(0, "thirty (30) days", "forty-five (45) days")])
    report = index.apply([RedlineEdit(0, "thirty (30) days", "ten (10) days")])
    assert reason(report) is Rejection.TARGET_ALREADY_STRUCK


def test_a_note_never_blocks_an_edit_on_the_same_span():
    rl = build("Payment is due within thirty (30) days.")
    index = ParagraphIndex(rl)
    report = index.apply(
        [
            ReviewNote(0, "thirty (30) days", "Too short."),
            RedlineEdit(0, "thirty (30) days", "forty-five (45) days"),
        ]
    )
    assert len(report.applied) == 2


# ---------------------------------------------------------------------------
# plan verification
# ---------------------------------------------------------------------------
def test_verify_plan_passes_on_an_untouched_document():
    rl = build("Payment is due within thirty (30) days.")
    index = ParagraphIndex(rl)
    assert verify_plan(index, index.fingerprint()) == index.fingerprint()


def test_verify_plan_refuses_a_plan_computed_against_older_text():
    rl = build("Payment is due within thirty (30) days.")
    index = ParagraphIndex(rl)
    stale = index.fingerprint()
    index.apply([RedlineEdit(0, "thirty (30) days", "forty-five (45) days")])
    with pytest.raises(StalePlanError, match="document moved on"):
        verify_plan(index, stale)


def test_fingerprint_ignores_attribution_and_tracks_only_text():
    one = ParagraphIndex(build("Alpha.", "Beta.")).fingerprint()
    two = ParagraphIndex(Redliner(_doc("Alpha.", "Beta."), author="Someone Else")).fingerprint()
    assert one == two


def _doc(*paragraphs):
    document = docx.Document()
    for text in paragraphs:
        document.add_paragraph(text)
    return document


# ---------------------------------------------------------------------------
# model output
# ---------------------------------------------------------------------------
def test_load_edits_rejects_a_malformed_batch_instead_of_dropping_it():
    with pytest.raises(RedlineError, match="does not accept 'replacment'"):
        load_edits([{"para_id": 0, "target": "x", "replacment": "y"}])


def test_load_edits_non_strict_keeps_the_good_items():
    items = load_edits(
        [
            {"para_id": 0, "target": "x", "replacement": "y"},
            {"kind": "wat", "para_id": 1, "target": "z"},
        ],
        strict=False,
    )
    assert [type(i).__name__ for i in items] == ["RedlineEdit"]


def test_validate_edits_flags_a_bad_severity():
    problems = validate_edits(
        [{"para_id": 0, "target": "x", "replacement": "y", "severity": "urgent"}]
    )
    assert problems and "severity" in problems[0]


def test_load_edits_builds_both_kinds():
    items = load_edits(
        [
            {"kind": "edit", "para_id": 0, "target": "x", "replacement": "y", "agent": "a"},
            {"kind": "note", "para_id": 1, "target": "z", "body": "look at this"},
        ]
    )
    assert [type(i).__name__ for i in items] == ["RedlineEdit", "ReviewNote"]
