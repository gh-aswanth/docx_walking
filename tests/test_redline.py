"""Behavioural tests.

The central invariant every test leans on: a correct redline must resolve both
ways -- ``accept`` gives the new document, ``reject`` gives the old one back.
"""

import docx
import pytest
from conftest import DATE, document_xml, resolved, roundtrip, texts
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from docx_redline import Redliner, redline_files
from docx_redline.editing.ops import apply_operations
from docx_redline.editing.redline import RedlineError
from docx_redline.oxml.diffing import diff_ops
from docx_redline.oxml.ns import qn


# ---------------------------------------------------------------------------
# text-level edits
# ---------------------------------------------------------------------------
def test_replace_text_round_trips(rl, original_text):
    assert rl.replace_text("thirty (30) days", "forty-five (45) days") == 1
    assert "forty-five (45) days" in "\n".join(resolved(rl, "accept"))
    assert resolved(rl, "reject") == original_text


def test_replace_spans_multiple_runs(rl, original_text):
    # "within thirty (30) days of" crosses three runs with different formatting
    assert rl.replace_text("within thirty (30) days of", "no later than 45 days after") == 1
    accepted = "\n".join(resolved(rl, "accept"))
    assert "no later than 45 days after the invoice date" in accepted
    assert resolved(rl, "reject") == original_text


def test_replace_all_occurrences(rl):
    assert rl.replace_text("Payment", "Remittance", count=None) == 2
    accepted = "\n".join(resolved(rl, "accept"))
    assert "Remittance is due" in accepted and "Late Remittance" in accepted


def test_replace_with_regex_backreference(rl):
    assert rl.replace_text(r"(\d+\.\d)% per month", r"\1% per annum", regex=True) == 1
    assert "1.5% per annum" in "\n".join(resolved(rl, "accept"))


def test_insert_and_delete_text(rl, original_text):
    rl.insert_text_after("invoice date", ", without setoff")
    rl.delete_matching("Interest accrues at 1.5% per month.")
    accepted = "\n".join(resolved(rl, "accept"))
    assert ", without setoff" in accepted
    assert "1.5% per month" not in accepted
    assert resolved(rl, "reject") == original_text


def test_deleted_text_uses_delText(rl):
    rl.delete_matching("Late Payment")
    assert "delText" in document_xml(rl)


def test_set_paragraph_text_is_minimal(rl):
    para = rl.find_paragraph(contains="Termination")
    rl.set_paragraph_text(
        para, "4. Termination. Either Party may terminate for material breach on notice."
    )
    summary = rl.summary()
    # a whole-paragraph rewrite would produce a delete; a word diff should not
    assert summary.counts["insert"] == 1
    assert summary.counts["delete"] == 0
    assert "breach on notice." in "\n".join(resolved(rl, "accept"))


def test_edit_inside_previously_inserted_text(rl, original_text):
    """A second pass must be able to edit text a first pass inserted."""
    rl.insert_text_after("Order Form", " and any Statement of Work")
    rl.replace_text("Statement of Work", "SOW")
    accepted = "\n".join(resolved(rl, "accept"))
    assert "and any SOW" in accepted
    assert resolved(rl, "reject") == original_text


# ---------------------------------------------------------------------------
# paragraph-level edits
# ---------------------------------------------------------------------------
def test_insert_paragraph_after(rl, original_text):
    rl.insert_paragraph_after(rl.find_paragraph(contains="Late Payment"), "3A. New clause.")
    accepted = resolved(rl, "accept")
    assert "3A. New clause." in accepted
    assert len(accepted) == len(original_text) + 1
    assert resolved(rl, "reject") == original_text


def test_insert_paragraph_before(rl, original_text):
    rl.insert_paragraph_before(rl.find_paragraph(contains="1. Fees"), "0. Preamble.")
    assert resolved(rl, "accept")[0] == "0. Preamble."
    assert resolved(rl, "reject") == original_text


def test_append_paragraph_at_end_of_body(rl, original_text):
    rl.append_paragraph("[Signature page follows.]")
    accepted = resolved(rl, "accept")
    assert accepted[-1] == "[Signature page follows.]" or "[Signature page follows.]" in accepted
    assert resolved(rl, "reject") == original_text


def test_delete_paragraph(rl, original_text):
    rl.delete_paragraph(rl.find_paragraph(contains="Late Payment"))
    accepted = resolved(rl, "accept")
    assert not any("Late Payment" in t for t in accepted)
    assert len(accepted) == len(original_text) - 1
    assert resolved(rl, "reject") == original_text


def test_delete_last_paragraph_of_story(rl, original_text):
    # the last body paragraph is inside the table; use the last non-table one
    target = rl.paragraphs(include_tables=False)[-1]
    rl.delete_paragraph(target)
    assert "material breach" not in "\n".join(resolved(rl, "accept"))
    assert resolved(rl, "reject") == original_text


def test_move_paragraph(rl, original_text):
    source = rl.find_paragraph(contains="Termination")
    destination = rl.find_paragraph(contains="1. Fees")
    rl.move_paragraph(source, destination)
    accepted = resolved(rl, "accept")
    assert accepted.index("4. Termination. Either Party may terminate for material breach.") == 1
    assert len(accepted) == len(original_text)
    assert resolved(rl, "reject") == original_text


def test_move_uses_move_revisions(rl):
    rl.move_paragraph(
        rl.find_paragraph(contains="Termination"), rl.find_paragraph(contains="1. Fees")
    )
    counts = rl.summary().counts
    assert counts["move-from"] and counts["move-to"]


# ---------------------------------------------------------------------------
# tables
# ---------------------------------------------------------------------------
def test_insert_table_row(rl, original_text):
    rl.insert_table_row(rl.tables()[0], values=["Enterprise", "500", "$180,000"])
    accepted = resolved(rl, "accept")
    assert "Enterprise" in accepted and "$180,000" in accepted
    assert resolved(rl, "reject") == original_text


def test_insert_table_row_at_index(rl):
    rl.insert_table_row(rl.tables()[0], index=1, values=["Trial", "5", "$0"])
    accepted = resolved(rl, "accept")
    assert accepted.index("Trial") < accepted.index("Standard")


def test_delete_table_row(rl, original_text):
    rl.delete_table_row(rl.tables()[0], 1)
    assert "Standard" not in resolved(rl, "accept")
    assert resolved(rl, "reject") == original_text


def test_deleted_row_keeps_geometry_on_reject(rl):
    table = rl.tables()[0]
    rl.delete_table_row(table, 2)
    other = Redliner(roundtrip(rl), track_changes=False)
    other.reject_all()
    assert len(other.tables()[0].rows) == 3


def test_update_cell_text(rl, original_text):
    table = rl.tables()[0]
    rl.set_cell_text(table.cell(1, 2), "$27,500")
    assert "$27,500" in resolved(rl, "accept")
    assert resolved(rl, "reject") == original_text


def test_row_index_out_of_range(rl):
    with pytest.raises(RedlineError):
        rl.delete_table_row(rl.tables()[0], 99)


# ---------------------------------------------------------------------------
# formatting revisions
# ---------------------------------------------------------------------------
def test_run_formatting_revision(rl):
    assert rl.format_matching("Late Payment", bold=True, color="C00000") == 1
    assert rl.summary().counts["format:rPrChange"] == 1
    other = Redliner(roundtrip(rl), track_changes=False)
    other.accept_all()
    run = next(
        r for r in other.find_paragraph(contains="Late Payment").runs if r.text == "Late Payment"
    )
    assert run.bold is True


def test_run_formatting_revision_reverts_on_reject(rl):
    rl.format_matching("Late Payment", bold=True)
    other = Redliner(roundtrip(rl), track_changes=False)
    other.reject_all()
    run = next(
        r for r in other.find_paragraph(contains="Late Payment").runs if "Late Payment" in r.text
    )
    assert not run.bold


def test_paragraph_formatting_revision(rl):
    para = rl.find_paragraph(contains="1. Fees")
    rl.format_paragraph(para, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(18))
    assert rl.summary().counts["format:pPrChange"] == 1
    accepted = Redliner(roundtrip(rl), track_changes=False)
    accepted.accept_all()
    assert accepted.find_paragraph(contains="1. Fees").alignment == WD_ALIGN_PARAGRAPH.CENTER
    rejected = Redliner(roundtrip(rl), track_changes=False)
    rejected.reject_all()
    assert rejected.find_paragraph(contains="1. Fees").alignment is None


def test_unknown_paragraph_property_raises(rl):
    with pytest.raises(RedlineError):
        rl.format_paragraph(rl.find_paragraph(contains="1. Fees"), nonsense=1)


# ---------------------------------------------------------------------------
# revision metadata
# ---------------------------------------------------------------------------
def test_author_and_date_are_stamped(rl):
    rl.replace_text("thirty (30) days", "45 days")
    xml = document_xml(rl)
    assert 'w:author="Tester"' in xml
    assert f'w:date="{DATE}"' in xml


def test_revision_ids_are_unique(rl):
    rl.replace_text("thirty (30) days", "45 days")
    rl.delete_paragraph(rl.find_paragraph(contains="Late Payment"))
    rl.insert_table_row(rl.tables()[0], values=["A", "B", "C"])
    ids = [
        el.get(qn("w:id"))
        for el in rl.document.element.body.iter()
        if el.tag in (qn("w:ins"), qn("w:del"))
    ]
    assert len(ids) == len(set(ids))


def test_ids_do_not_collide_with_a_second_pass(rl):
    rl.replace_text("thirty (30) days", "45 days")
    second = Redliner(roundtrip(rl), author="Second", date=DATE)
    second.delete_paragraph(second.find_paragraph(contains="Late Payment"))
    ids = [
        el.get(qn("w:id"))
        for el in second.document.element.body.iter()
        if el.tag in (qn("w:ins"), qn("w:del"))
    ]
    assert len(ids) == len(set(ids))


def test_track_changes_flag(rl):
    assert rl.document.settings.element.find(qn("w:trackChanges")) is not None
    rl.enable_track_changes(False)
    assert rl.document.settings.element.find(qn("w:trackChanges")) is None


def test_summary_reports_kinds(rl):
    rl.replace_text("thirty (30) days", "45 days")
    rl.delete_paragraph(rl.find_paragraph(contains="Late Payment"))
    counts = rl.summary().counts
    assert counts["insert"] and counts["delete"] and counts["paragraph-mark-delete"]
    assert "Tester" in rl.summary().authors


# ---------------------------------------------------------------------------
# combined / stress
# ---------------------------------------------------------------------------
def test_every_operation_together(rl, original_text):
    rl.replace_text("thirty (30) days", "forty-five (45) days")
    rl.insert_text_after("invoice date", ", without setoff")
    rl.delete_matching("at 1.5% per month")
    rl.set_paragraph_text(
        rl.find_paragraph(contains="Termination"),
        "4. Termination. Either Party may terminate for material breach on 30 days notice.",
    )
    rl.insert_paragraph_after(rl.find_paragraph(contains="Termination"), "5. Audit rights apply.")
    rl.delete_paragraph(rl.find_paragraph(contains="1. Fees"))
    rl.insert_table_row(rl.tables()[0], values=["Enterprise", "500", "$180,000"])
    rl.delete_table_row(rl.tables()[0], 1)
    rl.set_cell_text(rl.tables()[0].cell(0, 2), "Annual Fee")
    rl.format_matching("Audit rights", bold=True)
    rl.format_paragraph(rl.paragraphs(include_tables=False)[0], alignment=WD_ALIGN_PARAGRAPH.CENTER)

    assert len(rl.summary()) > 10
    assert resolved(rl, "reject") == original_text
    accepted = "\n".join(resolved(rl, "accept"))
    assert "forty-five (45) days" in accepted
    assert "5. Audit rights apply." in accepted
    assert "1. Fees" not in accepted
    assert "Enterprise" in accepted


def test_repeated_passes_by_different_authors(rl, original_text):
    rl.replace_text("thirty (30) days", "forty-five (45) days")
    second = Redliner(roundtrip(rl), author="Opposing Counsel", date=DATE)
    second.replace_text("Late Payment", "Overdue Amounts")
    third = Redliner(roundtrip(second), author="Client", date=DATE)
    third.insert_paragraph_after(third.find_paragraph(contains="Termination"), "5. Survival.")
    assert set(third.summary().authors) == {"Tester", "Opposing Counsel", "Client"}
    assert resolved(third, "reject") == original_text


def test_accept_then_reject_is_stable(rl, original_text):
    rl.replace_text("thirty (30) days", "45 days")
    accepted = Redliner(roundtrip(rl), track_changes=False)
    accepted.accept_all()
    again = Redliner(roundtrip(accepted), track_changes=False)
    again.reject_all()  # nothing left to reject
    assert texts(again.document) == texts(accepted.document)


# ---------------------------------------------------------------------------
# declarative plans
# ---------------------------------------------------------------------------
def test_json_plan(rl, original_text):
    apply_operations(
        rl,
        [
            {"op": "replace_text", "old": "thirty (30) days", "new": "forty-five (45) days"},
            {"op": "insert_paragraph", "match": "Late Payment", "text": "3A. Interest cap."},
            {"op": "delete_paragraph", "match": "Termination"},
            {"op": "insert_row", "table": 0, "values": ["Enterprise", "500", "$180,000"]},
            {"op": "update_cell", "table": 0, "row": 1, "col": 2, "text": "$27,500"},
            {"op": "format_text", "match": "Fees", "bold": True},
            {"op": "format_paragraph", "match": "1. Fees", "alignment": "CENTER"},
        ],
    )
    accepted = "\n".join(resolved(rl, "accept"))
    assert "forty-five (45) days" in accepted and "3A. Interest cap." in accepted
    assert resolved(rl, "reject") == original_text


def test_plan_validation_errors(rl):
    with pytest.raises(RedlineError, match="unknown op"):
        apply_operations(rl, [{"op": "nope"}])
    with pytest.raises(RedlineError, match="missing required key"):
        apply_operations(rl, [{"op": "replace_text", "old": "x"}])
    with pytest.raises(RedlineError, match="matched nothing"):
        apply_operations(rl, [{"op": "replace_text", "old": "not in doc", "new": "y"}])


def test_plan_lenient_mode(rl):
    results = apply_operations(
        rl, [{"op": "replace_text", "old": "absent", "new": "y"}], strict=False
    )
    assert results[0].applied == 0


# ---------------------------------------------------------------------------
# document compare
# ---------------------------------------------------------------------------
@pytest.fixture
def pair(tmp_path, contract_bytes):
    original = tmp_path / "original.docx"
    original.write_bytes(contract_bytes.getvalue())

    d = docx.Document(str(original))
    for para in d.paragraphs:
        if "thirty (30) days" in para.text:
            for run in para.runs:
                run.text = run.text.replace("thirty (30) days", "forty-five (45) days")
    late = next(p for p in d.paragraphs if "Late Payment" in p.text)
    late._p.getparent().remove(late._p)
    d.paragraphs[-1].insert_paragraph_before("3A. Insurance. Provider shall carry cyber cover.")
    revised = tmp_path / "revised.docx"
    d.save(str(revised))
    return original, revised


def _clean(document):
    return [t.strip() for t in texts(document) if t.strip()]


def test_compare_documents(tmp_path, pair):
    original, revised = pair
    out = tmp_path / "redlined.docx"
    stats = redline_files(original, revised, out, author="Compare", date=DATE)
    assert stats.paragraphs_changed or stats.paragraphs_inserted or stats.paragraphs_deleted

    accepted = Redliner(out, track_changes=False)
    accepted.accept_all()
    assert _clean(accepted.document) == _clean(docx.Document(str(revised)))

    rejected = Redliner(out, track_changes=False)
    rejected.reject_all()
    assert _clean(rejected.document) == _clean(docx.Document(str(original)))


def test_compare_identical_documents(tmp_path, pair):
    original, _ = pair
    out = tmp_path / "same.docx"
    redline_files(original, original, out, date=DATE)
    assert len(Redliner(out, track_changes=False).summary()) == 0


# ---------------------------------------------------------------------------
# diff engine
# ---------------------------------------------------------------------------
@pytest.mark.parametrize(
    "old,new",
    [
        ("", "brand new text"),
        ("all gone", ""),
        ("same", "same"),
        ("thirty (30) days", "forty-five (45) days"),
        ("a b c d e", "a x c y e"),
        ("word", "word word word"),
    ],
)
def test_diff_ops_reconstruct_the_target(old, new):
    result, shift = old, 0
    for op in diff_ops(old, new):
        start, end = op.start + shift, op.end + shift
        result = result[:start] + op.text + result[end:]
        shift += len(op.text) - (end - start)
    assert result.replace(" ", "") == new.replace(" ", "")


def test_diff_of_identical_text_is_empty():
    assert diff_ops("hello", "hello") == []


# ---------------------------------------------------------------------------
# lookup helpers
# ---------------------------------------------------------------------------
def test_find_helpers(rl):
    assert len(rl.find_paragraphs(contains="Payment")) == 2
    assert len(rl.find_paragraphs(regex=r"^\d\. ")) == 4
    assert rl.find_paragraph(startswith="3.").text.startswith("3.")
    assert len(rl.find_text("Payment")) == 2
    with pytest.raises(RedlineError):
        rl.find_paragraph(contains="does not exist")


def test_paragraphs_excludes_tables(rl):
    assert len(rl.paragraphs(include_tables=False)) == 4
    assert len(rl.paragraphs(include_tables=True)) == 13


def test_text_of_sees_insertions_and_hides_deletions(rl):
    para = rl.find_paragraph(contains="Late Payment")
    rl.insert_text_after("Late Payment", " (Overdue Amounts)")
    rl.delete_matching("Interest accrues")
    assert "(Overdue Amounts)" in rl.text_of(para)
    assert "Interest accrues" not in rl.text_of(para)
    # python-docx's own view is the one that gets this wrong
    assert "(Overdue Amounts)" not in para.text


def test_inserted_paragraph_is_a_real_python_docx_paragraph(rl):
    para = rl.insert_paragraph_after(
        rl.find_paragraph(contains="Late Payment"), "3A. New.", style="Heading 2"
    )
    assert para.style.name == "Heading 2"
    assert rl.text_of(para) == "3A. New."


def test_inserted_list_item_keeps_list_style(contract_bytes):
    import io

    d = docx.Document(contract_bytes)
    for text in ("First duty.", "Second duty."):
        d.add_paragraph(text, style="List Number")
    buf = io.BytesIO()
    d.save(buf)
    buf.seek(0)

    rl = Redliner(buf, author="Tester", date=DATE)
    rl.insert_paragraph_after(rl.find_paragraph(contains="First duty"), "Inserted duty.")
    accepted = Redliner(roundtrip(rl), track_changes=False)
    accepted.accept_all()
    inserted = accepted.find_paragraph(contains="Inserted duty")
    assert inserted.style.name == "List Number"


def test_delete_last_paragraph_removes_it_on_accept(rl, original_text):
    rl.delete_paragraph(rl.paragraphs(include_tables=False)[-1])
    accepted = resolved(rl, "accept")
    assert "material breach" not in "\n".join(accepted)
    assert resolved(rl, "reject") == original_text
