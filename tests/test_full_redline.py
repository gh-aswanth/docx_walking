"""``full_redline`` -- every layer over one document, in one call.

What is under test is mostly *ordering*: the compare has to land before the
scripted actions address the document, comments have to attach after every
edit, and clause renumbering has to run once at the end over the combined
result, measured against a baseline captured before any of it.
"""

import json

import docx
import pytest
from conftest import DATE, PLAN, SAMPLE, texts

from docx_redline import Redliner, full_redline
from docx_redline.errors import RedlineError
from docx_redline.planning.pipeline import _numbering_gaps
from docx_redline.structure.clauses import ClauseTree


@pytest.fixture
def revised(tmp_path, agreement_bytes):
    """A counterparty's edit of the same agreement, for the compare stage."""
    d = docx.Document(agreement_bytes)
    for para in d.paragraphs:
        if "thirty (30) days" in para.text:
            for run in para.runs:
                run.text = run.text.replace("thirty (30) days", "sixty (60) days")
        if "Notices must be in writing" in para.text:
            for run in para.runs:
                run.text = run.text.replace(
                    "Notices must be in writing.", "Notices must be in writing and couriered."
                )
    d.paragraphs[-1].insert_paragraph_before("Exhibit B. Security questionnaire is incorporated.")
    path = tmp_path / "revised.docx"
    d.save(str(path))
    return path


def accepted(path):
    rl = Redliner(path, track_changes=False)
    rl.accept_all()
    return rl


def rejected(path):
    rl = Redliner(path, track_changes=False)
    rl.reject_all()
    return rl


def rationale_texts(path) -> list[str]:
    """Comments the planner wrote to explain an action (``[AI-001 · high] ...``)."""
    return [t for t in comment_texts(path) if t.startswith("[")]


def review_texts(path) -> list[str]:
    """Comments a caller asked for explicitly, as opposed to rationales."""
    return [t for t in comment_texts(path) if not t.startswith("[")]


def comment_texts(path) -> list[str]:
    import re
    import zipfile

    with zipfile.ZipFile(path) as zf:
        if "word/comments.xml" not in zf.namelist():
            return []
        xml = zf.read("word/comments.xml").decode()
    return [
        "".join(re.findall(r"<w:t[^>]*>([^<]*)</w:t>", body))
        for body in re.findall(r"<w:comment [^>]*>(.*?)</w:comment>", xml, re.S)
    ]


# ---------------------------------------------------------------------------
# each source on its own
# ---------------------------------------------------------------------------
def test_compare_only(tmp_path, agreement, revised, agreement_bytes):
    out = tmp_path / "out.docx"
    result = full_redline(agreement, out, revised=revised, date=DATE)
    assert result.ok, result.format()
    assert result.compare.paragraphs_changed >= 2
    assert result.proposal.source == "none"

    text = "\n".join(texts(accepted(out).document))
    assert "sixty (60) days" in text
    assert "couriered" in text
    assert "Security questionnaire" in text
    assert texts(rejected(out).document) == texts(docx.Document(agreement_bytes))


def test_actions_only(tmp_path, agreement):
    out = tmp_path / "out.docx"
    result = full_redline(
        agreement,
        out,
        date=DATE,
        actions=[{"id": "A", "type": "move_clause", "clause": "4.1", "after_clause": "3.1"}],
    )
    assert result.ok
    assert result.proposal.source == "inline"
    assert {c["from"]: c["to"] for c in result.plan.renumbered}["4.1"] == "3.2"


def test_comments_only(tmp_path, agreement, agreement_bytes):
    out = tmp_path / "out.docx"
    result = full_redline(
        agreement,
        out,
        date=DATE,
        comments=[{"clause": "2.2", "text": "Check with finance."}],
    )
    assert result.ok
    assert review_texts(out) == ["Check with finance."]
    # a comment is an annotation, not an edit -- the text must be untouched
    assert texts(accepted(out).document) == texts(docx.Document(agreement_bytes))


def test_reviewer_only(tmp_path, agreement):
    out = tmp_path / "out.docx"
    result = full_redline(agreement, out, reviewer="rules", date=DATE)
    assert result.ok
    assert result.proposal.source == "rule-based"
    assert result.plan.applied >= 1


def test_nothing_to_do_is_an_error(tmp_path, agreement):
    with pytest.raises(RedlineError, match="needs something to do"):
        full_redline(agreement, tmp_path / "out.docx")


# ---------------------------------------------------------------------------
# all of it together
# ---------------------------------------------------------------------------
@pytest.fixture
def combined(tmp_path, agreement, revised):
    out = tmp_path / "combined.docx"
    result = full_redline(
        agreement,
        out,
        revised=revised,
        actions=[
            {
                "id": "A1",
                "type": "move_clause",
                "clause": "4.1",
                "after_clause": "3.1",
                "rationale": "governing law with term",
                "severity": "medium",
            },
            {
                "id": "A2",
                "type": "insert_clause",
                "after_clause": "2.1",
                "title": "Currency",
                "text": "All amounts are in U.S. dollars.",
                "rationale": "clarity",
                "severity": "low",
            },
        ],
        comments=[
            {"clause": "2.2", "text": "Confirm the payment window."},
            {"find": "couriered", "text": "Counterparty added courier delivery."},
        ],
        date=DATE,
        report_path=tmp_path / "report.json",
    )
    return result, out


def test_full_run_passes_every_check(combined):
    result, _ = combined
    assert result.ok, result.format()
    assert [name for name, _, _ in result.checks]
    assert all(passed for _, passed, _ in result.checks)


def test_full_run_stage_order(combined):
    """Compare must precede the plan; renumbering must come after both."""
    result, _ = combined
    assert [s.name for s in result.stages] == [
        "extract",
        "propose",
        "validate",
        "compare",
        "plan",
        "renumber",
        "apply",
        "verify",
        "report",
    ]


def test_full_run_applies_all_three_sources(combined):
    _result, out = combined
    text = "\n".join(texts(accepted(out).document))
    assert "sixty (60) days" in text  # from the compare
    assert "couriered" in text  # from the compare
    assert "All amounts are in U.S. dollars." in text  # from the action plan
    assert sorted(review_texts(out)) == [
        "Confirm the payment window.",
        "Counterparty added courier delivery.",
    ]


def test_full_run_renumbers_over_the_combined_result(combined):
    """One renumbering pass, covering the compare and the plan together."""
    result, out = combined
    mapping = {c["from"]: c["to"] for c in result.plan.renumbered}
    assert mapping["4.1"] == "3.2"  # the move
    assert mapping["2.2"] == "2.3"  # pushed down by the inserted clause
    tree = ClauseTree(accepted(out).document.element.body)
    assert not _numbering_gaps(tree)
    assert tree.get("2.2").title == "Currency"
    assert tree.get("3.2").title == "Governing Law"


def test_full_run_rejects_back_to_the_original(combined, agreement_bytes):
    _, out = combined
    restored = rejected(out)
    assert texts(restored.document) == texts(docx.Document(agreement_bytes))
    assert [c.label for c in ClauseTree(restored.document.element.body)] == [
        c.label for c in ClauseTree(docx.Document(agreement_bytes).element.body)
    ]


def test_full_run_writes_a_report(combined, tmp_path):
    _result, _ = combined
    payload = json.loads((tmp_path / "report.json").read_text())
    assert payload["ok"] is True
    assert payload["compare"]["paragraphs_changed"] >= 2
    assert payload["plan"]["summary"]["clauses_renumbered"] >= 2
    assert [s["stage"] for s in payload["stages"]][:4] == [
        "extract",
        "propose",
        "validate",
        "compare",
    ]


# ---------------------------------------------------------------------------
# ordering guarantees
# ---------------------------------------------------------------------------
def test_actions_address_the_document_the_compare_produced(tmp_path, agreement, revised):
    """The compare rewrites 2.2's text; a later action must see the new text."""
    out = tmp_path / "out.docx"
    result = full_redline(
        agreement,
        out,
        revised=revised,
        date=DATE,
        actions=[
            {
                "id": "A",
                "type": "replace_text",
                "clause": "2.2",
                "find": "sixty (60) days",
                "replace": "ninety (90) days",
            }
        ],
    )
    assert result.ok, result.format()
    assert "ninety (90) days" in "\n".join(texts(accepted(out).document))


def test_comments_anchor_after_the_edits(tmp_path, agreement, revised):
    """A comment may target text that only exists because of the compare."""
    out = tmp_path / "out.docx"
    result = full_redline(
        agreement,
        out,
        revised=revised,
        date=DATE,
        comments=[{"find": "couriered", "text": "New wording."}],
    )
    assert result.ok
    assert review_texts(out) == ["New wording."]


def test_comment_on_a_clause_moved_by_an_action(tmp_path, agreement):
    """The comment must follow the clause, not the number it used to have."""
    out = tmp_path / "out.docx"
    full_redline(
        agreement,
        out,
        date=DATE,
        actions=[{"id": "A", "type": "move_clause", "clause": "4.1", "after_clause": "3.1"}],
        comments=[{"clause": "4.1", "text": "Jurisdiction check."}],
    )
    assert review_texts(out) == ["Jurisdiction check."]


def test_unanchorable_comment_is_reported_not_raised(tmp_path, agreement):
    out = tmp_path / "out.docx"
    result = full_redline(
        agreement,
        out,
        date=DATE,
        comments=[{"find": "not in this document", "text": "..."}],
    )
    assert result.plan.failed == 1
    assert "cannot anchor a comment" in result.plan.results[0].detail
    assert not result.ok


# ---------------------------------------------------------------------------
# inputs
# ---------------------------------------------------------------------------
def test_actions_from_a_json_path(tmp_path, agreement):
    path = tmp_path / "actions.json"
    path.write_text(
        json.dumps(
            {
                "action_items": [
                    {"id": "A", "type": "move_clause", "clause": "4.1", "after_clause": "3.1"}
                ]
            }
        )
    )
    result = full_redline(agreement, tmp_path / "out.docx", actions=path, date=DATE)
    assert result.ok
    assert result.proposal.source == "file"


def test_renumbering_can_be_switched_off(tmp_path, agreement, revised):
    result = full_redline(
        agreement,
        tmp_path / "out.docx",
        revised=revised,
        date=DATE,
        renumber=False,
        actions=[{"id": "A", "type": "move_clause", "clause": "4.1", "after_clause": "3.1"}],
    )
    assert result.plan.renumbered == []
    assert [name for name, _, _ in result.checks] == [
        "redline reopens as a valid docx",
        "reject restores the original document",
        "accept changes the document",
        "no tracked changes survive accept",
        "reject restores the original numbering",
    ]


def test_strict_mode_raises_on_a_bad_action(tmp_path, agreement):
    with pytest.raises(RedlineError):
        full_redline(
            agreement,
            tmp_path / "out.docx",
            date=DATE,
            strict=True,
            actions=[{"id": "A", "type": "delete_clause", "clause": "99.9"}],
        )


# ---------------------------------------------------------------------------
# passing action items in
# ---------------------------------------------------------------------------
MOVE = {"id": "A", "type": "move_clause", "clause": "4.1", "after_clause": "3.1"}


def test_actions_as_a_list(tmp_path, agreement):
    result = full_redline(agreement, tmp_path / "out.docx", actions=[MOVE], date=DATE)
    assert result.ok and result.proposal.source == "inline"


def test_actions_as_a_path_object(tmp_path, agreement):
    path = tmp_path / "plan.json"
    path.write_text(json.dumps({"action_items": [MOVE]}))
    result = full_redline(agreement, tmp_path / "out.docx", actions=path, date=DATE)
    assert result.ok and result.proposal.source == "file"


def test_actions_as_a_bare_list_file(tmp_path, agreement):
    """A plan file may be a bare JSON array, not just {"action_items": [...]}."""
    path = tmp_path / "plan.json"
    path.write_text(json.dumps([MOVE]))
    result = full_redline(agreement, tmp_path / "out.docx", actions=str(path), date=DATE)
    assert result.ok and result.plan.applied == 1


def test_actions_file_is_never_rewritten(tmp_path, agreement):
    """A curated plan is an input; normalising it back over the user's file
    would silently reformat it."""
    path = tmp_path / "plan.json"
    original = json.dumps({"action_items": [MOVE]}, indent=4)
    path.write_text(original)
    full_redline(agreement, tmp_path / "out.docx", actions=path, date=DATE)
    assert path.read_text() == original


def test_supplied_actions_win_over_a_reviewer(tmp_path, agreement):
    result = full_redline(
        agreement, tmp_path / "out.docx", actions=[MOVE], reviewer="rules", date=DATE
    )
    assert result.proposal.source == "inline"
    assert [r.id for r in result.plan.results] == ["A"]


def test_malformed_action_items_never_reach_the_document(tmp_path, agreement):
    with pytest.raises(RedlineError, match="unknown action type"):
        full_redline(
            agreement,
            tmp_path / "out.docx",
            date=DATE,
            actions=[{"id": "X", "type": "nope"}],
        )
    assert not (tmp_path / "out.docx").exists()


# ---------------------------------------------------------------------------
# the `full` CLI subcommand
# ---------------------------------------------------------------------------
def run_cli(*argv) -> int:
    from docx_redline.cli import main

    return main([str(a) for a in argv])


def test_cli_full_with_a_plan_file(tmp_path, agreement, capsys):
    plan = tmp_path / "plan.json"
    plan.write_text(json.dumps({"action_items": [MOVE]}))
    out = tmp_path / "out.docx"
    assert run_cli("full", agreement, "-o", out, "--actions", plan, "--date", DATE) == 0
    assert out.exists()
    assert "4.1" in capsys.readouterr().out


def test_cli_full_with_inline_actions(tmp_path, agreement):
    out = tmp_path / "out.docx"
    code = run_cli(
        "full",
        agreement,
        "-o",
        out,
        "--action",
        json.dumps(MOVE),
        "--action",
        json.dumps({"type": "reorder_clauses", "section": "2", "order": ["2.3", "2.1", "2.2"]}),
        "--date",
        DATE,
    )
    assert code == 0
    tree = ClauseTree(accepted(out).document.element.body)
    assert tree.get("2.1").title == "Taxes"
    assert tree.get("3.2").title == "Governing Law"


def test_cli_full_combines_every_source(tmp_path, agreement, revised):
    out = tmp_path / "out.docx"
    report = tmp_path / "report.json"
    code = run_cli(
        "full",
        agreement,
        "-o",
        out,
        "--revised",
        revised,
        "--action",
        json.dumps(MOVE),
        "--comment",
        "2.2=Confirm the payment window.",
        "--comment",
        "find=couriered=Counterparty wording.",
        "--report",
        report,
        "--date",
        DATE,
    )
    assert code == 0
    payload = json.loads(report.read_text())
    assert payload["ok"] is True
    assert payload["compare"]["paragraphs_changed"] >= 1
    assert sorted(review_texts(out)) == ["Confirm the payment window.", "Counterparty wording."]


def test_cli_full_with_a_reviewer(tmp_path, agreement):
    out = tmp_path / "out.docx"
    assert run_cli("full", agreement, "-o", out, "--reviewer", "rules", "--date", DATE) == 0
    assert out.exists()


def test_cli_full_needs_something_to_do(tmp_path, agreement, capsys):
    assert run_cli("full", agreement, "-o", tmp_path / "out.docx") == 2
    assert "nothing to do" in capsys.readouterr().err


def test_cli_full_rejects_bad_inline_json(tmp_path, agreement):
    with pytest.raises(RedlineError, match="not valid JSON"):
        run_cli("full", agreement, "-o", tmp_path / "out.docx", "--action", "{not json}")


def test_cli_full_rejects_a_non_object_action(tmp_path, agreement):
    with pytest.raises(RedlineError, match="must be a JSON object"):
        run_cli("full", agreement, "-o", tmp_path / "out.docx", "--action", '["a"]')


def test_cli_full_fills_in_action_metadata(tmp_path, agreement):
    """Inline actions need not repeat id/rationale/severity by hand."""
    from docx_redline.cli import _inline_actions

    [item] = _inline_actions([json.dumps({"type": "delete_clause", "clause": "2.3"})])
    assert item["id"] == "CLI-001"
    assert item["severity"] == "medium"
    assert item["rationale"]


def test_cli_full_rejects_a_bad_comment_flag(tmp_path, agreement):
    with pytest.raises(RedlineError, match="CLAUSE=TEXT"):
        run_cli("full", agreement, "-o", tmp_path / "out.docx", "--comment", "no-equals-sign")


def test_cli_full_returns_nonzero_when_an_action_misses(tmp_path, agreement, capsys):
    out = tmp_path / "out.docx"
    code = run_cli(
        "full",
        agreement,
        "-o",
        out,
        "--action",
        json.dumps(
            {"type": "replace_text", "clause": "2.2", "find": "absent phrase", "replace": "x"}
        ),
        "--date",
        DATE,
    )
    assert code == 1
    assert "FAIL" in capsys.readouterr().out


# ---------------------------------------------------------------------------
# the shipped demo plan
# ---------------------------------------------------------------------------
DEMO_PLAN = PLAN


@pytest.fixture
def demo_items():
    if not DEMO_PLAN.exists():
        pytest.skip(f"{DEMO_PLAN.name} not present")
    return json.loads(DEMO_PLAN.read_text())["action_items"]


def test_demo_plan_covers_every_action_type(demo_items):
    """The shipped plan is the worked example -- it must not lose coverage."""
    from docx_redline.planning.actions import ACTION_SCHEMA

    assert {item["type"] for item in demo_items} == set(ACTION_SCHEMA)


def test_demo_plan_uses_every_action_field(demo_items):
    """Every field the model schema advertises is demonstrated at least once."""
    from docx_redline.planning.agent import ACTION_FIELDS

    used = {key for item in demo_items for key in item}
    assert set(ACTION_FIELDS) - used == set()


def test_demo_plan_is_schema_clean(demo_items):
    from docx_redline.planning.actions import validate_actions

    assert validate_actions(demo_items) == []


@pytest.mark.skipif(not SAMPLE.exists(), reason=f"{SAMPLE.name} not present")
def test_demo_plan_applies_cleanly(tmp_path, demo_items):
    """Every item must actually land on the real contract, not just validate."""
    out = tmp_path / "out.docx"
    result = full_redline(SAMPLE, out, actions=demo_items, date=DATE)
    assert result.plan.failed == 0, [r.detail for r in result.plan.results if r.status == "failed"]
    assert result.ok, result.format()
    assert result.plan.renumbered and result.plan.references
    explicit = review_texts(out)
    assert len(explicit) == sum(1 for i in demo_items if i["type"] == "comment"), explicit
    # every applied action with a rationale explains itself in the document
    assert len(rationale_texts(out)) == sum(
        1
        for r in result.plan.results
        if r.status == "applied" and r.rationale and r.type != "comment"
    )


# ---------------------------------------------------------------------------
# rationales in the document
# ---------------------------------------------------------------------------
EXPLAINED = [
    {
        "id": "AI-001",
        "type": "replace_text",
        "clause": "2.2",
        "find": "thirty (30) days",
        "replace": "forty-five (45) days",
        "rationale": "Standard 45-day cycle.",
        "severity": "high",
    },
    {
        "id": "AI-002",
        "type": "move_clause",
        "clause": "4.1",
        "after_clause": "3.1",
        "rationale": "Governing law with the term provisions.",
        "severity": "medium",
    },
]


def test_rationale_is_written_into_the_document(tmp_path, agreement):
    """The *why* has to reach the .docx, not just the JSON report."""
    out = tmp_path / "out.docx"
    full_redline(agreement, out, actions=EXPLAINED, date=DATE)
    assert sorted(rationale_texts(out)) == [
        "[AI-001 · high] Standard 45-day cycle.",
        "[AI-002 · medium] Governing law with the term provisions.",
    ]


def test_rationale_can_be_switched_off(tmp_path, agreement):
    out = tmp_path / "out.docx"
    full_redline(agreement, out, actions=EXPLAINED, date=DATE, explain=False)
    assert comment_texts(out) == []


def test_rationale_is_off_by_default_in_the_low_level_api(rl_agreement):
    """`apply_actions` is a primitive; it must not start writing comments."""
    from docx_redline.planning.actions import apply_actions

    apply_actions(rl_agreement, EXPLAINED)
    assert comment_texts_from(rl_agreement) == []
    apply_actions(rl_agreement, [], explain=True)  # no-op, still no comments
    assert comment_texts_from(rl_agreement) == []


def comment_texts_from(rl) -> list[str]:
    import io

    buf = io.BytesIO()
    rl.document.save(buf)
    buf.seek(0)
    return comment_texts(buf)


def test_actions_without_a_rationale_are_not_annotated(tmp_path, agreement):
    out = tmp_path / "out.docx"
    full_redline(
        agreement,
        out,
        date=DATE,
        actions=[{"id": "A", "type": "move_clause", "clause": "4.1", "after_clause": "3.1"}],
    )
    assert comment_texts(out) == []


def test_failed_actions_are_not_annotated(tmp_path, agreement):
    out = tmp_path / "out.docx"
    result = full_redline(
        agreement,
        out,
        date=DATE,
        actions=[
            {
                "id": "A",
                "type": "replace_text",
                "clause": "2.2",
                "find": "absent",
                "replace": "x",
                "rationale": "Should not appear.",
                "severity": "low",
            }
        ],
    )
    assert result.plan.failed == 1
    assert comment_texts(out) == []


def test_a_comment_action_is_not_explained_twice(tmp_path, agreement):
    """A rationale on a `comment` action would be a comment about a comment."""
    out = tmp_path / "out.docx"
    full_redline(
        agreement,
        out,
        date=DATE,
        comments=[{"clause": "2.2", "text": "Check with finance."}],
        actions=[
            {
                "id": "C",
                "type": "comment",
                "clause": "2.1",
                "text": "And here.",
                "rationale": "Explaining a comment helps nobody.",
                "severity": "low",
            }
        ],
    )
    assert sorted(comment_texts(out)) == ["And here.", "Check with finance."]


def test_rationale_for_a_deleted_clause_anchors_on_the_strikeout(tmp_path, agreement):
    """A struck paragraph has no visible runs; the comment still has to land."""
    out = tmp_path / "out.docx"
    result = full_redline(
        agreement,
        out,
        date=DATE,
        actions=[
            {
                "id": "D",
                "type": "delete_clause",
                "clause": "2.2",
                "rationale": "Superseded by the Order Form.",
                "severity": "medium",
            }
        ],
    )
    assert result.plan.failed == 0
    assert rationale_texts(out) == ["[D · medium] Superseded by the Order Form."]


def test_rationale_survives_renumbering(tmp_path, agreement):
    """Annotation runs after renumbering, so it cannot anchor on stale text."""
    out = tmp_path / "out.docx"
    full_redline(
        agreement,
        out,
        date=DATE,
        actions=[
            {
                "id": "M",
                "type": "move_clause",
                "clause": "4.1",
                "after_clause": "3.1",
                "rationale": "Reorder for readability.",
                "severity": "low",
            }
        ],
    )
    assert rationale_texts(out) == ["[M · low] Reorder for readability."]
    tree = ClauseTree(accepted(out).document.element.body)
    assert tree.get("3.2").title == "Governing Law"


def test_report_carries_the_rationale_too(tmp_path, agreement):
    report = tmp_path / "report.json"
    full_redline(agreement, tmp_path / "out.docx", actions=EXPLAINED, date=DATE, report_path=report)
    actions = json.loads(report.read_text())["plan"]["actions"]
    assert actions[0]["rationale"] == "Standard 45-day cycle."
    assert actions[0]["severity"] == "high"
