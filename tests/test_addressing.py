"""Clause addressing: the guarantees every other layer resolves through.

Each test here corresponds to a defect that reported ``applied`` while doing the
wrong thing. They are grouped by the invariant they protect rather than by the
function they touch, because the failures were only visible one layer up.
"""

import io
import itertools

import docx
import pytest
from conftest import DATE, texts

from docx_redline import Redliner
from docx_redline.oxml.ns import qn
from docx_redline.oxml.textmap import paragraph_text
from docx_redline.planning.actions import ActionPlanner, apply_actions
from docx_redline.structure.clauses import ClauseTree, parse_clause


def accepted(rl: Redliner) -> Redliner:
    buf = io.BytesIO()
    rl.document.save(buf)
    buf.seek(0)
    other = Redliner(buf, track_changes=False)
    other.accept_all()
    return other


def titles_of(document, section: str) -> list[str]:
    return [c.title for c in ClauseTree(document.element.body).get(section).children]


# ---------------------------------------------------------------------------
# a clause number resolves to the clause the reviewer meant
# ---------------------------------------------------------------------------
def test_insert_clause_makes_the_duplicate_visible(rl_agreement):
    """A new clause borrows its neighbour's number until renumbering runs."""
    planner = ActionPlanner(rl_agreement, renumber=False)
    planner.run([{"id": "I", "type": "insert_clause", "before_clause": "2.2", "text": "New."}])
    assert planner.tree.duplicates == {"2.2"}


def test_later_action_resolves_to_the_original_not_the_insertion(rl_agreement):
    """The bug: `tree.get` was first-match-wins, so the insert stole the number."""
    planner = ActionPlanner(rl_agreement, renumber=False)
    planner.run(
        [
            {"id": "I", "type": "insert_clause", "before_clause": "2.2", "text": "INSERTED."},
            {
                "id": "R",
                "type": "rewrite_clause",
                "clause": "2.2",
                "text": "Invoicing. Payment is due on receipt.",
            },
        ]
    )
    rewritten = planner.tree.get("2.2")
    assert "INSERTED." not in rewritten.text
    assert "due on receipt" in "\n".join(texts(accepted(rl_agreement).document))


def test_renumbering_warns_about_the_duplicate(rl_agreement):
    report = apply_actions(
        rl_agreement, [{"id": "I", "type": "insert_clause", "before_clause": "2.2", "text": "New."}]
    )
    assert any("carried by more than one paragraph" in w for w in report.warnings)


# ---------------------------------------------------------------------------
# an inserted clause is visible to the parser that has to renumber it
# ---------------------------------------------------------------------------
@pytest.mark.parametrize("anchor,expected_level", [("2", 1), ("2.1", 2)])
def test_inserted_clause_round_trips_through_the_parser(rl_agreement, anchor, expected_level):
    """`"4  "` has no trailing dot, so SECTION_RE could not read it back and the
    new paragraph was invisible to the clause tree -- and never renumbered."""
    planner = ActionPlanner(rl_agreement, renumber=False)
    planner.run(
        [
            {
                "id": "I",
                "type": "insert_clause",
                "after_clause": anchor,
                "title": "Currency",
                "text": "Amounts are in U.S. dollars.",
            }
        ]
    )
    added = [
        p
        for p in rl_agreement.document.element.body.iter(qn("w:p"))
        if "Amounts are in U.S. dollars." in paragraph_text(p)
    ]
    assert len(added) == 1
    clause = parse_clause(added[0])
    assert clause is not None, paragraph_text(added[0])
    assert clause.level == expected_level


def test_inserted_section_is_renumbered(rl_agreement):
    apply_actions(
        rl_agreement,
        [
            {
                "id": "I",
                "type": "insert_section",
                "after_section": "2",
                "title": "Insurance",
                "text": "Provider shall carry cyber cover.",
            }
        ],
    )
    tree = ClauseTree(accepted(rl_agreement).document.element.body)
    assert tree.get("3").title == "Insurance"
    assert [s.label for s in tree.sections] == ["1", "2", "3", "4", "5"]


def test_insert_section_without_an_anchor_needs_sections(tmp_path):
    d = docx.Document()
    d.add_paragraph("Unstructured prose with no numbering at all.")
    path = tmp_path / "flat.docx"
    d.save(str(path))
    rl = Redliner(path, author="T", date=DATE)
    report = apply_actions(rl, [{"id": "I", "type": "insert_section", "title": "Insurance"}])
    assert report.failed == 1
    assert "no numbered sections" in report.results[0].detail


# ---------------------------------------------------------------------------
# a delete removes what was asked for, and nothing else
# ---------------------------------------------------------------------------
def test_delete_refuses_to_swallow_a_clause_moved_into_it(rl_agreement):
    """Move X into section S, then delete S: X vanished silently, reporting
    `applied` on both actions."""
    report = apply_actions(
        rl_agreement,
        [
            {
                "id": "M",
                "type": "move_clause",
                "clause": "4.1",
                "into_section": "2",
                "position": "last",
            },
            {"id": "D", "type": "delete_section", "section": "2"},
        ],
    )
    assert report.results[0].status == "applied"
    assert report.results[1].status == "failed"
    assert "moved or inserted into it earlier" in report.results[1].detail
    surviving = [c.title for c in ClauseTree(accepted(rl_agreement).document.element.body)]
    assert "Governing Law" in surviving


def test_delete_refuses_to_swallow_a_clause_inserted_into_it(rl_agreement):
    report = apply_actions(
        rl_agreement,
        [
            {"id": "I", "type": "insert_clause", "into_section": "2", "text": "Newly added."},
            {"id": "D", "type": "delete_section", "section": "2"},
        ],
    )
    assert report.results[1].status == "failed"


def test_plain_delete_still_takes_its_own_subtree(rl_agreement):
    report = apply_actions(rl_agreement, [{"id": "D", "type": "delete_section", "section": "2"}])
    assert report.results[0].status == "applied"
    assert "+3 sub-clause(s)" in report.results[0].detail
    assert "Fees and Payment" not in "\n".join(texts(accepted(rl_agreement).document))


# ---------------------------------------------------------------------------
# reorder is positional, not label-keyed
# ---------------------------------------------------------------------------
@pytest.mark.parametrize("order", list(itertools.permutations(["2.1", "2.2", "2.3"])))
def test_every_permutation_lands_exactly(agreement_bytes, order):
    agreement_bytes.seek(0)
    rl = Redliner(agreement_bytes, author="T", date=DATE)
    before = {c.label: c.title for c in ClauseTree(rl.document.element.body).get("2").children}
    apply_actions(
        rl, [{"id": "R", "type": "reorder_clauses", "section": "2", "order": list(order)}]
    )
    assert titles_of(accepted(rl).document, "2") == [before[label] for label in order]


def test_reorder_is_unconfused_by_a_duplicate_label(rl_agreement):
    """With a label-keyed map, one clause moved twice and another never moved."""
    report = apply_actions(
        rl_agreement,
        [
            {"id": "I", "type": "insert_clause", "before_clause": "2.2", "text": "Inserted first."},
            {"id": "R", "type": "reorder_clauses", "section": "2", "order": ["2.3", "2.1", "2.2"]},
        ],
    )
    assert report.failed == 0
    tree = ClauseTree(accepted(rl_agreement).document.element.body)
    labels = [c.label for c in tree.get("2").children]
    assert labels == ["2.1", "2.2", "2.3", "2.4"]  # contiguous, nothing lost
    assert len({c.title for c in tree.get("2").children}) == 4  # nothing duplicated


# ---------------------------------------------------------------------------
# rationales anchor on text that survives
# ---------------------------------------------------------------------------
def test_rationale_prefers_a_visible_paragraph(rl_agreement):
    """`p.iter(w:r)` also matches struck runs, so a rationale could land on text
    the redline removes."""
    planner = ActionPlanner(rl_agreement, renumber=True, explain=True)
    planner.run(
        [
            {
                "id": "A",
                "type": "replace_text",
                "clause": "2.2",
                "find": "thirty (30) days",
                "replace": "forty-five (45) days",
                "rationale": "Longer terms.",
                "severity": "high",
            },
        ]
    )
    result = planner.report.results[0]
    anchor = ActionPlanner._live_anchor(result.anchors)
    from docx_redline.oxml.textmap import iter_visible_runs

    assert next(iter_visible_runs(anchor), None) is not None


def test_a_deleted_clause_still_gets_its_rationale(rl_agreement):
    """No visible runs remain, but the explanation still belongs on the strikeout."""
    report = apply_actions(
        rl_agreement,
        [
            {
                "id": "D",
                "type": "delete_clause",
                "clause": "2.2",
                "rationale": "Superseded by the Order Form.",
                "severity": "medium",
            }
        ],
        explain=True,
    )
    assert report.failed == 0
    assert not any("no paragraph left" in w for w in report.warnings)


# ---------------------------------------------------------------------------
# cross-references outside the body
# ---------------------------------------------------------------------------
def test_references_in_a_header_are_renumbered(tmp_path, agreement_bytes):
    d = docx.Document(agreement_bytes)
    header = d.sections[0].header
    header.is_linked_to_previous = False
    header.paragraphs[0].text = "Confidential — see Section 4.1 for governing law."
    path = tmp_path / "with_header.docx"
    d.save(str(path))

    rl = Redliner(path, author="T", date=DATE)
    report = apply_actions(
        rl, [{"id": "M", "type": "move_clause", "clause": "4.1", "after_clause": "3.1"}]
    )
    assert {(r["from"], r["to"]) for r in report.references} >= {("4.1", "3.2")}
    other = accepted(rl)
    assert "Section 3.2" in other.document.sections[0].header.paragraphs[0].text
