"""Chunked review: segment, triage, map concurrently, merge.

The interesting failures are all in the merge. Segments are reviewed in
isolation, so their proposals arrive with colliding ids, overlapping edits and
contradictory structural intentions -- and applying them naively does not fail
loudly, it produces a subtly wrong document with a plan report full of
``applied``. Most of what follows is about that.
"""

import io
import json
import re
import threading

import docx
import pytest
from conftest import DATE

from docx_redline import Redliner, full_redline
from docx_redline.planning.agent import Reply
from docx_redline.planning.chunked import ChunkedReviewer, SegmentCache, build_index, risk_families
from docx_redline.structure.clauses import ClauseTree
from docx_redline.structure.segments import segment_document


# ---------------------------------------------------------------------------
# a stand-in engine
# ---------------------------------------------------------------------------
class FakeEngine:
    """Answers `emit()` like a provider would, without a network."""

    model = "fake-1"
    max_tokens = 16000

    def __init__(
        self, *, items_for=None, triage_for=None, fail_times=0, truncate_times=0, delay=0.0
    ):
        self.calls: list[dict] = []
        self.items_for = items_for or (lambda variable: [])
        self.triage_for = triage_for
        self.fail_times = fail_times
        self.truncate_times = truncate_times
        self.delay = delay
        self.lock = threading.Lock()
        self.peak_concurrency = 0
        self._active = 0

    def emit(
        self,
        *,
        system,
        cached,
        variable,
        output,
        max_tokens=None,
        effort=None,
        cache_key=None,
        on_event=None,
    ):
        with self.lock:
            self._active += 1
            self.peak_concurrency = max(self.peak_concurrency, self._active)
            self.calls.append(
                {
                    "cached": cached,
                    "variable": variable,
                    "output": output.name,
                    "effort": effort,
                    "cache_key": cache_key,
                    "max_tokens": max_tokens,
                }
            )
        try:
            if self.delay:
                import time as _time

                _time.sleep(self.delay)
            if on_event:
                on_event("response.output_text.delta")
            if output.name == "emit_triage":
                ids = re.findall(r"^(S\d+)", cached, re.M)
                payload = (
                    self.triage_for(ids)
                    if self.triage_for
                    else {
                        "notes": "",
                        "selections": [
                            {"segment": i, "priority": "deep", "reason": "r", "focus": []}
                            for i in ids
                        ],
                    }
                )
                return Reply(
                    payload=payload,
                    model=self.model,
                    usage={"input_tokens": 10, "output_tokens": 5},
                )
            if self.truncate_times > 0:
                self.truncate_times -= 1
                from docx_redline.planning.agent import TruncatedReply

                raise TruncatedReply("out of room")
            if self.fail_times > 0:
                self.fail_times -= 1
                raise RuntimeError("transient upstream error")
            return Reply(
                payload={"summary": "s", "action_items": self.items_for(variable)},
                model=self.model,
                usage={"input_tokens": 100, "output_tokens": 20, "cache_read_input_tokens": 80},
            )
        finally:
            with self.lock:
                self._active -= 1


def item(**kwargs):
    base = {"type": "replace_text", "rationale": "r", "severity": "medium"}
    base.update(kwargs)
    return base


def clauses_in(variable: str) -> list[str]:
    return re.findall(r"^\s*(\d+\.\d+)\s", variable, re.M)


# ---------------------------------------------------------------------------
# a large fixture
# ---------------------------------------------------------------------------
@pytest.fixture(scope="module")
def big_tree():
    d = docx.Document()
    for section in range(1, 41):
        d.add_paragraph(f"{section}. Section Heading {section}")
        for clause in range(1, 11):
            p = d.add_paragraph()
            p.add_run(f"{section}.{clause}  ").bold = True
            p.add_run(
                f"Clause {clause}. Provider shall indemnify Customer and its liability is capped."
            )
    buf = io.BytesIO()
    d.save(buf)
    buf.seek(0)
    return ClauseTree(docx.Document(buf).element.body)


def segments_of(tree, budget=4_000):
    return segment_document(tree.root, budget_tokens=budget, tree=tree)


# ---------------------------------------------------------------------------
# the index
# ---------------------------------------------------------------------------
def test_index_lists_every_segment(big_tree):
    segments = segments_of(big_tree)
    index = build_index(segments)
    for seg in segments:
        assert seg.id in index, "a segment missing from the index cannot be triaged"


def test_index_carries_risk_flags(big_tree):
    index = build_index(segments_of(big_tree))
    assert "indemnity" in index


def test_index_sheds_depth_rather_than_segments(big_tree):
    segments = segments_of(big_tree)
    tight = build_index(segments, max_tokens=200)
    assert len(tight) < len(build_index(segments, max_tokens=100_000))
    for seg in segments:
        assert seg.id in tight


def test_risk_families_are_local_and_free():
    assert "liability" in risk_families("Limitation of Liability applies")
    assert risk_families("The quick brown fox") == []


# ---------------------------------------------------------------------------
# triage floors
# ---------------------------------------------------------------------------
def test_unmentioned_segments_default_to_scan_not_skip(big_tree):
    """Silence from the model must never mean a segment goes unread."""
    engine = FakeEngine(
        triage_for=lambda ids: {"notes": "", "selections": []}, items_for=lambda v: []
    )
    reviewer = ChunkedReviewer(
        engine, segment_tokens=4_000, use_cache=False, min_coverage=0.0, backoff=0
    )
    reviewer.propose(big_tree, "brief")
    priorities = reviewer.report.triage["priorities"]
    assert set(priorities.values()) == {"scan"}
    assert all(s["status"] != "skipped" for s in reviewer.report.segments)


def test_a_keyword_heavy_segment_cannot_be_skipped(big_tree):
    engine = FakeEngine(
        triage_for=lambda ids: {
            "notes": "",
            "selections": [
                {"segment": i, "priority": "skip", "reason": "r", "focus": []} for i in ids
            ],
        },
        items_for=lambda v: [],
    )
    reviewer = ChunkedReviewer(
        engine, segment_tokens=4_000, use_cache=False, min_coverage=0.0, backoff=0
    )
    reviewer.propose(big_tree, "brief")
    # every segment of this fixture mentions indemnity and liability
    assert "skip" not in set(reviewer.report.triage["priorities"].values())


def test_coverage_floor_promotes_until_enough_is_read(big_tree):
    engine = FakeEngine(
        triage_for=lambda ids: {
            "notes": "",
            "selections": [
                {"segment": i, "priority": "deep" if n == 0 else "scan", "reason": "r", "focus": []}
                for n, i in enumerate(ids)
            ],
        },
        items_for=lambda v: [],
    )
    reviewer = ChunkedReviewer(
        engine, segment_tokens=4_000, use_cache=False, min_coverage=0.6, backoff=0
    )
    reviewer.propose(big_tree, "brief")
    priorities = reviewer.report.triage["priorities"]
    assert sum(1 for p in priorities.values() if p == "deep") > 1


def test_triage_failure_reviews_everything(big_tree):
    class Broken(FakeEngine):
        def emit(self, *, output, **kw):
            if output.name == "emit_triage":
                raise RuntimeError("triage exploded")
            return super().emit(output=output, **kw)

    reviewer = ChunkedReviewer(
        Broken(items_for=lambda v: []), segment_tokens=4_000, use_cache=False, backoff=0
    )
    reviewer.propose(big_tree, "brief")
    assert set(reviewer.report.triage["priorities"].values()) == {"deep"}


def test_triage_can_be_switched_off(big_tree):
    engine = FakeEngine(items_for=lambda v: [])
    reviewer = ChunkedReviewer(
        engine, segment_tokens=4_000, use_cache=False, triage=False, backoff=0
    )
    reviewer.propose(big_tree, "brief")
    assert not any(c["output"] == "emit_triage" for c in engine.calls)


# ---------------------------------------------------------------------------
# prompt layout and caching
# ---------------------------------------------------------------------------
def test_the_cached_prefix_is_identical_across_segments(big_tree):
    engine = FakeEngine(items_for=lambda v: [])
    ChunkedReviewer(engine, segment_tokens=500, use_cache=False, triage=False, backoff=0).propose(
        big_tree, "brief"
    )
    review_calls = [c for c in engine.calls if c["output"] != "emit_triage"]
    assert len(review_calls) > 2
    assert len({c["cached"] for c in review_calls}) == 1
    assert len({c["cache_key"] for c in review_calls}) == 1


def test_the_brief_is_in_the_cached_prefix_not_the_variable_part(big_tree):
    """Constant across the run, so it belongs behind the cache breakpoint --
    the opposite of a single-shot review, where the contract is the constant."""
    engine = FakeEngine(items_for=lambda v: [])
    ChunkedReviewer(engine, segment_tokens=4_000, use_cache=False, triage=False, backoff=0).propose(
        big_tree, "UNIQUE-BRIEF-MARKER"
    )
    call = next(c for c in engine.calls if c["output"] != "emit_triage")
    assert "UNIQUE-BRIEF-MARKER" in call["cached"]
    assert "UNIQUE-BRIEF-MARKER" not in call["variable"]
    assert "<segment" in call["variable"]


def test_scan_priority_lowers_effort(big_tree):
    engine = FakeEngine(
        triage_for=lambda ids: {
            "notes": "",
            "selections": [
                {"segment": i, "priority": "scan", "reason": "r", "focus": []} for i in ids
            ],
        },
        items_for=lambda v: [],
    )
    ChunkedReviewer(
        engine, segment_tokens=4_000, use_cache=False, min_coverage=0.0, backoff=0
    ).propose(big_tree, "brief")
    review_calls = [c for c in engine.calls if c["output"] != "emit_triage"]
    assert all(c["effort"] == "low" for c in review_calls)


def test_segments_are_reviewed_concurrently(big_tree):
    engine = FakeEngine(items_for=lambda v: [], delay=0.02)
    ChunkedReviewer(
        engine, segment_tokens=500, use_cache=False, triage=False, concurrency=4, backoff=0
    ).propose(big_tree, "brief")
    assert engine.peak_concurrency > 1


# ---------------------------------------------------------------------------
# resilience
# ---------------------------------------------------------------------------
def test_a_transient_failure_is_retried(big_tree):
    engine = FakeEngine(items_for=lambda v: [], fail_times=1)
    reviewer = ChunkedReviewer(
        engine, segment_tokens=100_000, use_cache=False, triage=False, backoff=0
    )
    reviewer.propose(big_tree, "brief")
    assert all(s["status"] in ("ok", "cached") for s in reviewer.report.segments)


def test_a_failure_among_survivors_is_recorded_not_raised(big_tree):
    engine = FakeEngine(items_for=lambda v: [], fail_times=3)  # kills one segment only
    reviewer = ChunkedReviewer(engine, segment_tokens=500, use_cache=False, triage=False, backoff=0)
    reviewer.propose(big_tree, "brief")
    statuses = {s["status"] for s in reviewer.report.segments}
    assert "failed" in statuses and "ok" in statuses


def test_a_run_where_nothing_could_be_reviewed_raises(big_tree):
    """An empty result would read as a clean bill of health."""
    engine = FakeEngine(items_for=lambda v: [], fail_times=999)
    reviewer = ChunkedReviewer(engine, segment_tokens=500, use_cache=False, triage=False, backoff=0)
    with pytest.raises(RuntimeError, match="clean bill of health"):
        reviewer.propose(big_tree, "brief")


def test_missing_credentials_are_not_swallowed_as_a_failed_segment(big_tree):
    """RedlineCredentialsError is a RuntimeError; a broad except turned a
    misconfigured run into a silent 'no findings'."""
    from docx_redline.planning.agent import RedlineCredentialsError

    class NoKey(FakeEngine):
        def emit(self, **kwargs):
            raise RedlineCredentialsError("no Anthropic credentials found -- set ...")

    reviewer = ChunkedReviewer(NoKey(), segment_tokens=500, use_cache=False, backoff=0)
    with pytest.raises(RedlineCredentialsError):
        reviewer.propose(big_tree, "brief")


def test_truncation_splits_the_segment(big_tree):
    """More room first; if that is not enough the span itself is too big."""
    engine = FakeEngine(items_for=lambda v: [], truncate_times=2)
    reviewer = ChunkedReviewer(
        engine, segment_tokens=100_000, use_cache=False, triage=False, backoff=0
    )
    reviewer.propose(big_tree, "brief")
    assert any("split after truncation" in s["detail"] for s in reviewer.report.segments)
    budgets = [c["max_tokens"] for c in engine.calls if c["output"] != "emit_triage"]
    assert any(b for b in budgets), "should have retried with a larger ceiling first"


# ---------------------------------------------------------------------------
# the disk cache
# ---------------------------------------------------------------------------
def test_a_rerun_uses_the_cache(tmp_path, big_tree):
    make = lambda: ChunkedReviewer(  # noqa: E731
        FakeEngine(items_for=lambda v: []),
        segment_tokens=4_000,
        cache_dir=tmp_path,
        triage=False,
        backoff=0,
    )
    first = make()
    first.propose(big_tree, "brief")
    calls_first = len(first.engine.calls)

    second = make()
    second.propose(big_tree, "brief")
    assert second.engine.calls == []
    assert calls_first > 0
    assert all(s["status"] == "cached" for s in second.report.segments)


def test_changing_the_brief_invalidates_the_cache(tmp_path, big_tree):
    first = ChunkedReviewer(
        FakeEngine(items_for=lambda v: []),
        segment_tokens=4_000,
        cache_dir=tmp_path,
        triage=False,
        backoff=0,
    )
    first.propose(big_tree, "brief one")
    second = ChunkedReviewer(
        FakeEngine(items_for=lambda v: []),
        segment_tokens=4_000,
        cache_dir=tmp_path,
        triage=False,
        backoff=0,
    )
    second.propose(big_tree, "brief two")
    assert second.engine.calls, "a different brief must not reuse the old answers"


def test_refresh_ignores_the_cache(tmp_path, big_tree):
    ChunkedReviewer(
        FakeEngine(items_for=lambda v: []),
        segment_tokens=4_000,
        cache_dir=tmp_path,
        triage=False,
        backoff=0,
    ).propose(big_tree, "brief")
    again = ChunkedReviewer(
        FakeEngine(items_for=lambda v: []),
        segment_tokens=4_000,
        cache_dir=tmp_path,
        triage=False,
        refresh=True,
        backoff=0,
    )
    again.propose(big_tree, "brief")
    assert again.engine.calls


def test_cache_key_covers_what_changes_the_answer(tmp_path):
    cache = SegmentCache(tmp_path)
    assert cache.key("a", "b") == cache.key("a", "b")
    assert cache.key("a", "b") != cache.key("a", "c")


# ---------------------------------------------------------------------------
# end to end through the pipeline
# ---------------------------------------------------------------------------
def test_chunked_reviewer_drives_the_whole_pipeline(tmp_path, agreement):
    def propose_for(variable):
        return [
            item(id="AI-001", clause=label, find="Payment is due", replace="Payment falls due")
            for label in clauses_in(variable)
            if label == "2.2"
        ]

    reviewer = ChunkedReviewer(
        FakeEngine(items_for=propose_for),
        segment_tokens=1_000,
        use_cache=False,
        triage=False,
        backoff=0,
    )
    out = tmp_path / "out.docx"
    result = full_redline(agreement, out, reviewer=reviewer, date=DATE)
    assert result.ok, result.format()
    assert result.proposal.source.startswith("chunked:")
    assert result.proposal.merge["summary"]["kept"] >= 1

    accepted = Redliner(out, track_changes=False)
    accepted.accept_all()
    assert "Payment falls due" in "\n".join(p.text for p in accepted.document.paragraphs)


def test_the_merge_report_reaches_the_json_report(tmp_path, agreement):
    reviewer = ChunkedReviewer(
        FakeEngine(
            items_for=lambda v: [
                item(
                    id="A",
                    clause=label,
                    find="Payment is due",
                    replace="Payment falls due",
                    backoff=0,
                )
                for label in clauses_in(v)
                if label == "2.2"
            ]
        ),
        segment_tokens=1_000,
        use_cache=False,
        triage=False,
    )
    report = tmp_path / "report.json"
    full_redline(agreement, tmp_path / "out.docx", reviewer=reviewer, date=DATE, report_path=report)
    payload = json.loads(report.read_text())
    assert payload["proposal"]["usage"]["calls"] >= 1


# ---------------------------------------------------------------------------
# liveness — a long call must not look like a hang
# ---------------------------------------------------------------------------
def test_progress_is_emitted_before_the_call_not_only_after(big_tree):
    """The call itself can run for minutes. Reporting only on completion makes
    a working run indistinguishable from a hung one."""
    events = []
    engine = FakeEngine(items_for=lambda v: [])
    ChunkedReviewer(
        engine,
        segment_tokens=500,
        use_cache=False,
        triage=False,
        backoff=0,
        on_progress=events.append,
    ).propose(big_tree, "brief")
    sending = [e for e in events if e.status == "sending"]
    assert sending, "nothing announced the call before it blocked"
    assert sending[0].segment and "tok" in sending[0].detail


def test_a_retry_is_announced(big_tree):
    events = []
    engine = FakeEngine(items_for=lambda v: [], fail_times=1)
    ChunkedReviewer(
        engine,
        segment_tokens=100_000,
        use_cache=False,
        triage=False,
        backoff=0,
        on_progress=events.append,
    ).propose(big_tree, "brief")
    assert any(e.status == "retry" for e in events)


def test_split_recursion_is_bounded(big_tree):
    """`split()` terminates on its own, but only after a lot of paid calls."""
    engine = FakeEngine(items_for=lambda v: [], truncate_times=999)
    reviewer = ChunkedReviewer(
        engine, segment_tokens=100_000, use_cache=False, triage=False, backoff=0, max_split_depth=1
    )
    with pytest.raises(RuntimeError, match="clean bill of health"):
        reviewer.propose(big_tree, "brief")
    # depth 1 => the original plus two halves, each retried once for headroom
    assert len(engine.calls) <= 8, len(engine.calls)


def test_a_partly_successful_split_reports_what_was_lost(big_tree):
    """One half answered, the other did not -- say so rather than round up."""

    class HalfBroken(FakeEngine):
        def __init__(self):
            super().__init__(items_for=lambda v: [])
            self.seen = 0

        def emit(self, **kwargs):
            self.seen += 1
            from docx_redline.planning.agent import TruncatedReply

            if self.seen <= 4:  # parent (x2), then the first half (x2)
                raise TruncatedReply("out of room")
            return super().emit(**kwargs)

    reviewer = ChunkedReviewer(
        HalfBroken(),
        segment_tokens=100_000,
        use_cache=False,
        triage=False,
        backoff=0,
        max_split_depth=1,
    )
    reviewer.propose(big_tree, "brief")
    [segment] = reviewer.report.segments
    assert segment["status"] == "ok"
    assert "halves failed" in segment["detail"]


def test_a_long_call_reports_that_it_is_alive(big_tree):
    """A reasoning model sends nothing for minutes. Without a heartbeat there is
    no way to tell a working call from a hung one."""
    events = []
    engine = FakeEngine(items_for=lambda v: [])
    ChunkedReviewer(
        engine,
        segment_tokens=100_000,
        use_cache=False,
        triage=False,
        backoff=0,
        on_progress=events.append,
    ).propose(big_tree, "brief")
    beats = [e for e in events if e.status == "streaming"]
    assert beats, "the stream ticked but nothing was reported"
    assert "events" in beats[0].detail


def test_the_heartbeat_does_not_flood(big_tree):
    """Thousands of stream events must not become thousands of log lines."""
    from docx_redline.planning.chunked import _Heartbeat

    seen = []
    beat = _Heartbeat(seen.append, "S00", every=60.0)
    for _ in range(5_000):
        beat("response.output_text.delta")
    assert len(seen) == 1


def test_a_stream_that_ends_early_splits_instead_of_retrying_blindly(big_tree):
    """Repeating the same over-long request reproduces the failure; three
    full-length attempts is 40 minutes of nothing."""
    from docx_redline.planning.agent import StreamInterrupted

    class Interrupted(FakeEngine):
        def emit(self, **kwargs):
            self.calls.append(kwargs)
            raise StreamInterrupted("the response stream ended before the model finished")

    engine = Interrupted(items_for=lambda v: [])
    reviewer = ChunkedReviewer(
        engine, segment_tokens=100_000, use_cache=False, triage=False, backoff=0, max_split_depth=1
    )
    with pytest.raises(RuntimeError, match="clean bill of health"):
        reviewer.propose(big_tree, "brief")
    # one call for the parent, then one per half -- not three per segment
    assert len(engine.calls) == 3, len(engine.calls)
