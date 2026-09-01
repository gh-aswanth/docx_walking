"""Whole-document visibility: what the reviewer is actually shown.

`ClauseTree` answers "what is numbered here". These tests cover the other
question -- what the document says -- because the reviewer was previously sent
only the numbered clauses, truncated, with every exhibit and table dropped.
"""

import io

import docx
import pytest
from conftest import DATE, SAMPLE

from docx_redline import Redliner, full_redline
from docx_redline.planning.actions import allowed_types, apply_actions
from docx_redline.structure.clauses import ClauseTree, render_outline
from docx_redline.structure.segments import (
    count_headings,
    detect_strategy,
    iter_blocks,
    render_document,
    segment_document,
)


def body_of(document) -> tuple:
    body = document.element.body
    return body, ClauseTree(body)


def saved(build) -> docx.Document:
    d = docx.Document()
    build(d)
    buf = io.BytesIO()
    d.save(buf)
    buf.seek(0)
    return docx.Document(buf)


# ---------------------------------------------------------------------------
# fixtures for the four document shapes
# ---------------------------------------------------------------------------
def numbered_doc(d):
    for section in range(1, 5):
        d.add_paragraph(f"{section}. Heading {section}")
        for clause in range(1, 5):
            d.add_paragraph(f"{section}.{clause}  Clause {clause}. Body text that runs on.")


def heading_doc(d):
    for section in range(1, 5):
        d.add_heading(f"Heading {section}", level=1)
        for clause in range(1, 5):
            d.add_heading(f"Sub {clause}", level=2)
            d.add_paragraph("Clause body text that runs on for a while.")


def numpr_doc(d):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as Q

    def number(paragraph, ilvl):
        ppr = paragraph._p.get_or_add_pPr()
        numpr = OxmlElement("w:numPr")
        level = OxmlElement("w:ilvl")
        level.set(Q("w:val"), str(ilvl))
        numpr.append(level)
        num_id = OxmlElement("w:numId")
        num_id.set(Q("w:val"), "1")
        numpr.append(num_id)
        ppr.append(numpr)

    for section in range(1, 5):
        number(d.add_paragraph(f"Obligation {section}"), 0)
        for clause in range(1, 5):
            number(d.add_paragraph(f"Sub-obligation {clause} with body text."), 1)


def flat_doc(d):
    for index in range(20):
        d.add_paragraph(f"Unstructured prose paragraph {index} with some content in it.")


SHAPES = [
    ("numbered", numbered_doc, "clauses"),
    ("headings", heading_doc, "headings"),
    ("auto-numbered", numpr_doc, "headings"),
    ("flat", flat_doc, "windows"),
]


# ---------------------------------------------------------------------------
# every shape is visible
# ---------------------------------------------------------------------------
@pytest.mark.parametrize("name,build,expected", SHAPES)
def test_strategy_detection(name, build, expected):
    body, tree = body_of(saved(build))
    assert detect_strategy(body, tree) == expected


@pytest.mark.parametrize("name,build,expected", SHAPES)
def test_every_shape_renders_something(name, build, expected):
    """Three of four shapes previously rendered an empty string."""
    body, tree = body_of(saved(build))
    rendered = render_document(body, tree)
    assert rendered.strip()
    assert len(rendered) > 400


@pytest.mark.parametrize("name,build,expected", SHAPES)
def test_every_block_belongs_to_exactly_one_segment(name, build, expected):
    body, tree = body_of(saved(build))
    blocks = iter_blocks(body, tree)
    segments = segment_document(body, budget_tokens=200, tree=tree)
    owned = [id(b) for seg in segments for b in seg.blocks]
    assert len(owned) == len(set(owned)) == len(blocks)


def test_heading_levels_survive_lxml_proxy_churn():
    """An id()-keyed level map silently lost entries as proxies were freed."""
    body, _ = body_of(saved(numpr_doc))
    assert count_headings(body) == 20


def test_structural_actions_are_withheld_without_clause_numbers():
    body, tree = body_of(saved(flat_doc))
    [segment] = segment_document(body, budget_tokens=10_000, tree=tree)[:1]
    allowed = allowed_types(segment)
    assert "replace_text" in allowed
    assert "move_clause" not in allowed


def test_segments_respect_the_budget():
    body, tree = body_of(saved(numbered_doc))
    segments = segment_document(body, budget_tokens=60, tree=tree)
    assert len(segments) > 1
    # a unit is indivisible, so only a single oversized unit may exceed budget
    assert sum(1 for s in segments if s.approx_tokens > 60) <= 1


def test_segment_labels_cover_their_clauses():
    body, tree = body_of(saved(numbered_doc))
    segments = segment_document(body, budget_tokens=100, tree=tree)
    owned = set().union(*(s.labels for s in segments))
    assert owned == {c.label for c in tree}


# ---------------------------------------------------------------------------
# the real contract: what was previously dropped
# ---------------------------------------------------------------------------
@pytest.fixture
def sample_render():
    document = docx.Document(SAMPLE)
    body, tree = body_of(document)
    return render_document(body, tree)


@pytest.mark.parametrize(
    "content",
    [
        "SOFTWARE LICENSE AND SUBSCRIPTION SERVICES AGREEMENT",  # title
        "Client Co., Inc., a company organized under",  # recitals
        "IN WITNESS WHEREOF",  # signature block
        "Exhibit A — Order Form Summary",
        "Annual Fees: $186,000, payable annually in advance",  # the commercial terms
        "Exhibit B — Service Level Agreement",
        "99.0%–99.89% uptime: 5% service credit",  # the SLA table
        "Exhibit C — Data Processing Addendum",
    ],
)
def test_previously_invisible_content_is_rendered(sample_render, content):
    assert content in sample_render


def test_tables_are_rendered_with_addressable_indices(sample_render):
    assert "[table 0]" in sample_render and "[table 1]" in sample_render
    assert "row 0: Signature | Date" in sample_render


def test_clause_text_is_no_longer_truncated(sample_render):
    """render_outline caps every clause at 400 chars; 2.2 runs to 441."""
    assert "remove or obscure any proprietary notices" in sample_render
    assert "remove or obscure any proprietary notices" not in render_outline(
        ClauseTree(docx.Document(SAMPLE).element.body)
    )


def test_render_outline_still_summarises(sample_render):
    """The summary view is kept -- it is what the compact index will use."""
    tree = ClauseTree(docx.Document(SAMPLE).element.body)
    assert len(render_outline(tree)) < len(sample_render)
    subset = [c for c in tree if c.label.startswith("3.")]
    assert render_outline(tree, clauses=subset).strip().startswith("3.1")


# ---------------------------------------------------------------------------
# exhibit content is now actionable end to end
# ---------------------------------------------------------------------------
def test_an_exhibit_can_be_edited(tmp_path):
    out = tmp_path / "out.docx"
    result = full_redline(
        SAMPLE,
        out,
        date=DATE,
        actions=[
            {
                "id": "X",
                "type": "replace_text",
                "find": "Annual Fees: $186,000",
                "replace": "Annual Fees: $150,000",
                "rationale": "Negotiated fee reduction.",
                "severity": "high",
            }
        ],
    )
    assert result.ok, result.format()
    accepted = Redliner(out, track_changes=False)
    accepted.accept_all()
    assert "Annual Fees: $150,000" in "\n".join(p.text for p in accepted.document.paragraphs)


# ---------------------------------------------------------------------------
# ambiguity is an error, not a coin toss
# ---------------------------------------------------------------------------
def test_unscoped_edit_matching_many_paragraphs_is_refused(rl_agreement):
    report = apply_actions(
        rl_agreement,
        [{"id": "A", "type": "replace_text", "find": "Order Form", "replace": "Order"}],
    )
    assert report.failed == 1
    assert "appears in 2 paragraphs" in report.results[0].detail
    assert len(rl_agreement.summary()) == 0


def test_unscoped_edit_with_a_unique_quote_applies(rl_agreement):
    report = apply_actions(
        rl_agreement,
        [
            {
                "id": "A",
                "type": "replace_text",
                "find": "Delaware law governs",
                "replace": "New York law governs",
            }
        ],
    )
    assert report.failed == 0


def test_scoping_by_clause_still_takes_the_first_match(rl_agreement):
    report = apply_actions(
        rl_agreement,
        [
            {
                "id": "A",
                "type": "replace_text",
                "clause": "2.2",
                "find": "thirty (30) days",
                "replace": "forty-five (45) days",
            }
        ],
    )
    assert report.failed == 0


def test_all_true_bypasses_the_uniqueness_check(rl_agreement):
    report = apply_actions(
        rl_agreement,
        [
            {
                "id": "A",
                "type": "replace_text",
                "find": "Order Form",
                "replace": "Order",
                "all": True,
            }
        ],
    )
    assert report.failed == 0
    assert report.results[0].edits == 2


def test_insert_text_without_clause_or_anchor_is_refused(rl_agreement):
    """It appended the model's text to the document title."""
    report = apply_actions(
        rl_agreement, [{"id": "A", "type": "insert_text", "text": " (as amended)"}]
    )
    assert report.failed == 1
    assert "needs an anchor" in report.results[0].detail
